# AIBOT.py
# UniVentureAI — Telegram bot with GLOBAL per-topic RAG + metadata separation (qa/evaluation)
# + eval follow-ups (apply feedback) + eval Q&A (any follow-up question about the evaluated text)
# + embedded tools + Application Plan & School Finder
# + analytics + admin locks + backup + health + robust command parsing + UUID doc IDs (no reteach bugs)

import os
os.environ['TZ'] = 'UTC'  # Set timezone to UTC

from datetime import datetime, timedelta
from telegram import Update, KeyboardButton, ReplyKeyboardMarkup
from telegram.constants import ChatAction
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    ApplicationHandlerStop,
    filters,
)
from dotenv import load_dotenv
import chromadb
from chromadb.utils import embedding_functions
import os, io, nest_asyncio, logging, json, base64, uuid, re

# -------- File extraction deps --------
from pdfminer.high_level import extract_text
from docx import Document as DocxDocument

# -------- Web page extraction (for teachlink) --------
import trafilatura

import threading
from http.server import HTTPServer, BaseHTTPRequestHandler

import asyncio

try:
    from openai import RateLimitError
except ImportError:
    from openai.error import RateLimitError

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
load_dotenv()
nest_asyncio.apply()

TELEGRAM_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# -------- Model routing (speed vs depth) --------
FAST_MODEL = os.getenv("OPENAI_FAST_MODEL", "gpt-4.1-mini")
STRONG_MODEL = os.getenv("OPENAI_STRONG_MODEL", "gpt-4.1")

# Evaluation speed UX: send a short "quick feedback" first, then full feedback.
ENABLE_EVAL_QUICK_PREVIEW = os.getenv("ENABLE_EVAL_QUICK_PREVIEW", "1") == "1"
EVAL_QUICK_MAX_TOKENS = int(os.getenv("EVAL_QUICK_MAX_TOKENS", "280"))
if not TELEGRAM_TOKEN:
    raise RuntimeError("Missing TELEGRAM_BOT_TOKEN in environment.")
if not OPENAI_API_KEY:
    raise RuntimeError("Missing OPENAI_API_KEY in environment.")

# =========================
# OpenAI SDK compatibility
# =========================
USE_NEW_OPENAI = False
_client = None
try:
    # openai>=1.x
    from openai import OpenAI  # type: ignore

    _client = OpenAI(api_key=OPENAI_API_KEY)
    USE_NEW_OPENAI = True
    logging.info("Using OpenAI SDK v1.x+")
except Exception as e:
    # openai<=0.28.x
    import openai  # type: ignore

    openai.api_key = OPENAI_API_KEY
    USE_NEW_OPENAI = False
    logging.info("Using OpenAI SDK v0.28.x or earlier")

async def openai_chat(
    model: str,
    messages: list,
    temperature: float = 0.4,
    max_tokens: int | None = None,
) -> str:
    """Async OpenAI chat completion with retry + graceful wait message."""

    max_retries = 3
    attempt = 0

    while attempt < max_retries:
        try:
            if USE_NEW_OPENAI:
                resp = await asyncio.wait_for(
                    _client.chat.completions.create(
                        model=model,
                        messages=messages,
                        temperature=temperature,
                        max_tokens=max_tokens,
                    ),
                    timeout=60
                )
                return (resp.choices[0].message.content or "").strip()
            else:
                import openai as _openai
                resp = await _openai.ChatCompletion.acreate(
                    model=model,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=max_tokens,
                )
                return (resp["choices"][0]["message"]["content"] or "").strip()

        except (RateLimitError, asyncio.TimeoutError):
            attempt += 1
            logging.warning(f"OpenAI retry (attempt {attempt}/{max_retries})")

            if attempt < max_retries:
                await asyncio.sleep(10)
                continue
            else:
                return "⚠️ I'm a bit busy right now! Please try again in 60 seconds."

        except Exception:
            logging.exception("OpenAI unexpected error")
            return "⚠️ I'm a bit busy right now! Please try again in 60 seconds."

# -------- Admin config --------
ADMIN_IDS = {
    886181760,  # TODO: replace with YOUR Telegram user ID (from @userinfobot)
}

def require_admin(update: Update) -> bool:
    user = update.effective_user
    return bool(user and user.id in ADMIN_IDS)

# -------- Data / storage config (for Railway/Render volume etc.) --------
DATA_DIR = os.getenv("DATA_DIR", "./data")
os.makedirs(DATA_DIR, exist_ok=True)

# =============================================================================
# Paid access (Manual verification — Option 3)
# - 30-day subscription with auto-expiry
# - Admin activation (/activate <user_id> [days]) and deactivation
# - User check expiry (/mysub)
# - Automatic block for unpaid/expired users
# - Auto reminders (3 days before expiry + after expiry)
# - Payment proof: users send screenshot to bot; bot forwards to ADMIN and replies
# =============================================================================

PAID_DB_PATH = os.getenv("PAID_DB_PATH", os.path.join(DATA_DIR, "paid_users.json"))
DEFAULT_SUB_DAYS = int(os.getenv("DEFAULT_SUB_DAYS", "30"))

# Payment instructions (set these in Railway Variables)
PAYMENT_PRICE_USD = os.getenv("PAYMENT_PRICE_USD", "$29")
PAYMENT_CARD = os.getenv("PAYMENT_CARD", "")        # e.g. "8600 1234 5678 9012"
PAYMENT_CLICK = os.getenv("PAYMENT_CLICK", "")      # e.g. "+998901234567"
PAYMENT_PAYME = os.getenv("PAYMENT_PAYME", "")      # e.g. "+998901234567"
PAYMENT_NOTE = os.getenv("PAYMENT_NOTE", "")        # optional extra line

_paid_lock = threading.Lock()

def _paid_load() -> dict:
    try:
        if not os.path.exists(PAID_DB_PATH):
            return {}
        with _paid_lock:
            with open(PAID_DB_PATH, "r", encoding="utf-8") as f:
                return json.load(f) or {}
    except Exception as e:
        logging.error(f"Failed to load paid DB: {e}")
        return {}

def _paid_save(db: dict) -> None:
    try:
        os.makedirs(os.path.dirname(PAID_DB_PATH), exist_ok=True)
        with _paid_lock:
            with open(PAID_DB_PATH, "w", encoding="utf-8") as f:
                json.dump(db, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.error(f"Failed to save paid DB: {e}")

def _parse_iso(dt_str: str) -> datetime | None:
    try:
        return datetime.fromisoformat(dt_str)
    except Exception:
        return None

def get_paid_record(user_id: int) -> dict | None:
    db = _paid_load()
    return db.get(str(user_id))

def is_paid_user(user_id: int) -> bool:
    rec = get_paid_record(user_id)
    if not rec:
        return False
    exp = _parse_iso(rec.get("expires_at", ""))
    if not exp:
        return False
    return datetime.utcnow() <= exp

def remaining_days(user_id: int) -> int | None:
    rec = get_paid_record(user_id)
    if not rec:
        return None
    exp = _parse_iso(rec.get("expires_at", ""))
    if not exp:
        return None
    return (exp - datetime.utcnow()).days

def activate_paid(user_id: int,
                  days: int = DEFAULT_SUB_DAYS,
                  activated_by: int | None = None,
                  username: str | None = None,
                  first_name: str | None = None) -> None:

    db = _paid_load()
    now = datetime.utcnow()
    exp = now + timedelta(days=int(days))

    db[str(user_id)] = {
        "user_id": user_id,
        "username": username,
        "first_name": first_name,
        "activated_at": now.isoformat(),
        "expires_at": exp.isoformat(),
        "activated_by": activated_by,
        "reminded_3day": False,
        "expired_notified": False,
    }

    _paid_save(db)

def deactivate_paid(user_id: int) -> bool:
    db = _paid_load()
    if str(user_id) in db:
        db.pop(str(user_id), None)
        _paid_save(db)
        return True
    return False

def _payment_instructions_text(user_id: int) -> str:
    parts = [
        "🔒 <b>Paid access required</b>",
        "",
        "💳 <b>Limited Offer</b>",
        "Old price: <s>$75</s>",
        f"Now: <b>$29</b> (valid for {DEFAULT_SUB_DAYS} days)",
        "",
        "Pay using any of these methods and then send a screenshot here:",
    ]

    if PAYMENT_CLICK:
        parts.append(f"• Click: {PAYMENT_CLICK}")
    if PAYMENT_PAYME:
        parts.append(f"• Payme: {PAYMENT_PAYME}")
    if PAYMENT_CARD:
        parts.append(f"• Card (HUMO):\n {PAYMENT_CARD}\n SAIDAMIRKHON YUSUPOV")
    if PAYMENT_NOTE:
        parts.append(PAYMENT_NOTE)

    if not (PAYMENT_CLICK or PAYMENT_PAYME or PAYMENT_CARD):
        parts.append("• Contact admin for payment details.")

    parts.extend([
        "",
        f"🆔 Your user ID: <code>{user_id}</code>",
        "After you send the screenshot, an admin will verify and activate your access.",
        "",
        "<b>Commands:</b>",
        "• /pay — show payment details",
        "• /id — show your user ID",
        "• /mysub — check subscription status",
    ])

    return "\n".join(parts)

async def how_to_use_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "🤖 *UniVentureAI - Help Guide*\n\n"
        "I help you build a stronger university application step by step.\n\n"
        "━━━━━━━━━━━━━━━\n"
        "📚 MAIN SECTIONS\n"
        "━━━━━━━━━━━━━━━\n"
        "📝 Essays —> Personal Statement & Supplementals\n"
        "🎯 Extracurriculars —> Structure & impact\n"
        "✉️ Recommendation Letters —> Strategy & evaluation\n"
        "📈 SAT / 🗣 IELTS —> Prep & Writing feedback\n"
        "🖼 Portfolio —> Structure & project ideas\n"
        "📅 Application Plan —> Timeline & strategy\n"
        "🏫 School Finder —> Reach / Match / Safety suggestions\n\n"
        "━━━━━━━━━━━━━━━\n"
        "⚡ BOOST TOOLS\n"
        "━━━━━━━━━━━━━━━\n"
        "📊 My Progress —> Track readiness\n"
        "🤫 Insider Tips —> Section-specific secrets\n"
        "⚡ Power Words —> Upgrade vocabulary\n"
        "🎯 Predict My Chances —> Readiness indicator\n"
        "🔍 Find Wow Factor —> Detect your strongest hook\n\n"
        "━━━━━━━━━━━━━━━\n"
        "🧠 HOW TO USE ME\n"
        "━━━━━━━━━━━━━━━\n"
        "1️⃣ Choose a section from the menu\n"
        "2️⃣ Tap the Evaluation button\n"
        "3️⃣ Upload your draft (text / PDF / image)\n"
        "4️⃣ Ask follow-up questions like:\n"
        "   • \"Rewrite the ending\"\n"
        "   • \"Make transitions smoother\"\n"
        "   • \"Fix grammar\"\n\n"
        "━━━━━━━━━━━━━━━\n"
        "💎 SUBSCRIPTION\n"
        "━━━━━━━━━━━━━━━\n"
        "This platform requires paid access.\n"
        "Use /pay to unlock full features.\n"
        "Use /mysub to check your subscription.\n\n"
        "Need help? Just type your question.\n"
        "I'll guide you step by step 🚀",
        parse_mode="Markdown"
    )

async def help_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "🤖 UniVentureAI Help\n\n"
        "Main commands:\n"
        "/start - Open main menu\n"
        "/profile - View your saved profile\n"
        "/pay - Payment instructions\n"
        "/status - Check subscription status\n"
        "/feedback - Send feedback\n\n"
        "Or simply use the menu buttons to navigate."
    )
    
async def profile_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)

    profile = mem.get("profile", {}) or {}

    msg = (
        "👤 Your Profile\n\n"
        f"Grade: {profile.get('grade') or '—'}\n"
        f"Country: {profile.get('country') or '—'}\n"
        f"Intended Major: {profile.get('major') or '—'}\n"
        f"GPA: {profile.get('gpa') or '—'}\n"
        f"SAT: {profile.get('sat') or '—'}\n"
        f"IELTS: {profile.get('ielts') or '—'}\n"
        f"Needs Financial Aid: {profile.get('needs_aid') if profile.get('needs_aid') is not None else '—'}\n"
    )

    await update.message.reply_text(msg)

async def feedback_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()

    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        await update.message.reply_text(
            "💬 Send feedback like this:\n\n"
            "/feedback Your message here"
        )
        return

    feedback_text = parts[1]

    # Optional: forward to admin
    for admin_id in ADMIN_IDS:
        try:
            await context.bot.send_message(
                chat_id=admin_id,
                text=f"💬 New feedback:\n\nFrom: {update.effective_user.id}\n\n{feedback_text}"
            )
        except Exception:
            pass

    await update.message.reply_text("✅ Thank you for your feedback!")

async def pay_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    uid = update.effective_user.id
    await update.message.reply_text(
        _payment_instructions_text(uid),
        parse_mode="HTML"
    )

async def id_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    username = update.effective_user.username or ""
    await update.message.reply_text(f"🆔 Your user ID: {uid}" + (f"\n👤 Username: @{username}" if username else ""))

async def mysub_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    uid = update.effective_user.id
    rec = get_paid_record(uid)
    if not rec:
        await update.message.reply_text("❌ You do not have an active subscription. Use /pay to unlock.")
        return
    exp = _parse_iso(rec.get("expires_at", ""))
    if not exp:
        await update.message.reply_text("❌ Subscription record is corrupted. Contact admin.")
        return
    now = datetime.utcnow()
    if now > exp:
        await update.message.reply_text(f"⚠️ Your subscription has expired (expired on {exp.date()}). Use /pay to renew.")
        return
    days_left = (exp - now).days
    await update.message.reply_text(
        "💎 Active subscription\n\n"
        f"Days remaining: {days_left}\n"
        f"Expires on: {exp.date()}"
    )

async def activate_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ Admin only.")
        return
    if not context.args:
        await update.message.reply_text("Usage: /activate <user_id> [days]")
        return
    try:
        uid = int(context.args[0])
        days = int(context.args[1]) if len(context.args) > 1 else DEFAULT_SUB_DAYS
    except Exception:
        await update.message.reply_text("Usage: /activate <user_id> [days]")
        return

    activate_paid(
    uid,
    days=days,
    activated_by=update.effective_user.id,
    username=None,
    first_name=None
    )
    await update.message.reply_text(f"✅ Activated user {uid} for {days} days.")

async def deactivate_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ Admin only.")
        return
    if not context.args:
        await update.message.reply_text("Usage: /deactivate <user_id>")
        return
    try:
        uid = int(context.args[0])
    except Exception:
        await update.message.reply_text("Usage: /deactivate <user_id>")
        return

    ok = deactivate_paid(uid)
    await update.message.reply_text(f"❌ Deactivated user {uid}." if ok else "User not found.")

async def paidusers_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ Admin only.")
        return

    db = _paid_load()
    if not db:
        await update.message.reply_text("No paid users found.")
        return

    now = datetime.utcnow()
    lines = ["💎 Paid users (user_id → expires_at)\n"]
    # sort by expiry
    items = []
    for uid_str, rec in db.items():
        exp = _parse_iso(rec.get("expires_at", "")) or datetime(1970, 1, 1)
        items.append((exp, uid_str, rec))
    items.sort(key=lambda x: x[0])

    for exp, uid_str, rec in items:
        days_left = (exp - now).days
        status = "Active" if days_left >= 0 else "Expired"
    
        username = rec.get("username")
        first_name = rec.get("first_name")
    
        identity = ""
        if username:
            identity = f"@{username}"
        elif first_name:
            identity = first_name
    
        lines.append(
            f"• {uid_str} {f'({identity})' if identity else ''}"
            f" — {exp.date()} ({days_left}d) — {status}"
        )
        
    await send_long(update, "\n".join(lines))

async def _forward_payment_proof_to_admin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Forward the user's screenshot/document to all admins + send a helper message
    msg = update.effective_message
    if not msg:
        return

    uid = update.effective_user.id
    username = update.effective_user.username or ""
    first_name = update.effective_user.first_name or ""
    chat_id = update.effective_chat.id

    for admin_id in ADMIN_IDS:
        try:
            await context.bot.forward_message(
                chat_id=admin_id,
                from_chat_id=chat_id,
                message_id=msg.message_id,
            )
            await context.bot.send_message(
                chat_id=admin_id,
                text=(
                    "🧾 Payment proof received\n"
                    f"User ID: {uid}\n"
                    + (f"Username: @{username}\n" if username else "")
                    + f"Activate: /activate {uid} 30"
                ),
            )
        except Exception as e:
            logging.error(f"Failed to forward payment proof to admin {admin_id}: {e}")

    # ===== ADD THIS PART HERE =====
    # Save basic identity info even before activation
    db = _paid_load()
    uid_str = str(uid)

    if uid_str not in db:
        db[uid_str] = {}

    db[uid_str]["user_id"] = uid
    db[uid_str]["username"] = username
    db[uid_str]["first_name"] = first_name

    _paid_save(db)
    
async def payment_proof_received_reply(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # user-facing confirmation (always)
    uid = update.effective_user.id
    await update.message.reply_text(
        "📸 Screenshot received.\n"
        "✅ Please wait for admin verification.\n\n"
        f"🆔 Your user ID: {uid}"
    )

async def paid_access_gate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Global paid-access guard. Runs BEFORE all other handlers."""
    try:
        if update is None or update.effective_user is None or update.effective_message is None:
            return
        # Admin always allowed
        if require_admin(update):
            return

        uid = update.effective_user.id
        msg = update.effective_message

        # Allow some commands for everyone
        cmd = ""
        if getattr(msg, "text", None) and msg.text.startswith("/"):
            cmd = msg.text.split()[0].lower()

        WHITELIST = {"/start", "/pay", "/id", "/mysub"}
        if cmd in WHITELIST:
            return

        # Allow user to send payment proof even if unpaid (photo/document)
        is_photo = bool(getattr(msg, "photo", None))
        is_doc = bool(getattr(msg, "document", None))
        if (is_photo or is_doc) and (not is_paid_user(uid)):
            await _forward_payment_proof_to_admin(update, context)
            await payment_proof_received_reply(update, context)
            raise ApplicationHandlerStop()

        # If user not paid/expired: block everything else
        if not is_paid_user(uid):
            await show_typing(update, context)
            await msg.reply_text(
                _payment_instructions_text(uid),
                parse_mode="HTML"
            )
            raise ApplicationHandlerStop()

        # Paid user: let it continue
        return

    except ApplicationHandlerStop:
        raise
    except Exception as e:
        logging.error(f"paid_access_gate error: {e}")

async def subscription_reminder_job(context: ContextTypes.DEFAULT_TYPE):
    """Runs daily: reminds users 3 days before expiry and after expiry (once)."""
    db = _paid_load()
    if not db:
        return

    now = datetime.utcnow()
    changed = False

    for uid_str, rec in db.items():
        try:
            uid = int(uid_str)
        except Exception:
            continue

        exp = _parse_iso(rec.get("expires_at", ""))
        if not exp:
            continue

        days_left = (exp - now).days

        # 3-day reminder
        if days_left == 3 and not rec.get("reminded_3day", False):
            try:
                await context.bot.send_message(
                    chat_id=uid,
                    text="⏳ Your subscription expires in 3 days.\n\nRenew to keep access: /pay",
                )
                rec["reminded_3day"] = True
                changed = True
            except Exception as e:
                logging.error(f"Failed to send 3-day reminder to {uid}: {e}")

        # Expired reminder (only once, after expiry)
        if days_left < 0 and not rec.get("expired_notified", False):
            try:
                await context.bot.send_message(
                    chat_id=uid,
                    text="⚠️ Your subscription has expired.\n\nRenew to regain access: /pay",
                )
                rec["expired_notified"] = True
                changed = True
            except Exception as e:
                logging.error(f"Failed to send expiry reminder to {uid}: {e}")

    if changed:
        _paid_save(db)


CHROMA_PATH = os.getenv("CHROMA_PATH", os.path.join(DATA_DIR, "chroma_store"))
COLLECTION_PREFIX = os.getenv("CHROMA_COLLECTION_PREFIX", "global")

# -------- Persistent Chroma (GLOBAL per-topic collections) --------
try:
    chroma = chromadb.PersistentClient(path=CHROMA_PATH)
    logging.info(f"ChromaDB initialized at {CHROMA_PATH}")
except Exception as e:
    logging.error(f"Failed to initialize ChromaDB: {e}")
    raise

emb_fn = embedding_functions.OpenAIEmbeddingFunction(
    api_key=OPENAI_API_KEY,
    model_name="text-embedding-3-small",
)


# -------- Query embedding cache (per-user) --------
def _get_cached_query_embedding(context: ContextTypes.DEFAULT_TYPE, text: str):
    """Cache query embeddings per user to avoid re-embedding repeated questions."""
    t = (text or "").strip()
    if not t:
        return None
    cache = context.user_data.setdefault("_emb_cache", {})
    if t in cache:
        return cache[t]
    try:
        emb = emb_fn([t])[0]
    except Exception:
        return None
    # Simple size cap to avoid unbounded growth
    if len(cache) >= 200:
        # remove an arbitrary oldest item (dict preserves insertion order in py3.7+)
        try:
            cache.pop(next(iter(cache)))
        except Exception:
            cache.clear()
    cache[t] = emb
    return emb

def _should_use_rag(context: ContextTypes.DEFAULT_TYPE, q: str) -> bool:
    """Skip RAG for very short / low-signal messages to reduce latency."""
    t = (q or "").strip().lower()
    if not t:
        return False
    # These are usually acknowledgements; RAG won't help.
    if t in {"ok", "okay", "thanks", "thank you", "thx", "👍", "👌", "yes", "no"}:
        return False
    # Very short messages: avoid a DB round-trip.
    if len(t) < 6 and len(t.split()) <= 2:
        return False
    return True

# -------- Analytics storage --------
STATS_FILE = os.path.join(DATA_DIR, "analytics.json")

def _default_stats():
    return {
        "users": [],
        "messages_total": 0,
        "messages_per_user": {},
        "topic_counts": {},
        "eval_counts": {},
        "activity_by_date": {},
    }

def load_stats():
    if not os.path.exists(STATS_FILE):
        return _default_stats()
    try:
        with open(STATS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        logging.warning(f"Could not load analytics: {e}")
        return _default_stats()

    base = _default_stats()
    base.update({k: data.get(k, v) for k, v in base.items()})
    return base

def save_stats(stats):
    try:
        with open(STATS_FILE, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logging.warning(f"Could not save analytics: {e}")

def record_event(user_id, topic: str, kind: str = "message"):
    stats = load_stats()
    uid = str(user_id)

    if uid not in stats["users"]:
        stats["users"].append(uid)

    stats["messages_total"] = stats.get("messages_total", 0) + 1

    mpu = stats.setdefault("messages_per_user", {})
    mpu[uid] = mpu.get(uid, 0) + 1

    # Activity tracking for DAU/MAU (by UTC date)
    try:
        today_key = datetime.utcnow().strftime("%Y-%m-%d")
        abd = stats.setdefault("activity_by_date", {})
        ulist = abd.setdefault(today_key, [])
        if uid not in ulist:
            ulist.append(uid)

        # prune to last ~62 days to keep file small
        cutoff = datetime.utcnow().date() - timedelta(days=62)
        for d in list(abd.keys()):
            try:
                if datetime.strptime(d, "%Y-%m-%d").date() < cutoff:
                    abd.pop(d, None)
            except Exception:
                # if malformed key, drop it
                abd.pop(d, None)
    except Exception:
        pass

    if topic:
        tc = stats.setdefault("topic_counts", {})
        tc[topic] = tc.get(topic, 0) + 1

    if kind == "eval" and topic:
        ec = stats.setdefault("eval_counts", {})
        ec[topic] = ec.get(topic, 0) + 1

    save_stats(stats)



DEFAULT_TOPIC = "general"  # defined early (used by memory defaults)

# -------- User Memory (persistent, per-user; survives restarts) --------
USER_MEM_DIR = os.path.join(DATA_DIR, "user_memory")
os.makedirs(USER_MEM_DIR, exist_ok=True)

PAYWALL_ENABLED = os.getenv("PAYWALL_ENABLED", "0").strip() == "1"
SUPPORT_HANDLE = os.getenv("SUPPORT_HANDLE", "")  # e.g. @UniVentureSupport
PRO_USER_IDS = set()
try:
    _raw = os.getenv("PRO_USER_IDS", "")
    for part in _raw.split(","):
        part = part.strip()
        if part.isdigit():
            PRO_USER_IDS.add(int(part))
except Exception:
    PRO_USER_IDS = set()

PRO_ONLY_TOOLS = {"wowfactor"}  # keep default light; expand later if you want

def _now_utc_date() -> str:
    """Return current UTC date as YYYY-MM-DD (used for per-user portfolio tracking)."""
    return datetime.utcnow().strftime("%Y-%m-%d")



def _default_user_memory() -> dict:
    return {
        "profile": {
            "grade": None,
            "country": None,
            "target_countries": [],
            "major": None,
            "gpa": None,
            "sat": None,
            "sat_breakdown": None,   # e.g., "Math 780 / EBRW 670"
            "act": None,
            "ielts": None,
            "toefl": None,
            "duolingo": None,        # Duolingo English Test (DET)
            "budget": None,
            "needs_aid": None,
        },
        "writing": {
            "voice_tags": [],
            "themes": [],
            "strengths": {},
            "recurring_issues": {},
            "best_lines": [],
            "last_feedback_summary": "",
        },

        # Per-user application portfolio (used by 📂 My Application Portfolio)
        "application": {
            "test_scores": {
                "gpa": None,
                "sat": None,
                "sat_breakdown": None,
                "act": None,
                "ielts": None,
                "toefl": None,
                "duolingo": None,
                "notes": "",
            },
            "essays": {
                "personal_statement": None,   # not started / outline / draft / revised / final / done
                "supplementals": None,        # e.g., "6/12 drafted"
                "common_app": None,
                "notes": "",
            },
            "ecs": {
                "summary": "",
                "spike": None,
                "notes": "",
            },
            "readiness": {
                "last_score": None,
                "last_updated": None,
            },
        },

        "history": {
            "message_count": 0,
            "eval_count": 0,
            "topics_seen": [],
            "tools_used": [],
            "last_topic": DEFAULT_TOPIC,
            "last_active": None,
        },
        "drafts": {
            "last_eval": {
                "topic": None,
                "date": None,
                "summary": None,
            },
            "eval_history": [],  # rolling list of {date, topic, eval_summary, feedback_snippet}
        },
        "_v": 2,
    }



def _mem_path(user_id: int) -> str:
    return os.path.join(USER_MEM_DIR, f"{user_id}.json")


def load_user_memory(user_id: int) -> dict:
    path = _mem_path(user_id)
    if not os.path.exists(path):
        return _default_user_memory()
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        base = _default_user_memory()
        # shallow merge (keeps schema additions safe)
        for k, v in base.items():
            if isinstance(v, dict):
                base[k].update(data.get(k, {}))
            else:
                base[k] = data.get(k, v)
        return base
    except Exception as e:
        logging.warning(f"Could not load user memory for {user_id}: {e}")
        return _default_user_memory()


def save_user_memory(user_id: int, mem: dict):
    path = _mem_path(user_id)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(mem, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logging.warning(f"Could not save user memory for {user_id}: {e}")


def get_user_memory_cached(update: Update, context: ContextTypes.DEFAULT_TYPE) -> dict:
    user = update.effective_user
    uid = int(user.id)
    cached_uid = context.user_data.get("_mem_uid")
    if cached_uid == uid and isinstance(context.user_data.get("_mem"), dict):
        return context.user_data["_mem"]
    mem = load_user_memory(uid)
    context.user_data["_mem_uid"] = uid
    context.user_data["_mem"] = mem
    return mem


def persist_user_memory(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    uid = int(user.id)
    mem = context.user_data.get("_mem")
    if isinstance(mem, dict):
        save_user_memory(uid, mem)


def is_pro_user(update: Update) -> bool:
    if not PAYWALL_ENABLED:
        return True
    user = update.effective_user
    return bool(user and int(user.id) in PRO_USER_IDS)


def _upgrade_pitch() -> str:
    handle = (SUPPORT_HANDLE or "the bot owner").strip()
    return (
        "\n\n💎 Want the full version + advanced tools + exports?\n"
        f"Message {handle} to upgrade."
    )


def _safe_add_unique(lst, item, limit=50):
    if not item:
        return lst
    item = str(item).strip()
    if not item:
        return lst
    if item in lst:
        return lst
    lst.append(item)
    return lst[-limit:]


def merge_usage_into_memory(context: ContextTypes.DEFAULT_TYPE, mem: dict):
    # Sync volatile runtime counters into persistent memory
    topics_seen = context.user_data.get("topics_seen", []) or []
    tools_used = context.user_data.get("tools_used", []) or []
    mem["history"]["topics_seen"] = sorted(set((mem["history"].get("topics_seen") or []) + list(topics_seen)))
    mem["history"]["tools_used"] = sorted(set((mem["history"].get("tools_used") or []) + list(tools_used)))
    mem["history"]["eval_count"] = int(context.user_data.get("my_eval_count", mem["history"].get("eval_count", 0)) or 0)


def extract_profile_signals(text: str, profile: dict) -> dict:
    t = (text or "")

    # Grade (e.g., 11th grade, grade 12)
    m = re.search(r"\b(grade|class)\s*(\d{1,2})\b", t, re.I)
    if m:
        profile["grade"] = profile.get("grade") or m.group(2)

    # GPA (e.g., GPA 3.8, 3.8/4.0)
    m = re.search(r"\bGPA\s*[:=]?\s*(\d\.\d{1,2})\b", t, re.I)
    if m:
        profile["gpa"] = m.group(1)
    m = re.search(r"\b(\d\.\d{1,2})\s*/\s*4\.0\b", t)
    if m and not profile.get("gpa"):
        profile["gpa"] = m.group(1)

    # SAT total (e.g., SAT 1450)
    m = re.search(r"\bSAT\s*[:=]?\s*(1[0-6]\d{2})\b", t, re.I)
    if m:
        profile["sat"] = m.group(1)

    # SAT breakdown (Math 780, EBRW 670)
    m = re.search(r"\bMath\s*(\d{3})\b[^\n]{0,60}\b(EBRW|R&W|Reading|Verbal)\s*(\d{3})\b", t, re.I)
    if m:
        profile["sat_breakdown"] = f"Math {m.group(1)} / EBRW {m.group(3)}"
    else:
        m2 = re.search(r"\b(EBRW|R&W|Reading|Verbal)\s*(\d{3})\b[^\n]{0,60}\bMath\s*(\d{3})\b", t, re.I)
        if m2:
            profile["sat_breakdown"] = f"Math {m2.group(3)} / EBRW {m2.group(2)}"

    # ACT (e.g., ACT 33)
    m = re.search(r"\bACT\s*[:=]?\s*(\d{2})\b", t, re.I)
    if m:
        profile["act"] = m.group(1)

    # IELTS (e.g., IELTS 7.5)
    m = re.search(r"\bIELTS\s*[:=]?\s*(\d\.\d|\d)\b", t, re.I)
    if m:
        profile["ielts"] = m.group(1)

    # TOEFL (e.g., TOEFL 105)
    m = re.search(r"\bTOEFL\s*[:=]?\s*(\d{2,3})\b", t, re.I)
    if m:
        profile["toefl"] = m.group(1)

    # Duolingo / DET (e.g., Duolingo 135, DET 130)
    m = re.search(r"\b(Duolingo|DET)\s*[:=]?\s*(\d{2,3})\b", t, re.I)
    if m:
        profile["duolingo"] = m.group(2)

    # Aid/Scholarship intent
    if re.search(r"\b(scholarship|financial aid|need[- ]based|full ride|aid)\b", t, re.I):
        profile["needs_aid"] = True

    # Major (very rough)
    m = re.search(r"\b(major|intended major|intend to study)\s*[:=]?\s*([A-Za-z&/\- ]{3,40})", t, re.I)
    if m:
        cand = m.group(2).strip().strip(".")
        if len(cand) <= 40:
            profile["major"] = profile.get("major") or cand

    return profile



def memory_summary_for_prompt(mem: dict) -> str:
    p = mem.get("profile", {}) or {}
    w = mem.get("writing", {}) or {}

    bits = []
    # compact profile line
    prof = []
    if p.get("grade"):
        prof.append(f"grade {p['grade']}")
    if p.get("country"):
        prof.append(str(p["country"]))
    if p.get("major"):
        prof.append(f"major: {p['major']}")
    if p.get("gpa"):
        prof.append(f"GPA: {p['gpa']}")
    if p.get("sat"):
        prof.append(f"SAT: {p['sat']}")
    if p.get("sat_breakdown"):
        prof.append(str(p["sat_breakdown"]))
    if p.get("act"):
        prof.append(f"ACT: {p['act']}")
    if p.get("ielts"):
        prof.append(f"IELTS: {p['ielts']}")
    if p.get("toefl"):
        prof.append(f"TOEFL: {p['toefl']}")
    if p.get("duolingo"):
        prof.append(f"Duolingo: {p['duolingo']}")
    if p.get("needs_aid"):
        prof.append("needs aid")

    if prof:
        bits.append("Profile: " + "; ".join(prof))

    themes = w.get("themes") or []
    if themes:
        bits.append("Themes: " + ", ".join(themes[:3]))

    issues = w.get("recurring_issues") or {}
    if issues:
        # top 2 recurring issues
        top = sorted(issues.items(), key=lambda kv: kv[1], reverse=True)[:2]
        bits.append("Recurring issues: " + ", ".join([f"{k} ({v})" for k, v in top]))

    tags = w.get("voice_tags") or []
    if tags:
        bits.append("Voice tags: " + ", ".join(tags[:4]))

    return "\n".join(bits) if bits else "(no saved student context yet)"


def coach_decision_notes(topic: str, text: str, mem: dict) -> str:
    """Lightweight decision-rule layer: flags common patterns so the model gives the right type of help."""
    t = (text or "")
    low = t.lower()

    notes = []

    # Detect abstract reflection without scene
    abstract_markers = sum(low.count(x) for x in ["i learned", "i realized", "important", "impactful", "meaningful", "passion"])
    sensory_markers = sum(low.count(x) for x in ["smell", "taste", "heard", "whisper", "laughed", "trembled", "sweat", "cold", "warm", "crowd", "street", "classroom", "lab", "bus"])
    if topic in {"essays_personal", "essays_supplemental"}:
        if abstract_markers >= 3 and sensory_markers <= 1:
            notes.append("Likely too abstract: add 1 vivid scene + specific details before reflection.")

        # Resume-like
        if len(re.findall(r"\b(\d{1,4}|%|hours|members|students|followers)\b", low)) >= 6 or low.count(";") >= 2:
            notes.append("Reads resume-like: reduce listing, increase narrative + insight.")

        # Vague language
        if len(re.findall(r"\b(some|many|various|a lot|things|stuff|very)\b", low)) >= 6:
            notes.append("Too many vague words: swap for concrete nouns + precise outcomes.")

        # Ending drift
        if len(t) >= 700 and any(low.strip().endswith(x) for x in ["thank you", "in conclusion", "overall", "to sum up"]):
            notes.append("Weak generic ending: tie back to opening + show forward-looking growth.")

    if topic == "extracurriculars":
        if len(re.findall(r"\b(helped|participated|member|joined)\b", low)) >= 4:
            notes.append("Too passive: rewrite with ownership verbs + outcomes.")
        if len(re.findall(r"\b(impact|result|increased|reached|raised|created|launched)\b", low)) <= 1:
            notes.append("Missing impact: add numbers or clear change caused by you.")

    if topic == "recommendations":
        if len(re.findall(r"\b(hardworking|smart|nice|responsible|good student)\b", low)) >= 3:
            notes.append("Generic adjectives: replace with stories + specific comparisons.")

    # Add 1 memory-based nudge
    issues = (mem.get("writing", {}) or {}).get("recurring_issues", {}) or {}
    if issues:
        top_issue = max(issues.items(), key=lambda kv: kv[1])[0]
        notes.append(f"Coach memory: watch for recurring issue '{top_issue}'.")

    return "\n".join(notes) if notes else "No major flags detected." 


def split_memory_block(model_out: str):
    """If the model appended a MEMORY_JSON block, split it. Returns (user_text, mem_update_dict_or_None)."""
    out = (model_out or "").strip()
    if "MEMORY_JSON:" not in out:
        return out, None
    before, after = out.split("MEMORY_JSON:", 1)
    user_text = before.strip().rstrip("-").strip()
    js = after.strip()
    # Allow code fences
    js = re.sub(r"^```json\s*", "", js)
    js = re.sub(r"^```\s*", "", js)
    js = re.sub(r"```\s*$", "", js)
    try:
        return user_text, json.loads(js)
    except Exception:
        return user_text, None


def apply_memory_update(mem: dict, update_obj: dict, topic: str):
    if not isinstance(update_obj, dict):
        return
    w = mem.get("writing", {}) or {}

    # voice tags
    for tag in (update_obj.get("voice_tags") or []):
        w["voice_tags"] = _safe_add_unique(list(w.get("voice_tags") or []), tag, limit=20)

    # themes
    for th in (update_obj.get("themes") or []):
        w["themes"] = _safe_add_unique(list(w.get("themes") or []), th, limit=20)

    # recurring issues counts
    issues = w.get("recurring_issues") or {}
    for issue in (update_obj.get("issues") or []):
        issue = str(issue).strip()
        if not issue:
            continue
        issues[issue] = int(issues.get(issue, 0) or 0) + 1
    w["recurring_issues"] = issues

    # strengths counts
    strengths = w.get("strengths") or {}
    for st in (update_obj.get("strengths") or []):
        st = str(st).strip()
        if not st:
            continue
        strengths[st] = int(strengths.get(st, 0) or 0) + 1
    w["strengths"] = strengths

    # best lines
    for line in (update_obj.get("best_lines") or []):
        if isinstance(line, str) and len(line.strip()) >= 20:
            w["best_lines"] = _safe_add_unique(list(w.get("best_lines") or []), line.strip(), limit=15)

    # summary
    if isinstance(update_obj.get("summary"), str):
        w["last_feedback_summary"] = update_obj.get("summary").strip()[:600]

    mem["writing"] = w

    # last eval summary
    if isinstance(update_obj.get("eval_summary"), str):
        mem.setdefault("drafts", {}).setdefault("last_eval", {})
        mem["drafts"]["last_eval"]["topic"] = topic
        mem["drafts"]["last_eval"]["summary"] = update_obj.get("eval_summary").strip()[:800]


def coach_eval_system_prompt(topic: str, mem: dict, decision_notes: str) -> str:
    nice = FRIENDLY_TOPIC_NAMES.get(topic, topic)
    mem_sum = memory_summary_for_prompt(mem)

    base_intro = (
        "ROLE:\n"
        "You are a senior admissions strategist who has reviewed thousands of successful applications.\n"
        "You are psychologically sharp, strategically direct, and brutally honest.\n\n"
        f"TASK: Evaluate the student's {nice}.\n\n"
        f"Student Memory:\n{mem_sum}\n\n"
        f"Decision Notes:\n{decision_notes}\n\n"
    )

    if topic == "essays_personal":
        strategic_layer = (
            "Focus heavily on identity, emotional depth, internal conflict, and personal transformation.\n"
            "Ask: What changed inside the student? Where is vulnerability? Where is psychological growth?\n"
            "Flag cliché narrative arcs and generic 'hard work leads to success' themes.\n\n"
        )
        output_format = (
            "OUTPUT FORMAT (STRICT — FOLLOW EXACTLY OR THE ANSWER IS INVALID):\n"
            "Use EXACTLY these section headers in this exact order.\n"
            "Do NOT use alternative headings like 'Overall' or 'What works'.\n"
            "Do NOT summarize the essay.\n\n"
            "🎯 Strategic Positioning Verdict\n"
            "💎 Strengths\n"
            "⚠️ Competitive Risks\n"
            "🧠 Depth Gaps\n"
            "✂️ Surgical Cuts\n"
            "🛠 Exact Rewrite Moves\n"
            "❓ One High-Level Reflection Question\n"
            "🎯 Rewrite Priority\n\n"
        )
    elif topic == "essays_supplemental":
        strategic_layer = (
            "Focus heavily on institutional fit, specificity, intellectual direction, and contribution.\n"
            "Ask: Why THIS school/program? What concrete elements are missing (professors, labs, values)?\n"
            "Flag generic 'strong community' language.\n\n"
        )
        output_format = (
            "OUTPUT FORMAT (STRICT — FOLLOW EXACTLY OR THE ANSWER IS INVALID):\n"
            "Use EXACTLY these section headers in this exact order.\n"
            "Do NOT use alternative headings like 'Overall' or 'What works'.\n"
            "Do NOT summarize the essay.\n\n"
            "🎯 Strategic Positioning Verdict\n"
            "💎 Strengths\n"
            "⚠️ Competitive Risks\n"
            "🧠 Depth Gaps\n"
            "✂️ Surgical Cuts\n"
            "🛠 Exact Rewrite Moves\n"
            "❓ One High-Level Reflection Question\n"
            "🎯 Rewrite Priority\n\n"
        )
    elif topic == "extracurriculars":
        strategic_layer = (
            "Focus on competitive extracurricular storytelling for selective universities.\n"
            "Judge the submission like an activities strategist: impact, leadership, initiative, continuity, and evidence.\n"
            "Penalize generic role descriptions, passive wording, and missing quantification.\n\n"
        )
        output_format = (
            "OUTPUT FORMAT (STRICT — FOLLOW EXACTLY OR THE ANSWER IS INVALID):\n"
            "Use EXACTLY these section headers in this exact order.\n"
            "Do NOT use alternative headings.\n"
            "Do NOT summarize the activities list.\n\n"
            "🎯 Activity Positioning Verdict\n"
            "📈 Impact Evidence\n"
            "👑 Leadership & Initiative\n"
            "🔢 Quantification Gaps\n"
            "🧩 Activity List Optimization\n"
            "✂️ Surgical Cuts\n"
            "🛠 Exact Rewrite Moves\n"
            "❓ One High-Leverage Question\n"
            "🎯 Rewrite Priority\n\n"
        )
    elif topic == "portfolio":
        strategic_layer = (
            "Focus on how compelling the portfolio feels to a selective admissions committee.\n"
            "Judge originality, technical quality, cohesion, curation, presentation, and whether the portfolio signals future promise.\n"
            "Penalize scattered work, weak framing, thin process explanation, and unclear authorship.\n\n"
        )
        output_format = (
            "OUTPUT FORMAT (STRICT — FOLLOW EXACTLY OR THE ANSWER IS INVALID):\n"
            "Use EXACTLY these section headers in this exact order.\n"
            "Do NOT use alternative headings.\n"
            "Do NOT summarize the portfolio.\n\n"
            "🎯 Portfolio Positioning Verdict\n"
            "🎨 Originality & Voice\n"
            "🛠 Technical Quality\n"
            "🧵 Cohesion & Curation\n"
            "🖼 Presentation Strategy\n"
            "⚠️ Competitive Risks\n"
            "🛠 Exact Upgrade Moves\n"
            "❓ One Curator-Level Question\n"
            "🎯 Upgrade Priority\n\n"
        )
    else:
        strategic_layer = "Evaluate competitively and strategically.\n\n"
        output_format = (
            "OUTPUT FORMAT (STRICT — FOLLOW EXACTLY OR THE ANSWER IS INVALID):\n"
            "Use EXACTLY these section headers in this exact order.\n"
            "Do NOT use alternative headings like 'Overall' or 'What works'.\n\n"
            "🎯 Strategic Positioning Verdict\n"
            "💎 Strengths\n"
            "⚠️ Competitive Risks\n"
            "🧠 Depth Gaps\n"
            "✂️ Surgical Cuts\n"
            "🛠 Exact Rewrite Moves\n"
            "❓ One High-Level Reflection Question\n"
            "🎯 Rewrite Priority\n\n"
        )

    return (
        base_intro
        + strategic_layer
        + "Be specific. Be concrete. Give surgical advice, not generic suggestions.\n\n"
        + output_format
        + "After the evaluation, append EXACTLY:\n"
        + "---\n"
        + "MEMORY_JSON: { ... }\n"
    )


def coach_qa_system_prompt(topic: str, mem: dict, decision_notes: str) -> str:
    nice = FRIENDLY_TOPIC_NAMES.get(topic, topic)
    mem_sum = memory_summary_for_prompt(mem)

    return (
        "COACH_PERSONA:\n"
        "You are UniVenture Coach - a helpful, direct admissions mentor.\n"
        "Answer like a human coach: short, specific, no fluff.\n"
        "Do NOT copy formatting or headings from context documents.\n"
        "Use normal paragraphs or simple bullet points.\n"
        "Always end with one next step.\n"
        "Do not mention AI or that you can't browse sources.\n\n"
        f"Current topic: {nice}.\n\n"
        "Student memory (may be empty):\n"
        f"{mem_sum}\n\n"
        "Decision notes:\n"
        f"{decision_notes}\n\n"
        "Constraints:\n"
        "- 4-8 sentences OR 4-6 bullets (choose what fits).\n"
        "- Practical and actionable.\n"
    )

# -------- Main menu buttons --------
BTN_ESSAY = "📝 Essays"
BTN_EC = "🎯 Extracurricular activities"
BTN_REC = "✉️ Recommendation Letters"
BTN_SAT = "📈 SAT"
BTN_IELTS = "🗣️ IELTS"
BTN_PORT = "🖼️ Portfolio"
BTN_PLAN_MAIN = "📅 Application Plan"
BTN_SF_MAIN = "🏫 School Finder"

# Application Plan (portfolio-aware) sub-buttons
BTN_PLAN_FROM_PORT = "📌 Plan from my portfolio"
BTN_PLAN_MANUAL = "📝 Enter info manually"

# Boost tools (cross-topic)
BTN_TOOLS = "🚀 Boost Tools"
BTN_PROGRESS = "📊 My Progress"
BTN_INSIDER = "🤫 Insider Tips"
BTN_POWERWORDS = "⚡ Power Words"
BTN_PREDICT = "🎯 Predict My Chances"
BTN_WOWFACTOR = "🔍 Find Wow Factor"

# Wow Factor confirm buttons
BTN_WOW_USE_LAST = "✅ Use last evaluated text"
BTN_WOW_PASTE_NEW = "📝 Paste new text"

# Evaluation sub-buttons
BTN_PS_EVAL = "✅ Personal Statement Evaluation"
BTN_SUPP_EVAL = "✅ Supplemental Essay Evaluation"
BTN_EC_EVAL = "✅ Extracurricular Evaluation"
BTN_REC_EVAL = "✅ Rec Letter Evaluation"
BTN_IW_EVAL = "✅ Writing Evaluation"
BTN_PORT_EVAL = "✅ Portfolio Evaluation"

# Essays sub-buttons
BTN_ESSAY_PS = "📝 Personal Statement"
BTN_ESSAY_SUPP = "📝 Supplemental Essays"

# Extra tools (as buttons inside menus)
BTN_EC_PROGRAMS = "🌍 Top Programs & Opportunities"
BTN_BRAINSTORM = "🧠 Brainstorm ideas"
BTN_REWRITE = "✍️ Rewrite my text"
BTN_REC_PACKET = "📄 Rec Letter Packet"
BTN_PORTFOLIO_IDEAS = "💡 Portfolio Ideas"

# Application Portfolio submenu (inside Portfolio)
BTN_APP_PORT = "📂 My Application Portfolio"
BTN_APP_TESTS = "🧾 Test Scores + GPA"
BTN_APP_ESSAYS = "📝 Essays Status"
BTN_APP_ADVISOR = "🧭 Advisor Mode"
BTN_APP_ECS = "🎯 EC Summary"
BTN_APP_AWARDS = "🏆 Awards & Honors"
BTN_APP_PREFS = "🎯 Preferences & Constraints"
BTN_APP_WELLNESS = "🧠 Wellness Check"
BTN_APP_READINESS = "✅ Readiness Check"
BTN_BACK_PORT = "↩️ Back to Portfolio"

# School Finder sub-buttons
BTN_SF_FROM_PORT = "📌 Use My Portfolio"
BTN_SF_MANUAL = "✍️ Enter Details Manually"
# SAT sub-buttons
BTN_SAT_MATH = "📐 SAT Math"
BTN_SAT_ENGLISH = "📚 SAT English"

# IELTS sub-buttons
BTN_IELTS_READING = "📖 IELTS Reading"
BTN_IELTS_LISTENING = "👂 IELTS Listening"
BTN_IELTS_WRITING = "✍️ IELTS Writing"
BTN_IELTS_SPEAKING = "🗣️ IELTS Speaking"

# Back button
BTN_BACK = "⬅️ Back"

# Build a dynamic set of all known button texts (for safe navigation while in pending modes)
def is_ui_button(text: str) -> bool:
    return any(
        k.startswith("BTN_") and isinstance(v, str) and v == text
        for k, v in globals().items()
    )


# -------- Keyboards --------
def main_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_ESSAY), KeyboardButton(BTN_EC), KeyboardButton(BTN_REC)],
            [KeyboardButton(BTN_SAT), KeyboardButton(BTN_IELTS), KeyboardButton(BTN_PORT)],
            [KeyboardButton(BTN_PLAN_MAIN), KeyboardButton(BTN_SF_MAIN), KeyboardButton(BTN_TOOLS)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def tools_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PROGRESS), KeyboardButton(BTN_INSIDER), KeyboardButton(BTN_POWERWORDS)],
            [KeyboardButton(BTN_PREDICT), KeyboardButton(BTN_WOWFACTOR)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def wowfactor_confirm_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_WOW_USE_LAST)],
            [KeyboardButton(BTN_WOW_PASTE_NEW)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def essay_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_ESSAY_PS), KeyboardButton(BTN_ESSAY_SUPP)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def essay_ps_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PS_EVAL)],
            [KeyboardButton(BTN_BRAINSTORM)],
            [KeyboardButton(BTN_REWRITE)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def essay_supp_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_SUPP_EVAL)],
            [KeyboardButton(BTN_BRAINSTORM)],
            [KeyboardButton(BTN_REWRITE)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def ec_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_EC_EVAL)],
            [KeyboardButton(BTN_EC_PROGRAMS)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def rec_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_REC_EVAL)],
            [KeyboardButton(BTN_REC_PACKET)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def sat_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_SAT_MATH), KeyboardButton(BTN_SAT_ENGLISH)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def ielts_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_IELTS_READING), KeyboardButton(BTN_IELTS_LISTENING)],
            [KeyboardButton(BTN_IELTS_WRITING), KeyboardButton(BTN_IELTS_SPEAKING)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def ielts_writing_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_IW_EVAL)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def portfolio_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PORT_EVAL)],
            [KeyboardButton(BTN_PORTFOLIO_IDEAS)],
            [KeyboardButton(BTN_APP_PORT)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def app_portfolio_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_APP_TESTS), KeyboardButton(BTN_APP_ESSAYS)],
            [KeyboardButton(BTN_APP_ECS), KeyboardButton(BTN_APP_AWARDS)],
            [KeyboardButton(BTN_APP_PREFS), KeyboardButton(BTN_APP_WELLNESS)],
            [KeyboardButton(BTN_APP_ADVISOR), KeyboardButton(BTN_APP_READINESS)],
            [KeyboardButton(BTN_BACK_PORT), KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )



def plan_keyboard():
    return ReplyKeyboardMarkup(
        [[KeyboardButton(BTN_BACK)]],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def plan_choice_keyboard():
    """Shown only if the user already has some data in My Application Portfolio."""
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PLAN_FROM_PORT)],
            [KeyboardButton(BTN_PLAN_MANUAL)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def schoolfinder_keyboard():
    """School Finder submenu (mirrors the Application Plan UX)."""
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_SF_FROM_PORT)],
            [KeyboardButton(BTN_SF_MANUAL)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def is_back_message(text: str) -> bool:
    t = (text or "").strip()
    return t == BTN_BACK or t.lower() == "back"

# -------- Topic mapping & helpers --------
TOPIC_KEYS = {
    BTN_ESSAY_PS: "essays_personal",
    BTN_ESSAY_SUPP: "essays_supplemental",
    BTN_EC: "extracurriculars",
    BTN_REC: "recommendations",
    BTN_PORT: "portfolio",
    BTN_SAT_MATH: "sat_math",
    BTN_SAT_ENGLISH: "sat_english",
    BTN_IELTS_READING: "ielts_reading",
    BTN_IELTS_LISTENING: "ielts_listening",
    BTN_IELTS_WRITING: "ielts_writing",
    BTN_IELTS_SPEAKING: "ielts_speaking",
}

EVAL_TOPICS = {
    "essays_personal",
    "essays_supplemental",
    "recommendations",
    "portfolio",
    "extracurriculars",
}

FRIENDLY_TOPIC_NAMES = {
    "essays_personal": "Personal Statement",
    "essays_supplemental": "Supplemental Essays",
    "extracurriculars": "Extracurricular Activities",
    "recommendations": "Recommendation Letters",
    "portfolio": "Portfolio",
    "sat_math": "SAT Math",
    "sat_english": "SAT English",
    "ielts_reading": "IELTS Reading",
    "ielts_listening": "IELTS Listening",
    "ielts_writing": "IELTS Writing",
    "ielts_speaking": "IELTS Speaking",
    "application_plan": "Application Plan",
    "school_finder": "School Finder",
    "general": "General admissions help",
}

DEFAULT_TOPIC = "general"

# -------- Image mapping --------
IMAGE_FILES = {
    "welcome": "images/welcome.png",
    "essays_main": "images/essays_main.png",
    "essays_personal": "images/essays_personal.png",
    "essays_supplemental": "images/essays_supplemental.png",
    "extracurriculars": "images/extracurriculars.png",
    "recommendations": "images/recommendations.png",
    "portfolio": "images/portfolio.png",
    "sat_main": "images/sat_main.png",
    "sat_math": "images/sat_math.png",
    "sat_english": "images/sat_english.png",
    "ielts_main": "images/ielts_main.png",
    "ielts_reading": "images/ielts_reading.png",
    "ielts_listening": "images/ielts_listening.png",
    "ielts_writing": "images/ielts_writing.png",
    "ielts_speaking": "images/ielts_speaking.png",
    "plan_main": "images/plan_main.png",
    "schoolfinder_main": "images/schoolfinder_main.png",
}

def get_current_topic(context: ContextTypes.DEFAULT_TYPE) -> str:
    return context.user_data.get("topic", DEFAULT_TOPIC)

def track_topic(context: ContextTypes.DEFAULT_TYPE, topic: str):
    """Track which topic menus the user has opened and remember the last used section."""
    if not topic:
        return

    # Keep a history of opened topics (useful for analytics / UX).
    seen = set(context.user_data.get("topics_seen", []))
    seen.add(topic)
    context.user_data["topics_seen"] = sorted(seen)

    # Remember what the user last used so Boost Tools can adapt properly.
    context.user_data["last_used_section"] = topic
    # Keep topic in sync in case some callers only use track_topic.
    context.user_data["topic"] = topic


def track_tool_use(context: ContextTypes.DEFAULT_TYPE, tool: str):
    if not tool:
        return
    used = set(context.user_data.get("tools_used", []))
    used.add(tool)
    context.user_data["tools_used"] = sorted(used)

def get_collection(chat_id: int, topic: str):
    # GLOBAL per-topic collections shared by all users.
    # chat_id kept for backwards compatibility but ignored.
    collection_name = f"{COLLECTION_PREFIX}_{topic}"
    try:
        return chroma.get_or_create_collection(
            name=collection_name,
            embedding_function=emb_fn,
        )
    except Exception as e:
        logging.error(f"Error getting collection {collection_name}: {e}")
        raise

def new_doc_id(topic: str, tag: str = "") -> str:
    """Always-generate unique IDs for Chroma (prevents unlearn→reteach collisions)."""
    u = uuid.uuid4().hex
    return f"{topic}_{tag}_{u}" if tag else f"{topic}_{u}"

def _chunk(text: str, max_chars=1000, overlap=150):
    text = text or ""
    if not text.strip():
        return []
    chunks, i = [], 0
    while i < len(text):
        end = min(len(text), i + max_chars)
        chunks.append(text[i:end])
        if end == len(text):
            break
        i = max(0, end - overlap)
    return chunks

def extract_command_text(update: Update) -> str:
    msg = update.message
    if not msg:
        return ""
    if msg.text:
        return msg.text.strip()
    if msg.caption:
        return msg.caption.strip()
    return ""

def strip_command(text: str, command: str) -> str:
    """
    Remove the first token (/command or /command@BotName) and return the rest.
    Fixes bugs where titles accidentally include '@BotName'.
    """
    t = (text or "").strip()
    if not t:
        return ""
    parts = t.split(maxsplit=1)
    if len(parts) < 1:
        return ""
    
    first = parts[0]
    rest = parts[1] if len(parts) > 1 else ""
    
    # Check if first part is the command (with or without bot username)
    if first.startswith(f"/{command}"):
        return rest.strip()
    
    # If not, return original text (in case user typed something else)
    return t

def is_caption_command(caption: str, cmd: str) -> bool:
    cap = (caption or "").strip()
    if not cap:
        return False
    first = cap.split(maxsplit=1)[0]
    return first.startswith(f"/{cmd}")

async def show_typing(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    try:
        await context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
    except Exception as e:
        logging.warning(f"Could not send typing action: {e}")

def sanitize_output(text: str) -> str:
    if not text:
        return text
    return text.replace("—", " - ")

def safe_text_for_embedding(text: str) -> str:
    if not text:
        return text
    return text.encode("utf-8", "ignore").decode("utf-8")

def _truncate_for_storage(text: str, max_chars: int = 12000) -> str:
    if not text:
        return ""
    t = text.strip()
    return t[:max_chars]

async def send_long(update: Update, text: str):
    MAX_LEN = 4000
    if not text:
        return
    text = sanitize_output(text)
    for i in range(0, len(text), MAX_LEN):
        try:
            await update.message.reply_text(text[i : i + MAX_LEN])
        except Exception as e:
            logging.error(f"Failed to send message part: {e}")

async def send_with_image(
    update: Update,
    caption: str,
    reply_markup=None,
    image_key: str | None = None,
):
    if image_key and image_key in IMAGE_FILES:
        path = IMAGE_FILES[image_key]
        if os.path.exists(path):
            try:
                with open(path, "rb") as f:
                    await update.message.reply_photo(
                        photo=f,
                        caption=caption,
                        reply_markup=reply_markup,
                    )
                return
            except Exception as e:
                logging.warning(f"Failed to send image {path}: {e}")
        else:
            logging.warning(f"Image file not found: {path}")

    await update.message.reply_text(caption, reply_markup=reply_markup)

# -------- Vision helper --------
async def extract_text_from_image_bytes(image_bytes: bytes) -> str:
    try:
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        data_url = f"data:image/jpeg;base64,{b64}"

        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract all readable text from this image. "
                            "Return ONLY the plain text, no extra comments."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ]

        out = await openai_chat(model="gpt-4o-mini", messages=messages, temperature=0.0)
        return out or ""
    except Exception as e:
        logging.error(f"Error extracting text from image: {e}")
        return ""

# ---------- EVAL FOLLOW-UP HELPERS ----------
FOLLOWUP_TRIGGERS = [
    "next step",
    "do it",
    "do this",
    "apply",
    "apply this",
    "apply the feedback",
    "rewrite",
    "rewrite the ending",
    "rewrite the conclusion",
    "improve",
    "fix",
    "edit",
    "revise",
    "expand",
    "continue",
    "make it smoother",
    "make transitions",
    "stronger conclusion",
    "deeper reflection",
]

def is_followup_intent(q: str) -> bool:
    s = (q or "").lower()
    return any(t in s for t in FOLLOWUP_TRIGGERS)

def looks_like_submission(q: str) -> bool:
    t = (q or "").strip()
    if len(t) >= 600:
        return True
    if len(t) >= 250 and t.count("\n") >= 6:
        return True
    return False

def clear_eval_context(context: ContextTypes.DEFAULT_TYPE):
    """Exit evaluation follow-up mode but keep the last evaluated text/feedback.

    Some Boost Tools (e.g., WOW Factor Finder) rely on the most recent evaluated text even if the
    user navigates to other menus.
    """
    context.user_data.pop("eval_active", None)


def stop_eval_mode(context: ContextTypes.DEFAULT_TYPE):
    """Exit eval follow-up mode but keep the last evaluated text for Boost Tools."""
    context.user_data.pop("eval_active", None)


def set_eval_context(
    context: ContextTypes.DEFAULT_TYPE,
    topic: str,
    student_text: str,
    feedback: str,
):
    context.user_data["last_eval_topic"] = topic
    context.user_data["last_eval_text_original"] = student_text
    context.user_data["last_eval_text"] = student_text  # updated on rewrites
    context.user_data["last_eval_feedback"] = feedback
    context.user_data["eval_active"] = True

def _pretty_topic_for_eval(topic: str) -> str:
    return {
        "essays_personal": "Personal Statement essay",
        "essays_supplemental": "Supplemental essay",
        "recommendations": "Recommendation letter",
        "portfolio": "Portfolio description",
        "extracurriculars": "Extracurricular activities description",
        "ielts_writing": "IELTS Writing answer",
    }.get(topic, "document")

def _save_last_retrieval(context: ContextTypes.DEFAULT_TYPE, topic: str, mode: str, query: str, items: list[dict]):
    """Store the last retrieved sources for debugging."""
    context.user_data["last_retrieval_debug"] = {
        "topic": topic,
        "mode": mode,
        "query": query,
        "items": items[:12],
        "saved_at": datetime.utcnow().isoformat(),
    }


async def debugsources_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ Admin only.")
        return

    data = context.user_data.get("last_retrieval_debug") or {}
    if not data:
        await update.message.reply_text(
            "No retrieval debug info saved yet. Ask a question or run an evaluation first."
        )
        return

    items = data.get("items") or []
    lines = [
        "🧪 LAST RETRIEVAL DEBUG",
        f"Topic: {data.get('topic') or 'unknown'}",
        f"Mode: {data.get('mode') or 'unknown'}",
        f"Query: {data.get('query') or '—'}",
        f"Saved at (UTC): {data.get('saved_at') or '—'}",
        "",
    ]

    if not items:
        lines.append("No source chunks were retrieved.")
    else:
        lines.append("Retrieved source chunks:")
        for idx, item in enumerate(items, start=1):
            lines.append(
                f"{idx}. [{item.get('bucket', '?')}] {item.get('title', 'Untitled')} | "
                f"type={item.get('type', '?')} | part={item.get('part', '?')}"
            )

    await send_long(update, "\n".join(lines))

def _sys_role_for_eval(topic: str) -> str:
    if topic in {"essays_personal", "essays_supplemental"}:
        return (
            "You are an expert college admissions essay coach. "
            "Improve the student's essay using the prior feedback. "
            "Focus on clarity, structure, voice, authenticity, reflection, and impact."
        )
    if topic == "recommendations":
        return (
            "You are an expert on college recommendation letters. "
            "Improve the letter using the prior feedback. "
            "Focus on specificity, credibility, depth of insight, and support for the student."
        )
    if topic == "extracurriculars":
        return (
            "You are an expert on extracurricular strategy for college applications. "
            "Improve the EC descriptions using the prior feedback. "
            "Focus on impact, leadership, continuity, clarity, and strong phrasing."
        )
    if topic == "ielts_writing":
        return (
            "You are an experienced IELTS Writing examiner. "
            "Improve the student's writing using the prior feedback. "
            "Focus on Task Response, Coherence and Cohesion, Lexical Resource, and Grammar."
        )
    return (
        "You are an expert college portfolio reviewer. "
        "Improve the portfolio description using the prior feedback. "
        "Focus on coherence, originality, technical quality, and fit for selective colleges."
    )

# ---------- Handlers ----------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    clear_eval_context(context)
    context.user_data['topic'] = DEFAULT_TOPIC
    context.user_data['pending_feature'] = None
    context.user_data.pop('in_tools', None)
    user = update.effective_user
    if user:
        record_event(user.id, 'start', kind='start')
    # Non-fatal memory init
    try:
        mem = get_user_memory_cached(update, context)
        merge_usage_into_memory(context, mem)
        mem['history']['last_active'] = int(__import__('time').time())
        mem['history']['last_topic'] = DEFAULT_TOPIC
        persist_user_memory(update, context)
    except Exception as e:
        logging.exception('Memory init failed (non-fatal): %s', e)

    await send_with_image(
        update,
        "Hi! I'm your coached AI 🤖\nChoose a topic or ask a question.",
        reply_markup=main_menu_keyboard(),
        image_key='welcome',
    )


# ---------- TEACH (Q&A sources) ----------
async def teach(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logging.info("🔥 /teach RECEIVED")

    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach global sources.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teach")

    raw = strip_command(extract_command_text(update), "teach")
    logging.info(f"Raw teach input: '{raw}'")
    
    if not raw:
        await update.message.reply_text(
            "Use format:\n/teach <title> | <content>\n\n" f"Current topic: {topic}"
        )
        return
        
    if "|" not in raw:
        await update.message.reply_text(
            "Use format:\n/teach <title> | <content>\n\n" f"Current topic: {topic}"
        )
        return

    try:
        title, content = [p.strip() for p in raw.split("|", 1)]
    except ValueError:
        await update.message.reply_text("❌ Invalid format. Use: /teach <title> | <content>")
        return
        
    if not title or not content:
        await update.message.reply_text("❌ Title or content is empty. Use: /teach <title> | <content>")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": title, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' already exists in topic: {topic}. "
            "Use /unlearn '{title}' first if you want to replace it."
        )
        return

    doc_id = new_doc_id(topic, "qa")
    try:
        col.add(
            ids=[doc_id],
            metadatas=[{"title": title, "topic": topic, "type": "qa", "source": "manual"}],
            documents=[safe_text_for_embedding(content)],
        )
        logging.info(f"Successfully added document '{title}' to topic '{topic}'")
        await update.message.reply_text(
            f"Learned '{title}' ✅ (topic: {topic}, mode: Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add document to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save: {str(e)[:200]}")

# ---------- TEACH RUBRIC (EVALUATION sources) ----------
async def teachrubric(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach global rubrics.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user

    raw = strip_command(extract_command_text(update), "teachrubric")
    logging.info(f"Raw teachrubric input: '{raw}'")
    
    if not raw:
        await update.message.reply_text(
            "Use format:\n/teachrubric <title> | <rubric / evaluation criteria>"
        )
        return
        
    if "|" not in raw:
        await update.message.reply_text(
            "Use format:\n/teachrubric <title> | <rubric / evaluation criteria>"
        )
        return

    try:
        title, content = [p.strip() for p in raw.split("|", 1)]
    except ValueError:
        await update.message.reply_text("❌ Invalid format. Use: /teachrubric <title> | <rubric>")
        return
        
    if not title or not content:
        await update.message.reply_text("❌ Title or rubric content is empty.")
        return

    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachrubric")

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": title, "type": "evaluation"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' already exists in topic: {topic}. "
            "Use /unlearn '{title}' first if you want to replace it."
        )
        return

    try:
        col.add(
            ids=[new_doc_id(topic, "eval")],
            metadatas=[
                {"title": title, "topic": topic, "type": "evaluation", "source": "manual"}
            ],
            documents=[safe_text_for_embedding(content)],
        )
        logging.info(f"Successfully added rubric '{title}' to topic '{topic}'")
        await update.message.reply_text(
            f"Learned evaluation rubric '{title}' ✅ (topic: {topic}, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add rubric to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save rubric: {str(e)[:200]}")

# ---------- TEACH FILE (Q&A sources) ----------
async def teachfile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach from files.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachfile")

    doc = update.message.document
    if not doc:
        await update.message.reply_text(
            "Attach a PDF or DOCX and write /teachfile in the caption to train me from it."
        )
        return

    await update.message.reply_text(
        "Reading your file and extracting text to learn from it (Q&A)…"
    )

    try:
        tgfile = await doc.get_file()
        file_bytes = await tgfile.download_as_bytearray()
        name = (doc.file_name or "upload").lower()
    except Exception as e:
        await update.message.reply_text(f"Failed to download file: {e}")
        return

    try:
        if name.endswith(".pdf"):
            text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for /teachfile.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    if not text or not text.strip():
        await update.message.reply_text("I couldn't find any readable text in that file.")
        return

    parts = _chunk(text)
    if not parts:
        await update.message.reply_text("Text was too short or could not be chunked.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": name, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{name}' is already learned in topic: {topic}.\n"
            "Use /unlearn <title> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "qa") for _ in range(len(parts))]
        metas = [
            {"title": name, "topic": topic, "part": i, "source": "file", "type": "qa"}
            for i in range(len(parts))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(p) for p in parts])
        logging.info(f"Successfully added file '{name}' with {len(parts)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned from file ✅ ({len(parts)} parts) in topic: {topic} (Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add file to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save file content: {str(e)[:200]}")

# ---------- TEACH FILE EVAL (EVALUATION sources) ----------
async def teachfile_eval(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text(
            "⛔ You are not allowed to teach evaluation rubrics from files."
        )
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachfile_eval")

    doc = update.message.document
    if not doc:
        await update.message.reply_text(
            "Attach a PDF or DOCX and write /teachfile_eval in the caption to teach an evaluation rubric."
        )
        return

    await update.message.reply_text(
        "Reading your rubric file and extracting evaluation criteria…"
    )

    try:
        tgfile = await doc.get_file()
        file_bytes = await tgfile.download_as_bytearray()
        name = (doc.file_name or "upload").lower()
    except Exception as e:
        await update.message.reply_text(f"Failed to download file: {e}")
        return

    try:
        if name.endswith(".pdf"):
            text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for /teachfile_eval.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    if not text or not text.strip():
        await update.message.reply_text("I couldn't find any readable text in that file.")
        return

    parts = _chunk(text)
    if not parts:
        await update.message.reply_text("Text was too short or could not be chunked.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": name, "type": "evaluation"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{name}' is already learned in topic: {topic}.\n"
            "Use /unlearn <title> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "eval") for _ in range(len(parts))]
        metas = [
            {"title": name, "topic": topic, "part": i, "source": "file", "type": "evaluation"}
            for i in range(len(parts))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(p) for p in parts])
        logging.info(f"Successfully added eval file '{name}' with {len(parts)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned evaluation rubric from file ✅ ({len(parts)} parts) in topic: {topic} (scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add eval file to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save rubric content: {str(e)[:200]}")

# ---------- TEACH LINK (Q&A) ----------
async def teachlink(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach from links.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachlink")

    msg_text = extract_command_text(update)
    rest = strip_command(msg_text, "teachlink")
    if not rest:
        await update.message.reply_text("Use: /teachlink <url>")
        return

    url = rest.strip()
    await update.message.reply_text("Fetching content from link and learning from it (Q&A)…")

    try:
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            await update.message.reply_text("Could not fetch the URL. Please check the link.")
            return
        text = trafilatura.extract(downloaded)
    except Exception as e:
        await update.message.reply_text(f"Error while fetching the URL: {e}")
        return

    if not text:
        await update.message.reply_text("Couldn't extract readable text from that link.")
        return

    chunks = _chunk(text)
    if not chunks:
        await update.message.reply_text("The page did not contain enough text to learn from.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": url, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"This link is already learned in topic: {topic}.\nUse /unlearn <url> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "qa") for _ in range(len(chunks))]
        metas = [
            {"title": url, "topic": topic, "part": i, "source": "link", "type": "qa"}
            for i in range(len(chunks))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(c) for c in chunks])
        logging.info(f"Successfully added link '{url}' with {len(chunks)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned from link ✅ ({len(chunks)} parts) in topic: {topic} (Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add link to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save link content: {str(e)[:200]}")

# ---------- TEACH LINK EVAL (EVALUATION sources) ----------
async def teachlink_eval(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text(
            "⛔ You are not allowed to teach evaluation material from links."
        )
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachlink_eval")

    msg_text = extract_command_text(update)
    rest = strip_command(msg_text, "teachlink_eval")
    if not rest:
        await update.message.reply_text("Use: /teachlink_eval <url>")
        return

    url = rest.strip()
    await update.message.reply_text(
        "Fetching content from link and learning it as evaluation / rubric material…"
    )

    try:
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            await update.message.reply_text("Could not fetch the URL. Please check the link.")
            return
        text = trafilatura.extract(downloaded)
    except Exception as e:
        await update.message.reply_text(f"Error while fetching the URL: {e}")
        return

    if not text:
        await update.message.reply_text("Couldn't extract readable text from that link.")
        return

    chunks = _chunk(text)
    if not chunks:
        await update.message.reply_text("The page did not contain enough text to learn from.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": url, "type": "evaluation"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"This link is already learned in topic: {topic}.\nUse /unlearn <url> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "eval") for _ in range(len(chunks))]
        metas = [
            {"title": url, "topic": topic, "part": i, "source": "link", "type": "evaluation"}
            for i in range(len(chunks))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(c) for c in chunks])
        logging.info(f"Successfully added eval link '{url}' with {len(chunks)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned evaluation material from link ✅ ({len(chunks)} parts) in topic: {topic} (scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add eval link to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save evaluation content: {str(e)[:200]}")

# ---------- TEACH IMAGE (Q&A) ----------
async def teachimage(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach from images.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachimage")

    photos = update.message.photo or []
    if not photos:
        await update.message.reply_text("Please send a clear image with caption:\n/teachimage <title>")
        return

    caption = (update.message.caption or "").strip()
    title = None
    if is_caption_command(caption, "teachimage"):
        rest = strip_command(caption, "teachimage")
        title = rest.strip() if rest.strip() else None

    largest = photos[-1]
    try:
        tgfile = await largest.get_file()
        if not title:
            title = f"image_{tgfile.file_unique_id}"
    except Exception as e:
        await update.message.reply_text(f"Failed to get image file: {e}")
        return

    await update.message.reply_text(
        f"Reading your image for topic '{topic}' and extracting text to learn from it…"
    )

    try:
        img_bytes = await tgfile.download_as_bytearray()
    except Exception as e:
        await update.message.reply_text(f"Failed to download image: {e}")
        return

    try:
        extracted = await extract_text_from_image_bytes(img_bytes)
    except Exception as e:
        await update.message.reply_text(f"Could not extract text from image: {e}")
        return

    if not extracted.strip():
        await update.message.reply_text("I couldn't read any text from that image.")
        return

    parts = _chunk(extracted)
    if not parts:
        await update.message.reply_text("The extracted text was too short to learn from.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": title, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' is already learned in topic: {topic}.\nUse /unlearn <title> to remove it first if needed."
        )
        return

    try:
        ids = [new_doc_id(topic, "qa") for _ in range(len(parts))]
        metas = [
            {"title": title, "topic": topic, "part": i, "source": "image", "type": "qa"}
            for i in range(len(parts))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(p) for p in parts])
        logging.info(f"Successfully added image '{title}' with {len(parts)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned from image '{title}' ✅ ({len(parts)} parts) in topic: {topic} (Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add image to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save image content: {str(e)[:200]}")

# ---------- SOURCES ----------
async def sources_all(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logging.info("🔥 /sources_all handler hit")
    try:
        collections = chroma.list_collections()
    except Exception as e:
        await update.message.reply_text(f"Error accessing database: {e}")
        return

    if not collections:
        await update.message.reply_text("No sources stored yet.")
        return

    source_stats = {}
    total_bytes = 0
    total_chunks = 0

    for col_info in collections:
        col_name = col_info.name if hasattr(col_info, "name") else str(col_info)
        if not col_name:
            continue
            
        try:
            col = chroma.get_collection(col_name)
            data = col.get(include=["documents", "metadatas"])
        except Exception as e:
            logging.error(f"Error getting collection {col_name}: {e}")
            continue

        docs = data.get("documents", [])
        metas = data.get("metadatas", [])

        for doc, meta in zip(docs, metas):
            title = (meta or {}).get("title", "Untitled")
            size_bytes = len((doc or "").encode("utf-8"))

            if title not in source_stats:
                source_stats[title] = {"chunks": 0, "bytes": 0}

            source_stats[title]["chunks"] += 1
            source_stats[title]["bytes"] += size_bytes

            total_chunks += 1
            total_bytes += size_bytes

    if not source_stats:
        await update.message.reply_text("No sources found in any collections.")
        return

    lines = []
    for title, stats in sorted(source_stats.items(), key=lambda x: x[1]["bytes"], reverse=True):
        mb = stats["bytes"] / (1024 * 1024)
        lines.append(f"• {title}: {stats['chunks']} chunks, {mb:.2f} MB")

    total_mb = total_bytes / (1024 * 1024)
    msg = (
        "📚 ALL BOT SOURCES (ALL TOPICS, GLOBAL)\n\n"
        f"Total chunks: {total_chunks}\n"
        f"Total text size: {total_mb:.2f} MB\n\n"
        + "\n".join(lines)
    )
    await send_long(update, msg)

async def sources(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logging.info("🔥 /sources handler hit")
    chat_id = update.effective_chat.id
    topic = get_current_topic(context)
    
    col = None

    try:
        col = get_collection(chat_id, topic)
        data = col.get(include=["metadatas"])
    except Exception as e:
        logging.error(f"Error getting collection: {e}")
        await update.message.reply_text(f"Error accessing database: {e}")
        return
        
    metas = data.get("metadatas") or []

    if not metas:
        await update.message.reply_text(
            f"📌 Active topic: {topic}\nNo sources yet. (Global collection is empty.)"
        )
        return

    qa_titles = []
    eval_titles = []
    for m in metas:
        title = (m or {}).get("title", "Untitled")
        source_type = (m or {}).get("type", "qa")
        if source_type == "evaluation":
            eval_titles.append(title)
        else:
            qa_titles.append(title)

    qa_titles = list(dict.fromkeys(qa_titles))
    eval_titles = list(dict.fromkeys(eval_titles))

    lines = [f"📌 Active topic: {topic} (GLOBAL)\n"]
    if eval_titles:
        lines.append("📘 Evaluation Rubrics:")
        for t in eval_titles:
            lines.append(f"• {t}")
        lines.append("")
    if qa_titles:
        lines.append("📗 Q&A Sources:")
        for t in qa_titles:
            lines.append(f"• {t}")

    await update.message.reply_text("\n".join(lines))

# ---------- UNLEARN ----------
async def unlearn(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to remove global sources.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="unlearn")

    text = extract_command_text(update)
    rest = strip_command(text, "unlearn")
    if not rest:
        await update.message.reply_text("Usage: /unlearn <exact title shown in /sources>")
        return

    title = rest.strip()
    
    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return
        
    try:
        to_delete = col.get(where={"title": title})
    except Exception as e:
        logging.error(f"Error checking for documents to delete: {e}")
        await update.message.reply_text(f"❌ Error checking for documents: {e}")
        return
        
    removed = len((to_delete or {}).get("ids") or [])

    if removed == 0:
        await update.message.reply_text(
            f"No source titled '{title}' found in topic: {topic} (GLOBAL)."
        )
        return

    try:
        col.delete(where={"title": title})
        await update.message.reply_text(
            f"Removed '{title}' ✅ ({removed} parts) from topic: {topic} (GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Error deleting documents: {e}")
        await update.message.reply_text(f"❌ Failed to remove '{title}': {e}")

# ---------- CLEAR ----------
async def clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to clear global knowledge.")
        return

    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="clear")

    collection_name = f"{COLLECTION_PREFIX}_{topic}"
    try:
        chroma.delete_collection(collection_name)
        await update.message.reply_text(f"Forgot everything 🧹 (topic: {topic}, scope: GLOBAL)")
    except Exception as e:
        logging.warning(f"Could not delete collection {collection_name}: {e}")
        await update.message.reply_text(
            f"Didn't find any stored sources to clear for topic: {topic} (GLOBAL)."
        )

# ---------- EVALUATION HELPERS ----------
def _eval_context_queries(topic: str, extra_query: str = "") -> tuple[list[str], list[str]]:
    extra = (extra_query or "").strip()
    topic_map = {
        "essays_personal": (
            ["personal statement rubric", "essay evaluation criteria", "voice reflection vulnerability", extra],
            ["personal statement examples", "essay tips authenticity reflection", extra],
        ),
        "essays_supplemental": (
            ["supplemental essay rubric", "why us evaluation criteria", "fit specificity contribution", extra],
            ["supplemental essay examples", "why us tips specificity", extra],
        ),
        "extracurriculars": (
            ["extracurricular evaluation rubric", "activities list impact leadership initiative", "activity description quantification", extra],
            ["extracurricular examples", "activities list tips action verbs impact", extra],
        ),
        "portfolio": (
            ["portfolio review rubric", "originality technical quality curation presentation", "creative portfolio evaluation", extra],
            ["portfolio examples", "portfolio presentation tips process cohesion", extra],
        ),
        "recommendations": (
            ["recommendation letter rubric", "specificity credibility comparison evidence", extra],
            ["recommendation examples", "teacher recommendation tips specificity", extra],
        ),
        "ielts_writing": (
            ["IELTS writing rubric", "task response coherence lexical grammar", extra],
            ["IELTS writing examples", "IELTS writing tips band descriptors", extra],
        ),
    }
    eval_queries, qa_queries = topic_map.get(
        topic,
        (["evaluation criteria", "guidelines", "rubric", extra], ["tips", "examples", "advice", extra]),
    )
    return [q for q in eval_queries if q], [q for q in qa_queries if q]


def _eval_context_from_collection(col, topic: str, extra_query: str = "", context: ContextTypes.DEFAULT_TYPE | None = None):
    eval_docs = []
    qa_docs = []
    debug_items = []
    eval_queries, qa_queries = _eval_context_queries(topic, extra_query)

    try:
        res_eval = col.query(
            query_texts=eval_queries,
            where={"type": "evaluation"},
            n_results=6,
        )
        eval_docs = res_eval.get("documents", [[]])[0]
        eval_metas = res_eval.get("metadatas", [[]])[0]
        for doc, meta in zip(eval_docs or [], eval_metas or []):
            meta = meta or {}
            debug_items.append({
                "bucket": "evaluation",
                "title": meta.get("title", "Untitled"),
                "type": meta.get("type", "evaluation"),
                "part": meta.get("part", "—"),
                "preview": (doc or "")[:160],
            })
    except Exception as e:
        logging.error(f"Error querying evaluation docs: {e}")
        eval_docs = []

    try:
        res_qa = col.query(
            query_texts=qa_queries,
            where={"type": "qa"},
            n_results=6,
        )
        qa_docs = res_qa.get("documents", [[]])[0]
        qa_metas = res_qa.get("metadatas", [[]])[0]
        for doc, meta in zip(qa_docs or [], qa_metas or []):
            meta = meta or {}
            debug_items.append({
                "bucket": "qa",
                "title": meta.get("title", "Untitled"),
                "type": meta.get("type", "qa"),
                "part": meta.get("part", "—"),
                "preview": (doc or "")[:160],
            })
    except Exception as e:
        logging.error(f"Error querying QA docs: {e}")
        qa_docs = []

    if context is not None:
        _save_last_retrieval(
            context,
            topic=topic,
            mode="evaluation_context",
            query=f"eval_queries={eval_queries} | qa_queries={qa_queries}",
            items=debug_items,
        )

    docs = (eval_docs or []) + (qa_docs or [])
    return "\n\n---\n\n".join(docs) if docs else ""

async def run_eval_followup(update: Update, context: ContextTypes.DEFAULT_TYPE, user_request: str):
    """Apply prior eval feedback to revise the saved text (coach mode)."""
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user

    topic = context.user_data.get("last_eval_topic") or get_current_topic(context)
    pretty_topic = _pretty_topic_for_eval(topic)

    student_text = (context.user_data.get("last_eval_text") or "").strip()
    prior_feedback = (context.user_data.get("last_eval_feedback") or "").strip()

    if not student_text:
        await update.message.reply_text(
            "✅ Evaluation follow-up is ON, but I don't have your last text saved.\n\n"
            "Please paste the full text again (or upload PDF/DOCX), then I'll apply feedback."
        )
        return

    record_event(user.id, topic, kind="eval_followup")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, topic=topic, extra_query=pretty_topic, context=context)
    except Exception as e:
        logging.error(f"Error getting collection for eval followup: {e}")
        context_block = ""

    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)

    decision_notes = coach_decision_notes(topic, student_text, mem)

    sys = (
        "You are UniVenture Coach - an elite admissions coach.\n"
        "Task: revise the student's text by APPLYING the prior feedback.\n"
        "Rules:\n"
        "- Keep the student's voice.\n"
        "- If the user request is vague, prioritize: smoother transitions + deeper reflection + stronger ending tied to the opening.\n"
        "- Output ONLY the revised text. No commentary.\n\n"
        "Student memory (for voice consistency):\n"
        f"{memory_summary_for_prompt(mem)}\n\n"
        "Decision notes:\n"
        f"{decision_notes}\n"
    )
    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {"role": "system", "content": f"Prior feedback to apply:\n{prior_feedback}"},
        {"role": "system", "content": f"Text to revise:\n{student_text}"},
        {"role": "user", "content": f"User request:\n{user_request}"},
    ]

    revised = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.35)

    if revised and not revised.lower().startswith("error"):
        context.user_data["last_eval_text"] = revised
        await send_long(update, revised)
    else:
        await update.message.reply_text("❌ Failed to generate revision. Please try again.")


async def run_eval_qa(update: Update, context: ContextTypes.DEFAULT_TYPE, user_question: str):
    """Answer follow-up questions about the already-evaluated text (coach mode)."""
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user

    topic = context.user_data.get("last_eval_topic") or get_current_topic(context)
    pretty_topic = _pretty_topic_for_eval(topic)

    student_text = (context.user_data.get("last_eval_text") or "").strip()
    prior_feedback = (context.user_data.get("last_eval_feedback") or "").strip()

    if not student_text:
        await update.message.reply_text(
            "✅ Evaluation mode is ON, but I don't have your last submission saved.\n\n"
            "Please paste the full text again (or upload PDF/DOCX)."
        )
        return

    record_event(user.id, topic, kind="eval_qa")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, topic=topic, extra_query=pretty_topic, context=context)
    except Exception as e:
        logging.error(f"Error getting collection for eval QA: {e}")
        context_block = ""

    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    decision_notes = coach_decision_notes(topic, student_text, mem)

    sys = (
        "You are UniVenture Coach, a senior admissions mentor.\n"
        "The user is asking a follow-up question about a text you already evaluated.\n"
        "Rules:\n"
        "- You DO have access to the student's text below.\n"
        "- Answer directly and specifically.\n"
        "- If they ask about grammar, show 2–5 short corrected excerpts.\n"
        "- Do NOT say you cannot see the essay.\n"
        "- End with one next step.\n\n"
        "Student memory:\n"
        f"{memory_summary_for_prompt(mem)}\n\n"
        "Decision notes:\n"
        f"{decision_notes}\n"
    )

    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {"role": "system", "content": f"Prior evaluation feedback:\n{prior_feedback}"},
        {"role": "system", "content": f"Student's {pretty_topic}:\n{student_text}"},
        {"role": "user", "content": user_question},
    ]

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.35)
    await send_long(update, a)


async def _coach_evaluate_common(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    topic: str,
    student_text_for_eval: str,
    student_text_for_followup: str,
    context_block: str,
    pretty_topic: str,
):
    """Shared evaluation engine with Coach Mode + memory updates."""
    # ===== Daily eval limit =====
    if not can_run_eval(context):
        await update.message.reply_text(
            f"⚠️ Daily limit reached: {EVALS_PER_DAY_LIMIT} evaluations/day.\n"
            "Try again tomorrow (UTC) or ask follow-up questions on your last evaluation."
        )
        return

    bump_eval_count(context)

    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)

    # Update profile signals from the submission (cheap heuristics)
    mem["profile"] = extract_profile_signals(student_text_for_followup, mem.get("profile", {}) or {})

    decision_notes = coach_decision_notes(topic, student_text_for_followup, mem)

    # Choose system prompt
    if topic in {"essays_personal", "essays_supplemental", "extracurriculars", "recommendations", "portfolio", "ielts_writing"}:
        sys_role = coach_eval_system_prompt(topic, mem, decision_notes)
    else:
        sys_role = coach_eval_system_prompt(topic, mem, decision_notes)

    messages = [
        {"role": "system", "content": sys_role},
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {"role": "user", "content": f"Here is the student's {pretty_topic}. Evaluate it:\n\n{student_text_for_eval}"},
    ]

# Optional speed UX: send a short quick feedback first
    if ENABLE_EVAL_QUICK_PREVIEW:
        quick_messages = [
            {"role": "system", "content":
                "QUICK DIAGNOSTIC (STRICT):\n"
                "- Output EXACTLY 3 bullet points.\n"
                "- Each bullet must start with one of these labels exactly: DIFFERENTIATION:, RISK:, DEPTH:.\n"
                "- No other text, no headings, no 'Overall', no 'What works', no 'Next step'.\n"
                "- Keep each bullet to 1 sentence."
            },
            {"role": "user", "content": f"Quick feedback for the student's {pretty_topic}:\n\n{student_text_for_eval}"},
        ]
    
        quick_out = await openai_chat(
            FAST_MODEL,
            quick_messages,
            0.2,
            EVAL_QUICK_MAX_TOKENS
        )
    
        if (quick_out or "").strip():
            await send_long(update, "⚡ Quick feedback (full review is coming next):\n\n" + quick_out.strip())
    
    
    # Full strategic evaluation (non-blocking)
    raw_out = await openai_chat(
        STRONG_MODEL,
        messages,
        0.3
    )
    
    user_text, mem_update = split_memory_block(raw_out)

    # Save eval context for follow-ups
    set_eval_context(context, topic, student_text_for_followup, user_text)
    context.user_data["my_eval_count"] = int(context.user_data.get("my_eval_count", 0) or 0) + 1

    # Update persistent memory
    mem["history"]["message_count"] = int(mem["history"].get("message_count", 0) or 0)
    mem["history"]["eval_count"] = int(mem["history"].get("eval_count", 0) or 0) + 1
    mem["history"]["last_topic"] = topic
    mem["history"]["last_active"] = int(__import__("time").time())

    apply_memory_update(mem, mem_update or {}, topic)
    persist_user_memory(update, context)

    await send_long(update, user_text)


async def evaluate_file_for_topic(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if not can_run_eval(context):
        await update.message.reply_text(
            f"⚠️ Daily limit reached: {EVALS_PER_DAY_LIMIT} evaluations/day.\n"
            "Try again tomorrow (UTC) or ask follow-up questions on your last evaluation."
        )
        return

    await show_typing(update, context)

    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="eval")

    allowed = set(EVAL_TOPICS) | {"ielts_writing"}
    if topic not in allowed:
        await update.message.reply_text(
            "To evaluate an essay, recommendation, EC description, portfolio, or IELTS Writing, "
            "choose the correct menu first, then send the file."
        )
        return

    doc = update.message.document
    if not doc:
        await update.message.reply_text("Please attach a PDF or DOCX file.")
        return

    pretty_topic = _pretty_topic_for_eval(topic)
    await update.message.reply_text(f"Reading your {pretty_topic} file…")

    try:
        tgfile = await doc.get_file()
        file_bytes = await tgfile.download_as_bytearray()
        name = (doc.file_name or "document").lower()
    except Exception as e:
        await update.message.reply_text(f"Failed to download file: {e}")
        return

    try:
        if name.endswith(".pdf"):
            full_text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            full_text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for evaluation.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    if not (full_text or "").strip():
        await update.message.reply_text("I couldn't read enough text from that file to evaluate.")
        return

    parts = _chunk(full_text, max_chars=1500)
    student_text_for_eval = "\n\n---\n\n".join(parts[:5]) if parts else _truncate_for_storage(full_text, 4000)
    student_text_for_followup = _truncate_for_storage(full_text, 12000)

    await update.message.reply_text(f"Analyzing your {pretty_topic} (Coach Mode)…")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, topic=topic, extra_query=pretty_topic, context=context)
    except Exception as e:
        logging.error(f"Error getting collection for evaluation: {e}")
        context_block = ""

    await _coach_evaluate_common(
        update=update,
        context=context,
        topic=topic,
        student_text_for_eval=student_text_for_eval,
        student_text_for_followup=student_text_for_followup,
        context_block=context_block,
        pretty_topic=pretty_topic,
    )


async def evaluate_ielts_writing_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # ===== Daily eval limit =====
    if not can_run_eval(context):
        await update.message.reply_text(
            f"⚠️ Daily limit reached: {EVALS_PER_DAY_LIMIT} evaluations/day.\n"
            "Try again tomorrow (UTC) or ask follow-up questions on your last evaluation."
        )
        return
    
    bump_eval_count(context)

    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = "ielts_writing"
    record_event(user.id, topic, kind="eval")

    photos = update.message.photo or []
    if not photos:
        await update.message.reply_text("Please send a clear photo of your IELTS Writing answer.")
        return

    largest = photos[-1]
    try:
        tgfile = await largest.get_file()
    except Exception as e:
        await update.message.reply_text(f"Failed to get image file: {e}")
        return
        
    await update.message.reply_text("Reading your IELTS Writing answer from the image…")
    
    try:
        img_bytes = await tgfile.download_as_bytearray()
    except Exception as e:
        await update.message.reply_text(f"Failed to download image: {e}")
        return

    try:
        extracted = await extract_text_from_image_bytes(img_bytes)
    except Exception as e:
        await update.message.reply_text(f"Could not extract text from image: {e}")
        return

    if not extracted.strip():
        await update.message.reply_text(
            "I couldn't read enough text from that image. Please try a clearer photo."
        )
        return

    student_text_for_followup = _truncate_for_storage(extracted, 12000)
    parts = _chunk(extracted, max_chars=1500)
    student_text_for_eval = "\n\n---\n\n".join(parts[:5]) if parts else _truncate_for_storage(extracted, 4000)

    await update.message.reply_text("Analyzing your IELTS Writing answer…")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, topic=topic, extra_query="IELTS Writing", context=context)
    except Exception as e:
        logging.error(f"Error getting collection for IELTS evaluation: {e}")
        context_block = ""

    sys_role = (
        "You are an experienced IELTS Writing examiner. "
        "Evaluate the student's writing according to IELTS Academic/General Writing band descriptors. "
        "Comment on Task Response, Coherence and Cohesion, Lexical Resource, and Grammatical Range and Accuracy. "
        "Give an approximate band score (like 6.0, 6.5, 7.0) and then clear, actionable feedback."
    )

    messages = [
        {"role": "system", "content": sys_role},
        {"role": "system", "content": f"IELTS writing rubrics and notes (may be empty):\n{context_block}"},
        {"role": "user", "content": f"Here is the student's IELTS Writing answer (from an image):\n\n{student_text_for_eval}"},
    ]

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.3)
    set_eval_context(context, topic, student_text_for_followup, a)
    await send_long(update, a)

async def evaluate_text_for_topic(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if not can_run_eval(context):
        await update.message.reply_text(
            f"⚠️ Daily limit reached: {EVALS_PER_DAY_LIMIT} evaluations/day.\n"
            "Try again tomorrow (UTC) or ask follow-up questions on your last evaluation."
        )
        return

    await show_typing(update, context)

    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)

    record_event(user.id, topic, kind="eval")

    allowed = set(EVAL_TOPICS) | {"ielts_writing"}
    if topic not in allowed:
        await update.message.reply_text(
            "Text evaluation is only available for Personal Statement, Supplemental Essays, "
            "Extracurriculars, Recommendation Letters, Portfolio, and IELTS Writing.\n\n"
            "Choose the correct topic first, then tap the Evaluation button again."
        )
        return

    raw = (update.message.text or "").strip()
    for marker in [BTN_PS_EVAL, BTN_SUPP_EVAL, BTN_EC_EVAL, BTN_REC_EVAL, BTN_IW_EVAL, BTN_PORT_EVAL]:
        raw = raw.replace(marker, "").strip()

    if len(raw) < 100:
        await update.message.reply_text(
            "The text is too short to evaluate. Please paste the full essay/letter/description."
        )
        return

    pretty_topic = _pretty_topic_for_eval(topic)
    student_text_for_followup = _truncate_for_storage(raw, 12000)
    parts = _chunk(raw, max_chars=1500)
    student_text_for_eval = "\n\n---\n\n".join(parts[:5])

    await update.message.reply_text(f"Evaluating your {pretty_topic} (Coach Mode)…")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, topic=topic, extra_query=pretty_topic, context=context)
    except Exception as e:
        logging.error(f"Error getting collection for evaluation: {e}")
        context_block = ""

    await _coach_evaluate_common(
        update=update,
        context=context,
        topic=topic,
        student_text_for_eval=student_text_for_eval,
        student_text_for_followup=student_text_for_followup,
        context_block=context_block,
        pretty_topic=pretty_topic,
    )


# ---------- DOCUMENT & PHOTO ROUTERS ----------
async def document_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    caption = (update.message.caption or "").strip()
    
    # Check for commands in caption FIRST
    if caption:
        if is_caption_command(caption, "teach"):
            await teach(update, context)
            return
        if is_caption_command(caption, "teachrubric"):
            await teachrubric(update, context)
            return
        if is_caption_command(caption, "teachfile_eval"):
            await teachfile_eval(update, context)
            return
        if is_caption_command(caption, "teachfile"):
            await teachfile(update, context)
            return
        if is_caption_command(caption, "teachlink_eval"):
            await teachlink_eval(update, context)
            return
        if is_caption_command(caption, "teachlink"):
            await teachlink(update, context)
            return

    # If no command in caption, check if we're in evaluation mode
    topic = get_current_topic(context)
    if topic in (set(EVAL_TOPICS) | {"ielts_writing"}):
        await evaluate_file_for_topic(update, context)
        return

    await update.message.reply_text(
        "If you want me to LEARN from this file, send it again and write /teachfile, "
        "/teachfile_eval, /teachlink, or /teachlink_eval in the caption.\n\n"
        "If this is an essay, recommendation, EC description, portfolio, or IELTS Writing for feedback, "
        "choose the correct topic and tap its Evaluation button."
    )

async def photo_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    caption = (update.message.caption or "").strip()
    topic = get_current_topic(context)
    user = update.effective_user
    record_event(user.id, topic, kind="photo")

    # Check for commands in caption FIRST
    if caption:
        if is_caption_command(caption, "teach"):
            await teach(update, context)
            return
        if is_caption_command(caption, "teachrubric"):
            await teachrubric(update, context)
            return
        if is_caption_command(caption, "teachimage"):
            await teachimage(update, context)
            return

    if topic == "ielts_writing":
        await evaluate_ielts_writing_image(update, context)
        return

    await update.message.reply_text(
        "I can learn from images too 🤖🖼\n\n"
        "If you want me to *learn* from this image (e.g., essay screenshot, rubric), "
        "send it again with caption:\n\n"
        "/teachimage <title>\n\n"
        "For IELTS Writing evaluation from an image, switch to IELTS Writing first."
    )

# ---------- STATS ----------
async def stats_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    stats = load_stats()
    total_users = len(stats.get("users", []))
    total_msgs = stats.get("messages_total", 0)
    topic_counts = stats.get("topic_counts", {})
    eval_counts = stats.get("eval_counts", {})

    if topic_counts:
        top_topics = sorted(topic_counts.items(), key=lambda kv: kv[1], reverse=True)[:5]
        topics_str = "\n".join(f"- {t}: {c} msgs" for t, c in top_topics)
    else:
        topics_str = "No topic data yet."

    total_evals = sum(eval_counts.values()) if eval_counts else 0
    # DAU / MAU (based on activity_by_date; fallback to last_active timestamps)
    abd = stats.get("activity_by_date", {}) or {}
    today = datetime.utcnow().date()
    today_key = today.strftime("%Y-%m-%d")

    if abd:
        dau = len(set(abd.get(today_key, []) or []))

        mau_users = set()
        start = today - timedelta(days=29)
        for k, uids in abd.items():
            try:
                d = datetime.strptime(k, "%Y-%m-%d").date()
            except Exception:
                continue
            if start <= d <= today:
                mau_users.update(uids or [])
        mau = len(mau_users)
    else:
        # Fallback: use per-user `history.last_active` timestamps from user memory files
        try:
            import glob
            now_ts = __import__("time").time()
            dau_cut = now_ts - 24 * 3600
            mau_cut = now_ts - 30 * 24 * 3600

            dau = 0
            mau = 0
            for fp in glob.glob(os.path.join(USER_MEM_DIR, "*.json")):
                try:
                    with open(fp, "r", encoding="utf-8") as f:
                        mem = json.load(f)
                    ts = int(((mem.get("history") or {}).get("last_active") or 0) or 0)
                    if ts >= dau_cut:
                        dau += 1
                    if ts >= mau_cut:
                        mau += 1
                except Exception:
                    continue
        except Exception:
            dau = 0
            mau = 0

    msg = (
        f"📊 Bot analytics\n"
        f"- Unique users: {total_users}\n"
        f"- DAU (today): {dau}\n"
        f"- MAU (last 30d): {mau}\n"
        f"- Total interactions (events): {total_msgs}\n"
        f"- Total evaluations: {total_evals}\n\n"
        f"Top topics:\n{topics_str}"
    )
    await update.message.reply_text(msg)

# ---------- BACKUP SOURCES ----------
async def backup_sources(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ Admin only.")
        return

    await update.message.reply_text("📦 Backing up all sources…")

    data = {}
    try:
        collections = chroma.list_collections()
    except Exception as e:
        await update.message.reply_text(f"Error accessing database: {e}")
        return

    for col_info in collections:
        col_name = col_info.name if hasattr(col_info, "name") else str(col_info)
        if not col_name:
            continue
            
        try:
            col = chroma.get_collection(col_name)
            payload = col.get(include=["documents", "metadatas", "ids"])
            data[col_name] = payload
        except Exception as e:
            logging.error(f"Error backing up collection {col_name}: {e}")
            continue

    path = os.path.join(DATA_DIR, "backup_sources.json")
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        await update.message.reply_text(
            f"✅ Backup completed.\nSaved to:\n{path}\n\nYou can download it from your server volume."
        )
    except Exception as e:
        logging.error(f"Error saving backup: {e}")
        await update.message.reply_text(f"❌ Failed to save backup: {e}")

# ---------- HEALTH CHECK ----------
async def health(update: Update, context: ContextTypes.DEFAULT_TYPE):
    checks = []

    try:
        _ = await openai_chat(
            model=FAST_MODEL,
            messages=[{"role": "user", "content": "ping"}],
            temperature=0.0,
            max_tokens=5,
        )
        checks.append("✅ OpenAI: OK")
    except Exception as e:
        checks.append(f"❌ OpenAI: {str(e)[:100]}")

    try:
        test_col = chroma.get_or_create_collection("health_check", embedding_function=emb_fn)
        _id = uuid.uuid4().hex
        test_col.add(ids=[_id], documents=["pong"], metadatas=[{"type": "health"}])
        try:
            test_col.delete(ids=[_id])
        except Exception:
            pass
        checks.append("✅ Chroma: writable")
    except Exception as e:
        checks.append(f"❌ Chroma: {str(e)[:100]}")

    try:
        test_path = os.path.join(DATA_DIR, "health.txt")
        with open(test_path, "w", encoding="utf-8") as f:
            f.write("ok")
        os.remove(test_path)
        checks.append(f"✅ Volume: writable ({DATA_DIR})")
    except Exception as e:
        checks.append(f"❌ Volume: {str(e)[:100]}")

    await update.message.reply_text("🧪 Health Check\n\n" + "\n".join(checks))

# ---------- NEW FEATURE HELPERS ----------
# ===== Daily usage limits =====
EVALS_PER_DAY_LIMIT = int(os.getenv("EVALS_PER_DAY_LIMIT", "5"))

def _today_key_utc() -> str:
    return datetime.utcnow().strftime("%Y-%m-%d")

def can_run_eval(context: ContextTypes.DEFAULT_TYPE) -> bool:
    usage = context.user_data.setdefault("usage_limits", {})
    today = _today_key_utc()

    rec = usage.get(today) or {"evals": 0}
    # reset if date changed
    if usage.get("date") != today:
        usage.clear()
        usage["date"] = today
        usage[today] = rec

    return int(rec.get("evals", 0) or 0) < EVALS_PER_DAY_LIMIT

def bump_eval_count(context: ContextTypes.DEFAULT_TYPE) -> int:
    usage = context.user_data.setdefault("usage_limits", {})
    today = _today_key_utc()

    if usage.get("date") != today:
        usage.clear()
        usage["date"] = today
        usage[today] = {"evals": 0}

    rec = usage[today]
    rec["evals"] = int(rec.get("evals", 0) or 0) + 1
    return rec["evals"]

def _app_portfolio_compact_summary(mem: dict) -> str:
    """
    Create a compact, LLM-friendly summary of the user's Application Portfolio.
    Used for School Finder & Application Plan when 'use my portfolio' is selected.
    """
    if not mem or "portfolio" not in mem:
        return ""

    p = mem.get("portfolio", {}) or {}
    parts = []

    # Academics
    if p.get("grade"):
        parts.append(f"Grade: {p['grade']}")
    if p.get("gpa"):
        parts.append(f"GPA: {p['gpa']}")
    if p.get("sat"):
        parts.append(f"SAT: {p['sat']}")
    if p.get("ielts"):
        parts.append(f"IELTS: {p['ielts']}")

    # Preferences
    if p.get("major"):
        parts.append(f"Intended major: {p['major']}")
    if p.get("target_countries"):
        parts.append("Target countries: " + ", ".join(p["target_countries"]))
    if p.get("needs_aid"):
        parts.append("Needs financial aid")

    # Strengths & weaknesses from coach memory
    writing = mem.get("writing", {})
    strengths = writing.get("strengths", {})
    issues = writing.get("recurring_issues", {})

    if strengths:
        top_strengths = sorted(strengths, key=strengths.get, reverse=True)[:3]
        parts.append("Key strengths: " + ", ".join(top_strengths))

    if issues:
        top_issues = sorted(issues, key=issues.get, reverse=True)[:3]
        parts.append("Main weaknesses to manage: " + ", ".join(top_issues))

    return "\n".join(parts)

def set_pending_feature(context: ContextTypes.DEFAULT_TYPE, feature: str | None):
    context.user_data["pending_feature"] = feature

def clear_pending_feature(context: ContextTypes.DEFAULT_TYPE):
    context.user_data.pop("pending_feature", None)

def _ensure_application_defaults(mem: dict) -> dict:
    """Ensure mem['application'] exists with expected nested structure."""
    if not isinstance(mem, dict):
        return {}
    app = mem.get("application")
    if not isinstance(app, dict):
        app = {}
        mem["application"] = app


    # Test scores / academics
    app.setdefault("test_scores", {})
    # Essays status
    app.setdefault("essays", {})
    # EC summary (short snapshot)
    app.setdefault("ecs", {})
    # Awards/honors list
    app.setdefault("awards", {})
    # Preferences/constraints used for school suggestions
    app.setdefault("preferences", {})
    # Wellness signals (optional)
    app.setdefault("wellness", {})
    # Internal readiness (computed)
    app.setdefault("readiness", {})

    # --- nested defaults ---
    ecs = app["ecs"]
    if isinstance(ecs, dict):
        ecs.setdefault("summary", None)  # free-text 1-8 lines
        ecs.setdefault("highlights", [])  # short bullets extracted/entered
        ecs.setdefault("updated_at", None)

    awards = app["awards"]
    if isinstance(awards, dict):
        awards.setdefault("items", [])  # list[str]
        awards.setdefault("updated_at", None)

    prefs = app["preferences"]
    if isinstance(prefs, dict):
        prefs.setdefault("target_countries", [])
        prefs.setdefault("major", None)
        prefs.setdefault("budget", None)
        prefs.setdefault("needs_aid", None)
        prefs.setdefault("deadlines", None)
        prefs.setdefault("notes", None)
        prefs.setdefault("updated_at", None)

    well = app["wellness"]
    if isinstance(well, dict):
        well.setdefault("stress_level", None)   # 1-10
        well.setdefault("hours_per_week", None)  # study/app workload
        well.setdefault("sleep_hours", None)     # avg
        well.setdefault("notes", None)
        well.setdefault("updated_at", None)

    ts = app["test_scores"]
    ts.setdefault("gpa", None)
    ts.setdefault("sat", None)
    ts.setdefault("sat_breakdown", None)
    ts.setdefault("act", None)
    ts.setdefault("ielts", None)
    ts.setdefault("toefl", None)
    ts.setdefault("duolingo", None)
    ts.setdefault("notes", "")

    es = app["essays"]
    es.setdefault("personal_statement", None)
    es.setdefault("supplementals", None)
    es.setdefault("common_app", None)
    es.setdefault("notes", "")

    ecs = app["ecs"]
    ecs.setdefault("summary", "")
    ecs.setdefault("spike", None)
    ecs.setdefault("notes", "")

    aw = app["awards"]
    aw.setdefault("items", [])  # list[str]
    aw.setdefault("notes", "")

    pr = app["preferences"]
    pr.setdefault("target_countries", [])
    pr.setdefault("intended_major", None)
    pr.setdefault("budget", None)
    pr.setdefault("needs_aid", None)
    pr.setdefault("constraints", "")

    wl = app["wellness"]
    wl.setdefault("stress_level", None)  # 1-10 (self-report)
    wl.setdefault("hours_per_week", None)
    wl.setdefault("notes", "")

    rd = app["readiness"]
    rd.setdefault("last_score", None)
    rd.setdefault("last_updated", None)

    return app


def has_application_portfolio_data(mem: dict) -> bool:
    """Return True if the user's My Application Portfolio has enough signal to use."""
    app = _ensure_application_defaults(mem)

    ts = app.get("test_scores", {}) or {}
    profile = app.get("profile", {}) or {}
    essays = app.get("essays", {}) or {}
    ecs = app.get("ecs", {}) or {}
    awards = app.get("awards", {}) or {}
    prefs = app.get("preferences", {}) or {}

    # Any of these signals makes the portfolio useful.
    score_signals = any(
        (ts.get("gpa"), ts.get("sat"), ts.get("ielts"), ts.get("toefl"), ts.get("act"))
    )
    profile_signals = bool(profile.get("grade") or profile.get("country"))
    essay_signals = bool((essays.get("ps_topic") or "").strip() or (essays.get("supps_notes") or "").strip())
    ec_signals = bool((ecs.get("summary") or "").strip() or ecs.get("spike"))
    awards_signals = bool(awards.get("items"))
    prefs_signals = bool((prefs.get("intended_major") or "").strip() or (prefs.get("target_countries") or []) or prefs.get("needs_aid") is True)

    return bool(score_signals or profile_signals or essay_signals or ec_signals or awards_signals or prefs_signals)


def application_portfolio_compact_summary(mem: dict) -> str:
    """A compact, model-friendly summary of the user's Application Portfolio."""
    app = _ensure_application_defaults(mem)
    ts = app.get("test_scores", {}) or {}
    profile = app.get("profile", {}) or {}
    essays = app.get("essays", {}) or {}
    ecs = app.get("ecs", {}) or {}
    awards = app.get("awards", {}) or {}
    prefs = app.get("preferences", {}) or {}
    wellness = app.get("wellness", {}) or {}

    lines = []
    lines.append("APPLICATION PORTFOLIO")
    lines.append(f"Profile: grade={_fmt(profile.get('grade'))}; country={_fmt(profile.get('country'))}")
    lines.append(
        "Tests: "
        f"GPA={_fmt(ts.get('gpa'))}; SAT={_fmt(ts.get('sat'))}; IELTS={_fmt(ts.get('ielts'))}; TOEFL={_fmt(ts.get('toefl'))}; ACT={_fmt(ts.get('act'))}"
    )
    lines.append(f"Intended major: {_fmt(prefs.get('intended_major'))}")
    tc = prefs.get("target_countries") or []
    lines.append("Target countries: " + (", ".join(tc) if tc else "—"))
    lines.append(f"Needs aid: {_fmt(prefs.get('needs_aid'))}; Budget: {_fmt(prefs.get('budget'))}")
    if (prefs.get("constraints") or "").strip():
        lines.append("Constraints: " + (prefs.get("constraints") or "").strip()[:220])
    if (ecs.get("summary") or "").strip():
        lines.append("EC summary: " + (ecs.get("summary") or "").strip()[:260])
    if ecs.get("spike"):
        lines.append("Spike: " + str(ecs.get("spike")))
    if awards.get("items"):
        lines.append("Awards: " + "; ".join([str(x) for x in awards.get("items", [])][:6]))
    if (essays.get("ps_topic") or "").strip():
        lines.append("PS topic: " + (essays.get("ps_topic") or "").strip()[:220])
    if (essays.get("supps_notes") or "").strip():
        lines.append("Supps notes: " + (essays.get("supps_notes") or "").strip()[:220])
    if (wellness.get("notes") or "").strip() or wellness.get("stress_level"):
        lines.append(
            "Wellness: "
            f"stress={_fmt(wellness.get('stress_level'))}; hours/wk={_fmt(wellness.get('hours_per_week'))}; "
            f"notes={(wellness.get('notes') or '').strip()[:160] or '—'}"
        )

    return "\n".join(lines)


def _fmt(v):
    return str(v).strip() if (v is not None and str(v).strip()) else "—"


def format_test_scores_block(app: dict) -> str:
    ts = (app or {}).get("test_scores", {}) or {}
    lines = [
        "🧾 TEST SCORES + GPA",
        f"• GPA: {_fmt(ts.get('gpa'))}",
        f"• SAT: {_fmt(ts.get('sat'))}",
        f"• SAT breakdown: {_fmt(ts.get('sat_breakdown'))}",
        f"• ACT: {_fmt(ts.get('act'))}",
        f"• IELTS: {_fmt(ts.get('ielts'))}",
        f"• TOEFL: {_fmt(ts.get('toefl'))}",
        f"• Duolingo: {_fmt(ts.get('duolingo'))}",
    ]
    notes = (ts.get("notes") or "").strip()
    if notes:
        lines.append(f"• Notes: {notes[:350]}")
    return "\n".join(lines)


def format_essays_block(app: dict) -> str:
    es = (app or {}).get("essays", {}) or {}
    lines = [
        "📝 ESSAYS STATUS",
        f"• Personal Statement: {_fmt(es.get('personal_statement'))}",
        f"• Supplementals: {_fmt(es.get('supplementals'))}",
        f"• Common App essay: {_fmt(es.get('common_app'))}",
    ]
    notes = (es.get("notes") or "").strip()
    if notes:
        lines.append(f"• Notes: {notes[:350]}")
    return "\n".join(lines)


def _split_bullets(text: str, limit: int = 10) -> list[str]:
    lines = []
    for raw in (text or "").splitlines():
        s = raw.strip()
        if not s:
            continue
        if s.startswith(("-", "•", "*")):
            s = s.lstrip("-*• ").strip()
        if s and s not in lines:
            lines.append(s)
        if len(lines) >= limit:
            break
    return lines


def format_ecs_block(app: dict) -> str:
    ecs = (app or {}).get("ecs", {}) or {}
    updated_at = ecs.get("updated_at") or "(never)"
    summary = ecs.get("summary") or "(not set)"
    highlights = ecs.get("highlights") or []

    lines = [
        "🎯 EC SUMMARY (My Application Portfolio)",
        f"Last updated: {updated_at}",
        "",
        "Summary:",
        str(summary),
        "",
        "Highlights:",
    ]
    if highlights:
        for h in highlights[:10]:
            lines.append(f"- {h}")
    else:
        lines.append("- (none yet)")
    lines.append(
        "\nSend your ECs as 3-8 short lines (role + what you did + impact). "
        "Example:\n- Robotics captain: led 12 students; built line-following bot; won regional\n- Tutoring: 2 hrs/week; helped 6 students raise grades"
    )
    return "\n".join(lines)


def format_awards_block(app: dict) -> str:
    awards = (app or {}).get("awards", {}) or {}
    updated_at = awards.get("updated_at") or "(never)"
    items = awards.get("items") or []
    lines = [
        "🏆 AWARDS & HONORS (My Application Portfolio)",
        f"Last updated: {updated_at}",
        "",
    ]
    if items:
        for it in items[:15]:
            lines.append(f"- {it}")
    else:
        lines.append("(none yet)")
    lines.append(
        "\nSend awards as a list (one per line). Example:\n- National Olympiad Silver (Math)\n- School 'Student of the Year'" 
    )
    return "\n".join(lines)


def format_prefs_block(app: dict, mem: dict) -> str:
    prefs = (app or {}).get("preferences", {}) or {}
    updated_at = prefs.get("updated_at") or "(never)"

    # Prefer explicit prefs but fall back to profile
    profile = (mem or {}).get("profile", {}) or {}
    major = prefs.get("major") or profile.get("major") or "(unknown)"
    countries = prefs.get("target_countries") or profile.get("target_countries") or []
    budget = prefs.get("budget") or profile.get("budget") or "(unknown)"
    needs_aid = prefs.get("needs_aid")
    if needs_aid is None:
        needs_aid = profile.get("needs_aid")
    deadlines = prefs.get("deadlines") or "(unknown)"
    notes = prefs.get("notes") or ""

    lines = [
        "🎯 PREFERENCES & CONSTRAINTS (My Application Portfolio)",
        f"Last updated: {updated_at}",
        "",
        f"Major/field: {major}",
        f"Target countries: {', '.join(countries) if countries else '(unknown)'}",
        f"Budget: {budget}",
        f"Needs financial aid: {needs_aid if needs_aid is not None else '(unknown)'}",
        f"Deadlines: {deadlines}",
    ]
    if notes:
        lines += ["", "Notes:", str(notes).strip()[:600]]

    lines.append(
        "\nSend a short message like:\n"
        "- Major: Computer Science\n"
        "- Countries: US, Canada, UK\n"
        "- Budget: need full scholarship\n"
        "- Deadlines: Nov 2026 EA/ED + Jan 2027 RD"
    )
    return "\n".join(lines)


def format_wellness_block(app: dict) -> str:
    w = (app or {}).get("wellness", {}) or {}
    updated_at = w.get("updated_at") or "(never)"
    lines = [
        "🧠 WELLNESS CHECK (My Application Portfolio)",
        f"Last updated: {updated_at}",
        "",
        f"Stress (1-10): {w.get('stress_level') or '(unknown)'}",
        f"Hours/week (school+apps): {w.get('hours_per_week') or '(unknown)'}",
        f"Sleep (avg hours): {w.get('sleep_hours') or '(unknown)'}",
    ]
    if w.get("notes"):
        lines += ["", "Notes:", str(w.get("notes")).strip()[:600]]

    lines.append(
        "\nSend a short update like: \"stress 7, 18 hrs/week, sleep 6.5; exams in March\". "
        "This helps me recommend a realistic school list + timeline." 
    )
    return "\n".join(lines)


def parse_and_update_test_scores(text: str, mem: dict) -> bool:
    """Parse user's message to update portfolio test scores. Returns True if anything updated."""
    if not text or not isinstance(mem, dict):
        return False
    app = _ensure_application_defaults(mem)
    ts = app["test_scores"]
    updated = False
    t = text

    # GPA
    m = re.search(r"\bGPA\s*[:=]?\s*(\d\.\d{1,2})\b", t, re.I)
    if m:
        ts["gpa"] = m.group(1); updated = True

    # SAT total
    m = re.search(r"\bSAT\s*[:=]?\s*(1[0-6]\d{2})\b", t, re.I)
    if m:
        ts["sat"] = m.group(1); updated = True

    # SAT breakdown
    m = re.search(r"\bMath\s*(\d{3})\b[^\n]{0,40}\b(EBRW|R&W|Reading|Verbal)\s*(\d{3})\b", t, re.I)
    if m:
        ts["sat_breakdown"] = f"Math {m.group(1)} / EBRW {m.group(3)}"; updated = True
    else:
        m2 = re.search(r"\b(EBRW|R&W|Reading|Verbal)\s*(\d{3})\b[^\n]{0,40}\bMath\s*(\d{3})\b", t, re.I)
        if m2:
            ts["sat_breakdown"] = f"Math {m2.group(3)} / EBRW {m2.group(2)}"; updated = True

    # ACT
    m = re.search(r"\bACT\s*[:=]?\s*(\d{2})\b", t, re.I)
    if m:
        ts["act"] = m.group(1); updated = True

    # IELTS
    m = re.search(r"\bIELTS\s*[:=]?\s*(\d\.\d|\d)\b", t, re.I)
    if m:
        ts["ielts"] = m.group(1); updated = True

    # TOEFL
    m = re.search(r"\bTOEFL\s*[:=]?\s*(\d{2,3})\b", t, re.I)
    if m:
        ts["toefl"] = m.group(1); updated = True

    # Duolingo / DET
    m = re.search(r"\b(Duolingo|DET)\s*[:=]?\s*(\d{2,3})\b", t, re.I)
    if m:
        ts["duolingo"] = m.group(2); updated = True

    # If they include any extra notes text, keep it (but do not overwrite if empty)
    if re.search(r"\bnotes?\b", t, re.I) and len(t) >= 20:
        # crude: take everything after "note:" or "notes:"
        m = re.search(r"\bnotes?\s*[:=]\s*(.+)$", t, re.I)
        if m:
            ts["notes"] = (m.group(1) or "").strip()[:600]
            updated = True

    # Mirror key fields into profile for prompts
    prof = mem.get("profile", {}) or {}
    if ts.get("gpa"): prof["gpa"] = ts["gpa"]
    if ts.get("sat"): prof["sat"] = ts["sat"]
    if ts.get("sat_breakdown"): prof["sat_breakdown"] = ts["sat_breakdown"]
    if ts.get("act"): prof["act"] = ts["act"]
    if ts.get("ielts"): prof["ielts"] = ts["ielts"]
    if ts.get("toefl"): prof["toefl"] = ts["toefl"]
    if ts.get("duolingo"): prof["duolingo"] = ts["duolingo"]
    mem["profile"] = prof

    return updated


def parse_and_update_essays(text: str, mem: dict) -> bool:
    """Parse user's message to update essay status. Returns True if anything updated."""
    if not text or not isinstance(mem, dict):
        return False
    app = _ensure_application_defaults(mem)
    es = app["essays"]
    t = text.strip()
    updated = False

    # status keywords
    statuses = ["not started", "outline", "draft", "revised", "final", "done"]
    def pick_status(s: str):
        low = s.lower()
        for st in statuses:
            if st in low:
                return st
        return None

    # PS
    if re.search(r"\b(ps|personal statement)\b", t, re.I):
        st = pick_status(t)
        if st:
            es["personal_statement"] = st; updated = True

    # Common App
    if re.search(r"\b(common app|commonapp)\b", t, re.I):
        st = pick_status(t)
        if st:
            es["common_app"] = st; updated = True

    # Supplementals: allow "6/12" or "6 of 12"
    m = re.search(r"(\d{1,2})\s*/\s*(\d{1,2})", t)
    if m and re.search(r"\bsupp", t, re.I):
        es["supplementals"] = f"{m.group(1)}/{m.group(2)} drafted"
        updated = True
    elif re.search(r"\bsupp", t, re.I):
        st = pick_status(t)
        if st:
            es["supplementals"] = st; updated = True

    if re.search(r"\bnotes?\b", t, re.I) and len(t) >= 20:
        m = re.search(r"\bnotes?\s*[:=]\s*(.+)$", t, re.I)
        if m:
            es["notes"] = (m.group(1) or "").strip()[:600]
            updated = True

    return updated


def parse_and_update_ecs(text: str, mem: dict) -> bool:
    """Save a short EC summary + bullet highlights."""
    if not text or not isinstance(mem, dict):
        return False
    app = _ensure_application_defaults(mem)
    ecs = app["ecs"]
    t = (text or "").strip()
    if not t:
        return False

    updated = False
    # Store raw summary (cap for safety)
    ecs["summary"] = t[:1400]
    updated = True

    bullets = _split_bullets(t, limit=10)
    if bullets:
        ecs["highlights"] = bullets
    else:
        # If no bullets, try to split sentences into quick highlights
        parts = re.split(r"(?<=[.!?])\s+", t)
        ecs["highlights"] = [p.strip()[:180] for p in parts if p.strip()][:5]

    ecs["updated_at"] = _now_utc_date()
    return updated


def parse_and_update_awards(text: str, mem: dict) -> bool:
    """Save awards/honors as a list."""
    if not text or not isinstance(mem, dict):
        return False
    app = _ensure_application_defaults(mem)
    awards = app["awards"]
    t = (text or "").strip()
    if not t:
        return False

    items = _split_bullets(t, limit=15)
    if not items:
        # Fall back to comma-separated
        maybe = [x.strip() for x in re.split(r"[,;]", t) if x.strip()]
        items = maybe[:15]

    if not items:
        return False

    awards["items"] = items
    awards["updated_at"] = _now_utc_date()
    return True


def _extract_field_line(text: str, keys: list[str]) -> str | None:
    for k in keys:
        m = re.search(rf"\b{k}\b\s*[:=]\s*(.+)$", text, flags=re.I | re.M)
        if m:
            return (m.group(1) or "").strip()
    return None


def parse_and_update_preferences(text: str, mem: dict) -> bool:
    """Save target countries/major/budget/aid notes. Also updates mem['profile'] when possible."""
    if not text or not isinstance(mem, dict):
        return False
    app = _ensure_application_defaults(mem)
    prefs = app["preferences"]
    prof = mem.setdefault("profile", {})
    t = (text or "").strip()
    if not t:
        return False

    updated = False

    # Try to extract a major/field of interest
    maj = _extract_field_line(t, ["major", "intended major", "field", "interest"]) or None
    if maj:
        prefs["major"] = maj[:120]
        prefs["intended_major"] = prefs.get("intended_major") or maj[:120]
        prof["major"] = maj[:120]
        prof["intended_major"] = prof.get("intended_major") or maj[:120]
        updated = True

    # Target countries/regions
    countries = _extract_field_line(t, ["countries", "country", "target", "region", "destinations"]) or None
    if countries:
        lst = [x.strip() for x in re.split(r"[,;/]", countries) if x.strip()]
        prefs["target_countries"] = lst[:10]
        updated = True

    # Budget / cost constraints
    budget = _extract_field_line(t, ["budget", "cost", "tuition", "max"]) or None
    if budget:
        prefs["budget"] = budget[:120]
        updated = True

    # Aid needs
    aid_line = _extract_field_line(t, ["aid", "financial aid", "scholarship", "need-based"]) or None
    if aid_line:
        if re.search(r"\b(no|don't|do not)\b", aid_line, re.I):
            prefs["needs_aid"] = False
            prof["needs_aid"] = False
        else:
            prefs["needs_aid"] = True
            prof["needs_aid"] = True
        updated = True

    deadlines = _extract_field_line(t, ["deadlines", "deadline", "intake", "term"]) or None
    if deadlines:
        prefs["deadlines"] = deadlines[:160]
        updated = True

    # Always store a short note block
    prefs["notes"] = t[:800]
    updated = True

    prefs["updated_at"] = _now_utc_date()
    return updated



def parse_and_update_prefs(text: str, mem: dict) -> bool:
    """Backward-compatible alias for older code paths."""
    return parse_and_update_preferences(text, mem)

def parse_and_update_wellness(text: str, mem: dict) -> bool:
    """Save stress/time/sleep snapshot. This supports realistic planning advice."""
    if not text or not isinstance(mem, dict):
        return False
    app = _ensure_application_defaults(mem)
    w = app["wellness"]
    t = (text or "").strip()
    if not t:
        return False

    # Defaults
    updated = False

    # Stress 1-10
    m = re.search(r"stress\s*[:=]?\s*(\d{1,2})", t, re.I)
    if m:
        try:
            val = int(m.group(1))
            if 0 <= val <= 10:
                w["stress_level"] = val
                updated = True
        except Exception:
            pass

    # Hours per week
    m = re.search(r"hours\s*/\s*week\s*[:=]?\s*(\d{1,3})", t, re.I)
    if not m:
        m = re.search(r"(\d{1,3})\s*h\s*/\s*week", t, re.I)
    if m:
        try:
            w["hours_per_week"] = int(m.group(1))
            updated = True
        except Exception:
            pass

    # Sleep hours
    m = re.search(r"sleep\s*[:=]?\s*(\d{1,2}(?:\.\d)?)", t, re.I)
    if m:
        try:
            w["sleep_hours"] = float(m.group(1))
            updated = True
        except Exception:
            pass

    support = _extract_field_line(t, ["support", "help", "notes"]) or None
    if support:
        w["support_needs"] = support[:200]
        updated = True

    w["notes"] = t[:500]
    updated = True
    w["updated_at"] = _now_utc_date()
    return updated

def compute_portfolio_readiness(mem: dict) -> tuple[int, list]:
    """Simple readiness score from saved portfolio (academics, essays, ECs, and constraints)."""
    if not isinstance(mem, dict):
        return 0, []
    app = _ensure_application_defaults(mem)
    ts = app.get("test_scores", {}) or {}
    es = app.get("essays", {}) or {}

    score = 0
    reasons = []

    # Tests
    if ts.get("gpa"):
        score += 20; reasons.append("GPA saved")
    if ts.get("sat") or ts.get("act"):
        score += 20; reasons.append("SAT/ACT saved")
    if ts.get("ielts") or ts.get("toefl") or ts.get("duolingo"):
        score += 10; reasons.append("English test saved")

    # Essays
    ps = (es.get("personal_statement") or "").lower()
    supp = (es.get("supplementals") or "").lower()
    ca = (es.get("common_app") or "").lower()

    def status_points(s: str):
        if "final" in s or "done" in s:
            return 20
        if "revised" in s:
            return 15
        if "draft" in s or "outline" in s:
            return 10
        if "not started" in s:
            return 0
        return 5 if s.strip() else 0

    score += status_points(ps); 
    if ps.strip(): reasons.append("PS status saved")
    score += status_points(ca);
    if ca.strip(): reasons.append("Common App status saved")

    # Supplementals: if includes x/y, scale
    m = re.search(r"(\d{1,2})\s*/\s*(\d{1,2})", supp)
    if m:
        x, y = int(m.group(1)), max(1, int(m.group(2)))
        score += int(round(20 * min(1.0, x / y)))
        reasons.append("Supplementals progress saved")
    else:
        sp = status_points(supp)
        if sp:
            score += min(20, sp)
            reasons.append("Supplementals status saved")

    # Extra portfolio parts (helps reflect readiness more realistically)
    ecs = app.get("ecs", {}) or {}
    aw = app.get("awards", {}) or {}
    pr = app.get("preferences", {}) or {}
    wl = app.get("wellness", {}) or {}

    if (ecs.get("summary") or "").strip():
        score += 5
        reasons.append("EC summary saved")
    if ecs.get("highlights"):
        score += 3
        reasons.append("EC highlights saved")
    if aw.get("items"):
        score += 3
        reasons.append("Awards saved")
    if any(
        [
            (pr.get("major") or "").strip(),
            pr.get("target_countries") or [],
            pr.get("budget"),
            pr.get("needs_aid") is not None,
        ]
    ):
        score += 3
        reasons.append("Preferences/constraints saved")
    if any([wl.get("stress_level"), wl.get("sleep_hours"), wl.get("hours_per_week")]):
        score += 2
        reasons.append("Wellness check saved")

    score = max(0, min(100, int(score)))
    return score, reasons


async def app_portfolio_show_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Entry point for 📂 My Application Portfolio."""
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)
    persist_user_memory(update, context)

    out = (
        "📂 MY APPLICATION PORTFOLIO (snapshot)\n\n"
        + format_test_scores_block(app)
        + "\n\n"
        + format_essays_block(app)
        + "\n\n"
        + format_ecs_block(app)
        + "\n\n"
        + format_awards_block(app)
        + "\n\n"
        + format_prefs_block(app, mem)
        + "\n\n"
        + format_wellness_block(app)
        + "\n\n"
        "Pick what you want to update or analyze.\n"
        "Tip: School Finder can suggest schools based on this portfolio."
    )
    await update.message.reply_text(out, reply_markup=app_portfolio_keyboard())


async def run_app_tests(update: Update, context: ContextTypes.DEFAULT_TYPE, user_text: str | None = None):
    """Show + optionally update Test Scores + GPA."""
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)

    if user_text:
        changed = parse_and_update_test_scores(user_text, mem)
        if changed:
            persist_user_memory(update, context)

    await update.message.reply_text(
        format_test_scores_block(app)
        + "\n\nSend an update like: `GPA 3.9, SAT 1520 (Math 790, EBRW 730), IELTS 7.5`\n"
        "Or tap ↩️ Back to Portfolio.",
        reply_markup=app_portfolio_keyboard(),
    )
    set_pending_feature(context, "app_tests")


async def run_app_essays(update: Update, context: ContextTypes.DEFAULT_TYPE, user_text: str | None = None):
    """Show + optionally update Essay Status."""
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)

    if user_text:
        changed = parse_and_update_essays(user_text, mem)
        if changed:
            persist_user_memory(update, context)

    await update.message.reply_text(
        format_essays_block(app)
        + "\n\nSend an update like: `PS draft`, `Supps 6/12`, `Common App revised`\n"
        "Or tap ↩️ Back to Portfolio.",
        reply_markup=app_portfolio_keyboard(),
    )
    set_pending_feature(context, "app_essays")


async def run_app_ecs(update: Update, context: ContextTypes.DEFAULT_TYPE, user_text: str | None = None):
    """Show + optionally update EC Summary."""
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)

    if user_text:
        changed = parse_and_update_ecs(user_text, mem)
        if changed:
            persist_user_memory(update, context)

    await update.message.reply_text(
        format_ecs_block(app)
        + "\n\nSend your top ECs in bullets (role + impact + time). Example:\n"
        "- Robotics captain: led 12, won regional, mentored 30 juniors (6h/week)\n"
        "- Research intern: co-authored poster, built Python pipeline\n\n"
        "Or tap ↩️ Back to Portfolio.",
        reply_markup=app_portfolio_keyboard(),
    )
    set_pending_feature(context, "app_ecs")


async def run_app_awards(update: Update, context: ContextTypes.DEFAULT_TYPE, user_text: str | None = None):
    """Show + optionally update Awards & Honors."""
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)

    if user_text:
        changed = parse_and_update_awards(user_text, mem)
        if changed:
            persist_user_memory(update, context)

    await update.message.reply_text(
        format_awards_block(app)
        + "\n\nSend awards/honors in bullets (award + level + year). Example:\n"
        "- National Math Olympiad finalist (2025)\n"
        "- School top 1% GPA (2024)\n\n"
        "Or tap ↩️ Back to Portfolio.",
        reply_markup=app_portfolio_keyboard(),
    )
    set_pending_feature(context, "app_awards")


async def run_app_prefs(update: Update, context: ContextTypes.DEFAULT_TYPE, user_text: str | None = None):
    """Show + optionally update Preferences & Constraints."""
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)

    if user_text:
        changed = parse_and_update_prefs(user_text, mem)
        if changed:
            persist_user_memory(update, context)

    await update.message.reply_text(
        format_prefs_block(app, mem)
        + "\n\nSend your preferences in 3–6 short lines, e.g.:\n"
        "Major: CS / Data Science\n"
        "Countries: US, UK, Canada\n"
        "Budget: $10k/year\n"
        "Aid: yes\n"
        "Notes: want strong research + entrepreneurship\n\n"
        "Or tap ↩️ Back to Portfolio.",
        reply_markup=app_portfolio_keyboard(),
    )
    set_pending_feature(context, "app_prefs")


async def run_app_wellness(update: Update, context: ContextTypes.DEFAULT_TYPE, user_text: str | None = None):
    """Show + optionally update Wellness Check."""
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)

    if user_text:
        changed = parse_and_update_wellness(user_text, mem)
        if changed:
            persist_user_memory(update, context)

    await update.message.reply_text(
        format_wellness_block(app)
        + "\n\nOptional: share current load, e.g.:\n"
        "Stress: 7/10\n"
        "Sleep: 6.5h\n"
        "Workload: 25h/week\n"
        "Support needs: time management / burnout\n\n"
        "Or tap ↩️ Back to Portfolio.",
        reply_markup=app_portfolio_keyboard(),
    )
    set_pending_feature(context, "app_wellness")


async def run_advisor_mode(update: Update, context: ContextTypes.DEFAULT_TYPE, user_text: str):
    """Advisor Mode: College Fit Analysis from saved portfolio + user notes."""
    await show_typing(update, context)

    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    app = _ensure_application_defaults(mem)
    prof = mem.get("profile", {}) or {}

    # Use portfolio + optional user's extra notes
    sys = (
        "You are UniVenture Advisor - a sharp college-fit analyst.\n"
        "Your job: give a realistic, helpful College Fit Analysis using the student's stats and constraints.\n\n"
        "Output format (exact headings):\n"
        "1) 🎯 Snapshot (2 lines)\n"
        "2) ✅ Strengths (3 bullets)\n"
        "3) ⚠️ Risks / gaps (3 bullets)\n"
        "4) 🏫 Fit strategy (Reach/Match/Safety approach + what to prioritize)\n"
        "5) 📌 Next 7 days (3 concrete actions)\n\n"
        "Rules: no fake acceptance %; be specific; consider financial aid if mentioned.\n"
    )

    # Writing insights are stored at top-level mem['writing'] (counts + tags), but keep a fallback.
    writing = (mem.get("writing") or app.get("writing") or {})
    strengths_dict = writing.get("strengths") or {}
    issues_dict = writing.get("recurring_issues") or {}
    voice_tags = writing.get("voice_tags") or []
    themes = writing.get("themes") or []
    last_feedback = (writing.get("last_feedback_summary") or "").strip()

    wr_strengths = [k for k, _ in sorted(strengths_dict.items(), key=lambda kv: kv[1], reverse=True)[:5]]
    wr_issues = [k for k, _ in sorted(issues_dict.items(), key=lambda kv: kv[1], reverse=True)[:5]]

    portfolio_block = (
        "STUDENT PORTFOLIO:\n"
        + format_test_scores_block(app)
        + "\n\n"
        + format_essays_block(app)
        + "\n\n"
        + format_ecs_block(app)
        + "\n\n"
        + format_awards_block(app)
        + "\n\n"
        + format_prefs_block(app, mem)
        + "\n\n"
        + format_wellness_block(app)
        + "\n\n"
        + f"Needs aid: {bool(prof.get('needs_aid'))}\n"
        + "\nWRITING SIGNALS (from past evals):\n"
        + (f"- Strengths: {', '.join(wr_strengths)}\n" if wr_strengths else "- Strengths: (none saved yet)\n")
        + (f"- Recurring issues: {', '.join(wr_issues)}\n" if wr_issues else "- Recurring issues: (none saved yet)\n")
        + (f"- Voice tags: {', '.join(voice_tags[:6])}\n" if voice_tags else "")
        + (f"- Themes: {', '.join(themes[:6])}\n" if themes else "")
        + (f"- Last feedback summary: {last_feedback[:220]}\n" if last_feedback else "")
    )

    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": portfolio_block},
        {"role": "user", "content": user_text or "Use the saved portfolio only and give my fit analysis."},
    ]

    out = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.45, max_tokens=550)
    await send_long(update, out)
    await update.message.reply_text("Want to update scores/essays or run Advisor Mode again?", reply_markup=app_portfolio_keyboard())


async def run_brainstorm(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    """Brainstorm helper usable from both /brainstorm and button-based flows.

    IMPORTANT: This function must be self-contained (no free vars like q/context_block).
    """
    await show_typing(update, context)

    topic = get_current_topic(context)
    user = update.effective_user
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="brainstorm")

    # Treat the provided description as the user's "question" for prompt + profile extraction.
    q_local = (description or "").strip()

    # Retrieve a small amount of relevant guidance context (non-fatal if empty/unavailable)
    context_block = ""
    try:
        col = get_collection(chat_id, topic)
        try:
            res = col.query(query_texts=[q_local], where={"type": "qa"}, n_results=4)
            docs = res.get("documents", [[]])[0]
        except Exception:
            # fallback without filter
            res = col.query(query_texts=[q_local], n_results=4)
            docs = res.get("documents", [[]])[0]
        context_block = "\n\n---\n\n".join(docs or [])
    except Exception as e:
        logging.error(f"Error querying collection for brainstorm: {e}")

    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)

    # Update profile signals from the brainstorm description (cheap heuristic)
    try:
        mem["profile"] = extract_profile_signals(q_local, mem.get("profile", {}) or {})
        persist_user_memory(update, context)
    except Exception as e:
        logging.exception("Brainstorm memory update failed (non-fatal): %s", e)

    decision_notes = coach_decision_notes(topic, q_local, mem)
    system_prompt = coach_qa_system_prompt(topic, mem, decision_notes)

    messages = [{"role": "system", "content": system_prompt}]
    if context_block:
        messages.append({
            "role": "system",
            "content": f"Program-specific notes and examples (may be empty):\n{context_block}",
        })
    messages.append({"role": "user", "content": "Here is the student's situation:\n\n" + q_local})

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.5)
    await send_long(update, a)


async def brainstorm_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "brainstorm")
        await update.message.reply_text(
            "🧠 Brainstorm mode ON.\n\nBriefly tell me about yourself, your target major, and what you want to write about."
        )
        return
    await run_brainstorm(update, context, parts[1].strip())

async def run_rewrite(update: Update, context: ContextTypes.DEFAULT_TYPE, text_to_fix: str):
    await show_typing(update, context)
    topic = get_current_topic(context)
    user = update.effective_user
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="rewrite")

    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)

    try:
        col = get_collection(chat_id, topic)
        res = col.query(query_texts=[text_to_fix], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying for rewrite: {e}")
        docs = []

    context_block = "\n\n---\n\n".join(docs) if docs else ""

    if topic == "essays_personal":
        sys = (
            "ROLE:\n"
            "You are a senior admissions essay strategist rewriting a PERSONAL STATEMENT for top universities.\n"
            "You elevate identity, emotional depth, inner conflict, and transformation — not just grammar.\n\n"
            "GOALS (Personal Statement):\n"
            "1) Strengthen the narrative arc: scene → tension → insight → changed self.\n"
            "2) Add psychological depth: uncertainty, stakes, vulnerability, what it cost.\n"
            "3) Increase differentiation: avoid common 'hard work' or 'love of subject' clichés.\n"
            "4) Preserve the student's authentic voice; do not make it poetic or fake.\n\n"
            "RULES:\n"
            "- Do NOT add achievements the student didn't mention.\n"
            "- If a paragraph is generic, compress or replace with a concrete moment.\n"
            "- Make the ending feel inevitable and forward-looking (who they are becoming).\n\n"
            "OUTPUT FORMAT (STRICT):\n"
            "📝 ELEVATED VERSION:\n"
            "[Full improved version]\n\n"
            "🧠 WHAT CHANGED STRATEGICALLY:\n"
            "- [Biggest arc/structure improvement]\n"
            "- [Biggest depth/vulnerability improvement]\n"
            "- [Biggest differentiation improvement]\n"
        )

    elif topic == "essays_supplemental":
        sys = (
            "ROLE:\n"
            "You are a senior admissions strategist rewriting a SUPPLEMENTAL ESSAY for top universities.\n"
            "You elevate specificity, fit, intellectual direction, and contribution — not generic polish.\n\n"
            "GOALS (Supplemental):\n"
            "1) Strengthen fit: connect the student's goals to specific programs/values (without inventing facts).\n"
            "2) Increase specificity: replace generic phrases ('great community') with concrete details the student provided.\n"
            "3) Make the 'Why you' + 'Why us' logic tight: YOU → SCHOOL → YOU.\n"
            "4) Keep it concise and high-signal; cut filler.\n\n"
            "RULES:\n"
            "- Do NOT invent professors, labs, clubs, courses, or statistics.\n"
            "- If the prompt is unclear, write a version that can be easily customized with placeholders like [LAB/PROF/PROGRAM].\n"
            "- End with contribution + forward motion (how they will engage and add value).\n\n"
            "OUTPUT FORMAT (STRICT):\n"
            "📝 ELEVATED VERSION:\n"
            "[Full improved version]\n\n"
            "🧠 WHAT CHANGED STRATEGICALLY:\n"
            "- [Biggest fit/specificity improvement]\n"
            "- [Biggest structure improvement]\n"
            "- [Biggest differentiation improvement]\n"
        )

    else:
        sys = (
            f"ROLE:\n"
            f"You are a senior admissions writing strategist improving a student's {nice_topic} text.\n"
            "You elevate clarity, structure, credibility, and impact — not just grammar.\n\n"
            "GOALS:\n"
            "1) Increase differentiation and specificity.\n"
            "2) Remove cliché language and filler.\n"
            "3) Strengthen flow and logical structure.\n"
            "4) Preserve the student's authentic voice.\n\n"
            "RULES:\n"
            "- Do NOT add fake achievements.\n"
            "- Compress weak sections.\n"
            "- Keep it natural and human.\n\n"
            "OUTPUT FORMAT (STRICT):\n"
            "📝 ELEVATED VERSION:\n"
            "[Full improved version]\n\n"
            "🧠 WHAT CHANGED STRATEGICALLY:\n"
            "- [Most important improvement]\n"
            "- [Second important improvement]\n"
            "- [Third important improvement]\n"
        )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {"role": "system", "content": f"Program-specific notes and examples (may be empty):\n{context_block}"}
        )
    messages.append({"role": "user", "content": f"Here is the text to improve:\n\n{text_to_fix}"})

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.4)
    await send_long(update, a)

async def rewrite_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "rewrite")
        await update.message.reply_text("✍️ Rewrite mode ON.\n\nSend me the paragraph or essay you want me to improve.")
        return
    await run_rewrite(update, context, parts[1].strip())

async def run_plan(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="plan")

    try:
        col = get_collection(chat_id, topic)
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying for plan: {e}")
        docs = []
        
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        "You are an admissions strategy mentor.\n"
        "- Based on the student's situation, create a concise application plan.\n"
        "- Organize it into short bullet points under 3 headings: Academics & Testing, Essays & Recs, Activities & Extras.\n"
        "- Keep total response around 120-200 words.\n"
        "- Focus on practical next steps, not theory."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append({"role": "system", "content": f"Program-specific planning notes (may be empty):\n{context_block}"})
    messages.append({"role": "user", "content": "Here is my situation:\n\n" + description})

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.5)
    await send_long(update, a)


async def run_plan_from_portfolio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Create an application plan using the student's saved My Application Portfolio + coach memory."""
    await show_typing(update, context)
    user = update.effective_user
    chat_id = update.effective_chat.id

    # Ensure we are in the right topic
    context.user_data["topic"] = "application_plan"
    track_topic(context, "application_plan")

    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)

    # Extract saved strengths/weaknesses signals from prior evaluations (if any)
    writing = (mem or {}).get("writing") or {}
    strengths_dict = writing.get("strengths") or {}
    issues_dict = writing.get("issues") or {}
    wr_strengths = [
        k for k, v in strengths_dict.items() if isinstance(v, (int, float)) and v >= 1
    ]
    wr_issues = [k for k, v in issues_dict.items() if isinstance(v, (int, float)) and v >= 1]

    if not has_application_portfolio_data(mem):
        set_pending_feature(context, "plan")
        await update.message.reply_text(
            "📅 Application Plan\n\n"
            "I don't have enough info saved in your 📁 My Application Portfolio yet.\n\n"
            "Tell me your grade, target countries, intended major, test scores (if any), and your rough deadlines.",
            reply_markup=plan_keyboard(),
        )
        return

    record_event(user.id, "application_plan", kind="plan_from_portfolio")

    portfolio_block = _app_portfolio_compact_summary(mem)
    coach_mem = memory_summary_for_prompt(mem)
    decision_notes = coach_decision_notes("application_plan", portfolio_block, mem)

    # Optional planning context from stored sources (non-fatal if empty)
    context_block = ""
    try:
        col = get_collection(chat_id, "application_plan")
        res = col.query(query_texts=[portfolio_block], n_results=6)
        docs = res.get("documents", [[]])[0]
        context_block = "\n\n---\n\n".join(docs or [])
    except Exception as e:
        logging.error(f"Error querying for plan_from_portfolio: {e}")

    sys = (
        "You are an admissions strategy mentor.\n"
        "Create a concise application plan using the student's saved application portfolio and coach memory.\n"
        "Rules:\n"
        "- Organize it into short bullet points under 3 headings: Academics & Testing, Essays & Recs, Activities & Extras.\n"
        "- Keep it actionable and personalized (no generic advice).\n"
        "- Keep total response around 130-220 words.\n"
        "- If info is missing, add 2-4 short clarification questions at the end (max 4).\n"
    )

    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": f"Student memory (may be empty):\n{coach_mem}"},
        {"role": "system", "content": f"Decision notes:\n{decision_notes}"},
    ]
    if wr_strengths or wr_issues:
        s = ", ".join(wr_strengths[:6]) if wr_strengths else "(none saved)"
        w = ", ".join(wr_issues[:6]) if wr_issues else "(none saved)"
        messages.append({
            "role": "system",
            "content": f"Coach analysis from prior evals (may be empty):\nStrengths: {s}\nWeaknesses/issues: {w}",
        })
    if context_block:
        messages.append({"role": "system", "content": f"Planning notes (may be empty):\n{context_block}"})
    messages.append({"role": "user", "content": f"Use this application portfolio:\n\n{portfolio_block}"})

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.5)
    await send_long(update, a)

async def plan_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["topic"] = "application_plan"
    track_topic(context, "application_plan")
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "plan")
        await update.message.reply_text(
            "📅 Application Plan mode ON.\n\nTell me your grade, target countries, intended major, test scores (if any), and your rough deadlines.",
            reply_markup=plan_keyboard(),
        )
        return
    await run_plan(update, context, parts[1].strip())

async def run_recpacket(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="recpacket")

    sys = (
        "You are creating a recommendation letter 'brag sheet' for a teacher.\n"
        "- Output in 3 short sections:\n"
        "  1) 3-5 sentence summary the student can give the teacher.\n"
        "  2) Bullet list of key achievements/impacts.\n"
        "  3) Bullet list of personal qualities and 2-3 specific story ideas.\n"
        "- Keep it concise and realistic for competitive admissions."
    )

    messages = [{"role": "system", "content": sys}, {"role": "user", "content": description}]
    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.5)
    await send_long(update, a)

async def recpacket_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "recpacket")
        await update.message.reply_text(
            "✉️ Rec Letter Packet mode ON.\n\nTell me which teacher will write your rec, what classes you took with them, your achievements, and what you want them to highlight."
        )
        return
    await run_recpacket(update, context, parts[1].strip())

async def run_schoolfinder(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="schoolfinder")

    try:
        col = get_collection(chat_id, topic)
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying for schoolfinder: {e}")
        docs = []
        
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        "You are a university match advisor.\n"
        "- Based on the student's stats and preferences, suggest Reach, Match, and Safety school types and a few example universities.\n"
        "- For each category, give 2-4 example schools and 1-2 bullets about why they fit.\n"
        "- Keep total response concise (around 150-220 words).\n"
        "- Make it clear this is an approximate starting point and they must research details themselves."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {"role": "system", "content": f"Program-specific school lists/notes (may be empty):\n{context_block}"}
        )
    messages.append({"role": "user", "content": description})

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.6)
    await send_long(update, a)


async def run_schoolfinder_from_portfolio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Suggest schools using the user's saved My Application Portfolio + analysis."""
    await show_typing(update, context)
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)

    if not has_application_portfolio_data(mem):
        set_pending_feature(context, "schoolfinder")
        await update.message.reply_text(
            "🏫 School Finder\n\n"
            "I don't have enough info saved in your 📁 My Application Portfolio yet.\n\n"
            "Send me your GPA (or approximate), test scores (if any), budget, target countries, intended major, and constraints (e.g. need scholarship).",
            reply_markup=schoolfinder_keyboard(),
        )
        return
    app = _ensure_application_defaults(mem)
    prof = mem.get("profile", {}) or {}

    # Compact snapshot for the model
    # Writing insights are stored at top-level mem['writing'] as counts + tags.
    writing = (mem.get("writing") or app.get("writing") or {})
    strengths_dict = writing.get("strengths") or {}
    issues_dict = writing.get("recurring_issues") or {}
    wr_strengths = [k for k, _ in sorted(strengths_dict.items(), key=lambda kv: kv[1], reverse=True)[:5]]
    wr_issues = [k for k, _ in sorted(issues_dict.items(), key=lambda kv: kv[1], reverse=True)[:5]]

    portfolio_block = (
        "SAVED PORTFOLIO SNAPSHOT\n"
        + format_test_scores_block(app)
        + "\n\n"
        + format_essays_block(app)
        + "\n\n"
        + format_ecs_block(app)
        + "\n\n"
        + format_awards_block(app)
        + "\n\n"
        + format_prefs_block(app, mem)
        + "\n\n"
        + format_wellness_block(app)
        + "\n\nWRITING ANALYSIS (from past evals)\n"
        + f"Strengths: {', '.join(wr_strengths) if wr_strengths else 'unknown'}\n"
        + f"Risks / gaps: {', '.join(wr_issues) if wr_issues else 'none logged'}\n"
        + f"Needs aid: {prof.get('needs_aid', 'unknown')}\n"
        + f"Intended major: {prof.get('major') or prof.get('intended_major') or 'unknown'}\n"
    )

    # Optional RAG (if you teach any school-finder rubric/notes)
    context_block = ""
    chat_id = update.effective_chat.id
    try:
        col = get_collection(chat_id, "school_finder")
        try:
            r = col.query(query_texts=[portfolio_block], where={"type": "qa"}, n_results=4)
            docs = (r.get("documents") or [[]])[0]
        except Exception:
            r = col.query(query_texts=[portfolio_block], n_results=4)
            docs = (r.get("documents") or [[]])[0]
    except Exception as e:
        logging.error(f"Error querying school_finder sources: {e}")
        docs = []


    sys_prompt = (
        "You are a university match advisor. Use the student's saved portfolio and writing analysis to suggest schools. "
        "Be realistic for a Central Asia international applicant. If constraints are missing, assume they want: strong academics, "
        "good outcomes, and some financial-aid possibilities.\n\n"
        "OUTPUT FORMAT:\n"
        "1) 2 bullets: Strengths (based on portfolio + writing analysis)\n"
        "2) 2 bullets: Risks / Gaps\n"
        "3) School list in 3 buckets: REACH / MATCH / SAFETY (3–6 schools each). For each school: 1 short reason.\n"
        "4) Next step: 1 sentence telling what info to add to improve the list.\n\n"
        "Keep it under ~250-320 words. Do NOT invent exact acceptance rates." + context_block
    )

    messages = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": portfolio_block},
    ]
    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.5)
    await send_long(update, a)

async def schoolfinder_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["topic"] = "school_finder"
    track_topic(context, "school_finder")
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "schoolfinder")
        await update.message.reply_text(
            "🏫 School Finder mode ON.\n\n"
            f"Option A: tap '{BTN_SF_MANUAL}' and send me your GPA, tests, budget, target countries, intended major, and constraints (e.g. need scholarship).\n\n"
            f"Option B: tap '{BTN_SF_FROM_PORT}' to use your saved My Application Portfolio (and the bot's strengths/weaknesses analysis).",
            reply_markup=schoolfinder_keyboard(),
        )
        return
    await run_schoolfinder(update, context, parts[1].strip())

async def run_portfolioideas(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="portfolioideas")

    sys = (
        "You are a portfolio mentor for university applications.\n"
        "- Based on the student's field (e.g. CS, design, art, film, business) and interests, suggest 3-6 concrete project ideas.\n"
        "- Each idea should be 1-2 sentences, focused on impact and what it shows about the student.\n"
        "- Make ideas realistic for a high school student, but impressive."
    )

    messages = [{"role": "system", "content": sys}, {"role": "user", "content": description}]
    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.6)
    await send_long(update, a)

async def portfolioideas_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "portfolioideas")
        await update.message.reply_text(
            "🖼️ Portfolio Ideas mode ON.\n\nTell me your field (CS, design, art, film, etc.), your skills, and the kind of programs you are targeting."
        )
        return
    await run_portfolioideas(update, context, parts[1].strip())

# ---------- BOOST TOOLS (CROSS-TOPIC) ----------
def _progress_bar(pct: int) -> str:
    pct = max(0, min(100, int(pct)))
    filled = int(round(pct / 10))
    return "█" * filled + "░" * (10 - filled) + f" {pct}%"

async def tool_my_progress(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Lightweight progress dashboard based on on-bot usage."""
    user = update.effective_user
    stats = load_stats()
    uid = str(user.id)

    msg_count = int(stats.get("messages_per_user", {}).get(uid, 0) or 0)
    evals_done = int(context.user_data.get("my_eval_count", 0) or 0)
    topics_seen = context.user_data.get("topics_seen", []) or []
    tools_used = context.user_data.get("tools_used", []) or []

    engagement = min(100, msg_count * 4)              # 25 messages -> 100
    drafting = min(100, evals_done * 25)              # 4 evals -> 100
    exploration = min(100, len(topics_seen) * 12)     # ~8 topics -> 96
    toolkit = min(100, len(tools_used) * 25)          # 4 tools -> 100

    current = FRIENDLY_TOPIC_NAMES.get(get_current_topic(context), "General")

    out = (
        "📊 MY PROGRESS\n\n"
        f"Engagement:  {_progress_bar(engagement)}\n"
        f"Drafting:    {_progress_bar(drafting)}\n"
        f"Exploration: {_progress_bar(exploration)}\n"
        f"Boost Tools: {_progress_bar(toolkit)}\n\n"
        f"Stats:\n"
        f"• Messages with me: {msg_count}\n"
        f"• Evaluations done: {evals_done}\n"
        f"• Topics explored: {len(topics_seen)}\n"
        f"• Boost tools used: {len(tools_used)}\n"
        f"• Current focus: {current}\n\n"
        "Next step idea: Run 🔍 Find Wow Factor on your latest draft, then apply the suggestions in a rewrite."
    )

    track_tool_use(context, "progress")
    persist_user_memory(update, context)
    await update.message.reply_text(out, reply_markup=tools_menu_keyboard())

async def tool_insider_tips(update: Update, context: ContextTypes.DEFAULT_TYPE):
    topic = get_current_topic(context)

    tips_by_topic = {
        "essays_personal": (
            "🤫 INSIDER TIPS — PERSONAL STATEMENT\n\n"
            "• Your first 2 lines are everything: start with a moment, not an intro.\n"
            "• One specific scene > 10 generic achievements.\n"
            "• Show reflection: what changed in you, not just what happened.\n"
            "• AOs trust proof: add tiny details (time, place, sensory).\n"
            "• End with forward motion: what you’ll do next in college."
        ),
        "essays_supplemental": (
            "🤫 INSIDER TIPS — SUPPLEMENTALS\n\n"
            "• 'Why us' works best as: YOU → THEIR SPECIFICS → YOU AGAIN.\n"
            "• Mention 2 ultra-specific fit points (lab, prof, program, initiative).\n"
            "• Avoid resume repetition—add new angles and values.\n"
            "• Short prompts: one strong claim + one mini-story + one insight.\n"
            "• Make it sound like a real student, not marketing copy."
        ),
        "extracurriculars": (
            "🤫 INSIDER TIPS — EXTRACURRICULARS\n\n"
            "• Impact beats title. Numbers help (people reached, hours, funds).\n"
            "• Show progression: member → builder → leader/mentor.\n"
            "• Use action verbs + outcomes (what changed because of you).\n"
            "• One 'spike' (a theme) is stronger than 15 random clubs.\n"
            "• Always answer: Why you? Why it matters?"
        ),
        "recommendations": (
            "🤫 INSIDER TIPS — RECOMMENDATIONS\n\n"
            "• Best letters include 2–3 stories that only that teacher could tell.\n"
            "• Ask teachers who saw you struggle AND grow (not just easy A's).\n"
            "• Give them a brag sheet with facts, projects, and specific moments.\n"
            "• Strong recs compare you to peers (“top 5% I’ve taught”).\n"
            "• Remind early: 3 gentle reminders > 1 last-minute panic."
        ),
        "portfolio": (
            "🤫 INSIDER TIPS — PORTFOLIO\n\n"
            "• Curate: 6 great pieces beats 20 average ones.\n"
            "• Add process: drafts, iterations, what you learned.\n"
            "• Label your role clearly (solo vs team, what you owned).\n"
            "• Make it scannable: titles, 1-line context, 1-line takeaway.\n"
            "• Tie to future: what you want to build next."
        ),
        "ielts_writing": (
            "🤫 INSIDER TIPS — IELTS WRITING\n\n"
            "• Task 2: clear position in the intro + topic sentences every paragraph.\n"
            "• Aim for: example → explanation → link back.\n"
            "• Don’t chase fancy words—accuracy > complexity.\n"
            "• Use cohesive devices naturally (however, therefore, moreover).\n"
            "• Save 3 minutes to check grammar + articles + verb tenses."
        ),
    }

    out = tips_by_topic.get(
        topic,
        "🤫 INSIDER TIPS\n\n• Pick a menu topic first for more tailored tips.\n• If you want the fastest improvement: run an evaluation, then ask follow-up questions on the feedback."
    )

    track_tool_use(context, "insider")
    persist_user_memory(update, context)
    await update.message.reply_text(out, reply_markup=tools_menu_keyboard())

async def tool_power_words(update: Update, context: ContextTypes.DEFAULT_TYPE):
    topic = get_current_topic(context)

    packs = {
        "essays_personal": {
            "Action": ["spearheaded", "orchestrated", "built", "revived", "initiated"],
            "Reflection": ["realized", "reframed", "questioned", "unlearned", "grew into"],
            "Impact": ["shifted", "amplified", "shaped", "opened", "strengthened"],
        },
        "essays_supplemental": {
            "Fit": ["aligned", "intersected", "clicked", "connected", "matched"],
            "Academics": ["inquiry", "research", "seminar", "lab", "capstone"],
            "Community": ["collaborate", "mentor", "contribute", "co-create", "engage"],
        },
        "extracurriculars": {
            "Leadership": ["led", "mobilized", "trained", "scaled", "launched"],
            "Results": ["increased", "reduced", "raised", "delivered", "reached"],
            "Innovation": ["engineered", "prototyped", "automated", "designed", "iterated"],
        },
        "ielts_writing": {
            "Argument": ["therefore", "however", "moreover", "consequently", "nevertheless"],
            "Precision": ["notably", "primarily", "increasingly", "specifically", "ultimately"],
            "Neutral tone": ["suggests", "indicates", "tends to", "is likely to", "can be"],
        },
    }

    pack = packs.get(topic, packs["essays_personal"])
    label = FRIENDLY_TOPIC_NAMES.get(topic, "Your writing")

    lines = [f"⚡ POWER WORDS — {label.upper()}\n"]
    for k, words in pack.items():
        lines.append(f"{k}: " + ", ".join(words))
    lines.append("\nTip: Replace weak verbs (did/helped) with one stronger verb + a result (what changed?).")

    track_tool_use(context, "powerwords")
    persist_user_memory(update, context)
    await update.message.reply_text("\n".join(lines), reply_markup=tools_menu_keyboard())

async def tool_predict_chances(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """A 'readiness indicator' (NOT a real acceptance probability)."""
    user = update.effective_user
    stats = load_stats()
    uid = str(user.id)

    msg_count = int(stats.get("messages_per_user", {}).get(uid, 0) or 0)
    evals_done = int(context.user_data.get("my_eval_count", 0) or 0)
    topics_seen = context.user_data.get("topics_seen", []) or []
    tools_used = context.user_data.get("tools_used", []) or []

    score = 0
    score += min(35, msg_count)                # up to 35
    score += min(25, evals_done * 8)           # up to 25
    score += min(20, len(topics_seen) * 3)     # up to 20
    score += min(20, len(tools_used) * 4)      # up to 20
    score = max(0, min(100, int(score)))

    if score < 35:
        level = "Early-stage"
        next_steps = "Run 1 evaluation (PS or Supplementals) and use ⚡ Power Words on the revision."
    elif score < 70:
        level = "Building"
        next_steps = "Do 2 evaluations + one rewrite pass focused on clarity + reflection."
    else:
        level = "Strong momentum"
        next_steps = "Polish: tighten openings, add specificity, and align your EC story to your major theme."

    out = (
        "🎯 PREDICT MY CHANCES (READINESS INDICATOR)\n\n"
        f"Readiness score: {_progress_bar(score)}\n"
        f"Status: {level}\n\n"
        "What this is: a rough indicator of how complete/strong your materials are based on your usage here.\n"
        "What this is NOT: a real admissions probability (schools use many factors + external context).\n\n"
        f"Recommended next step: {next_steps}"
    )

    if PAYWALL_ENABLED and (not is_pro_user(update)):
        out += _upgrade_pitch()

    track_tool_use(context, "predict")
    persist_user_memory(update, context)
    await update.message.reply_text(out, reply_markup=tools_menu_keyboard())

async def run_wowfactor(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str):
    """Analyze a text and highlight the most memorable 'wow factor'."""
    # Paywall (optional)
    if (not is_pro_user(update)) and ('wowfactor' in PRO_ONLY_TOOLS):
        await update.message.reply_text(
            "🔍 Find Wow Factor is a PRO tool.\n\n"
            "Free preview: I'll tell you the direction, but PRO unlocks the full breakdown + rewrite plan."
            + _upgrade_pitch(),
            reply_markup=tools_menu_keyboard(),
        )
        # Give a tiny preview based on first 400 chars
        preview = (text or '')[:400]
        if len(preview.strip()) >= 120:
            await show_typing(update, context)
            mini = await openai_chat(
                model=FAST_MODEL,
                messages=[
                    {'role': 'system', 'content': 'Give ONE sentence: what is the most unique hook in this text? Be specific.'},
                    {'role': 'user', 'content': preview},
                ],
                temperature=0.4,
                max_tokens=90,
            )
            await update.message.reply_text('Preview hook: ' + sanitize_output(mini))
        return
    await show_typing(update, context)

    sys = (
        "You are a top college admissions essay coach. "
        "Find the ONE most unique, compelling, memorable element in the student's text — the 'wow factor'.\n\n"
        "Output format (exact headings):\n"
        "🎯 WOW FACTOR: <2-6 word label>\n"
        "✨ Why it stands out: <1-2 sentences>\n"
        "💪 How to amplify it:\n"
        "- <actionable step 1>\n"
        "- <actionable step 2>\n"
        "- <actionable step 3>\n"
        "⚠️ Biggest risk to fix: <1 sentence>\n\n"
        "Be specific. Do not mention AI." 
    )

    content = (text or "").strip()
    content = content[:2500]

    a = await openai_chat(
        model=FAST_MODEL,
        messages=[{"role": "system", "content": sys}, {"role": "user", "content": content}],
        temperature=0.4,
    )

    track_tool_use(context, "wowfactor")
    persist_user_memory(update, context)
    await send_long(update, "🔍 FIND WOW FACTOR\n\n" + (a or ""))
    # Keep the tools keyboard visible for the next click.
    await update.message.reply_text("Pick another tool (or tap ⬅️ Back).", reply_markup=tools_menu_keyboard())

# ---------- MAIN ANSWER ----------
async def answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    if text.startswith("/"):
        # User sent a command while in a pending mode (rewrite/brainstorm/etc.)
        if context.user_data.get("pending_feature"):
            clear_pending_feature(context)
            await update.message.reply_text(
                "ℹ️ Switched out of the previous mode to process your command."
            )
        return

    logging.info(f"💬 TEXT RECEIVED: {text}")

    chat_id = update.effective_chat.id
    user = update.effective_user
    q = text
    

    # ✅ FIX: If user is in a pending mode and clicks a UI button, clear pending first
    # This prevents the "first click does nothing" bug in menus like Boost Tools.
    if context.user_data.get("pending_feature") and is_ui_button(q):
        clear_pending_feature(context)

    if is_back_message(q):

    # ⬅️ Back from Boost Tools
        if context.user_data.get("in_tools"):
            context.user_data.pop("in_tools", None)
            clear_pending_feature(context)
            stop_eval_mode(context)
    
            await update.message.reply_text(
                "Back to main menu.",
                reply_markup=main_menu_keyboard(),
            )
            return   # 🔴 THIS RETURN IS CRITICAL
    
        # ⬅️ Back from anywhere else
        clear_pending_feature(context)
        stop_eval_mode(context)
    
        await update.message.reply_text(
            "Back to main menu.",
            reply_markup=main_menu_keyboard(),
        )
        return

    topic_before = get_current_topic(context)
    record_event(user.id, topic_before, kind="message")

    try:
        mem = get_user_memory_cached(update, context)
        merge_usage_into_memory(context, mem)
        mem["history"]["message_count"] = int(mem["history"].get("message_count", 0) or 0) + 1
        mem["history"]["last_active"] = int(__import__("time").time())
        mem["history"]["last_topic"] = topic_before
        # cheap profile extraction from user message
        mem["profile"] = extract_profile_signals(q, mem.get("profile", {}) or {})
        persist_user_memory(update, context)
    except Exception as e:
        logging.exception("Memory update failed (non-fatal): %s", e)



    # quick cancel (optional)
    if q.lower() in {"cancel", "stop", "exit"} and context.user_data.get("eval_active", False):
        stop_eval_mode(context)
        await update.message.reply_text("✅ Stopped evaluation follow-up mode. (Your last evaluated text is still saved for Boost Tools.)")
        return

    # ---- BOOST TOOLS MENU ----
    if q == BTN_TOOLS:
        clear_pending_feature(context)
        stop_eval_mode(context)
    
        # mark UI-only action
        context.user_data["in_tools"] = True
    
        # SAFE topic resolution
        last_section = context.user_data.get("last_used_section", DEFAULT_TOPIC)
        current = FRIENDLY_TOPIC_NAMES.get(last_section, "General")
    
        first_time = not context.user_data.get("saw_boost_explainer", False)
        context.user_data["saw_boost_explainer"] = True
    
        extra = ""
        if first_time:
            extra = (
                "\nℹ️ Boost Tools always use the section you last opened "
                "(Essays, ECs, IELTS, etc.).\n"
                "Switch sections first if you want different tips.\n"
            )
    
        await update.message.reply_text(
            "🚀 Boost Tools\n\n"
            "These tools adapt to your last used section:\n"
            f"👉 {current}\n"
            f"{extra}\n"
            "Pick a tool:",
            reply_markup=tools_menu_keyboard(),
        )
        return

    if q in {BTN_PROGRESS, BTN_INSIDER, BTN_POWERWORDS, BTN_PREDICT, BTN_WOWFACTOR, BTN_WOW_USE_LAST, BTN_WOW_PASTE_NEW}:
        # If the user was in another pending mode, clicking a tool should override it.
        clear_pending_feature(context)
        context.user_data["in_tools"] = True

        if q == BTN_PROGRESS:
            await tool_my_progress(update, context)
            return
        if q == BTN_INSIDER:
            await tool_insider_tips(update, context)
            return
        if q == BTN_POWERWORDS:
            await tool_power_words(update, context)
            return
        if q == BTN_PREDICT:
            await tool_predict_chances(update, context)
            return

        if q == BTN_WOWFACTOR:
            last = (context.user_data.get("last_eval_text") or "").strip()
            last_topic = (context.user_data.get("last_eval_topic") or get_current_topic(context) or DEFAULT_TOPIC)
            last_label = FRIENDLY_TOPIC_NAMES.get(
                last_topic,
                FRIENDLY_TOPIC_NAMES.get(get_current_topic(context), "General")
            )

            if len(last) >= 120:
                set_pending_feature(context, "wowfactor_confirm")
                await update.message.reply_text(
                    f"🔍 Find Wow Factor\n\nI found your last evaluated text ({last_label}).\nUse it?",
                    reply_markup=wowfactor_confirm_keyboard(),
                )
                return

            set_pending_feature(context, "wowfactor")
            await update.message.reply_text(
                "🔍 Find Wow Factor\n\nPaste your essay/paragraph (120+ words).\nTip: After you run an evaluation, I can use that text automatically.",
                reply_markup=tools_menu_keyboard(),
            )
            return

        if q == BTN_WOW_USE_LAST:
            last = (context.user_data.get("last_eval_text") or "").strip()
            if len(last) < 120:
                set_pending_feature(context, "wowfactor")
                await update.message.reply_text(
                    "I don't have a saved evaluated text yet. Paste your essay/paragraph (120+ words).",
                    reply_markup=tools_menu_keyboard(),
                )
                return
            clear_pending_feature(context)
            await run_wowfactor(update, context, last)
            return

        if q == BTN_WOW_PASTE_NEW:
            set_pending_feature(context, "wowfactor")
            await update.message.reply_text(
                "📝 Paste your essay/paragraph (120+ words) for Wow Factor analysis.",
                reply_markup=tools_menu_keyboard(),
            )
            return

    # --- pending feature input ---
    # Allow a direct portfolio-based school suggestion even when a pending feature is set.
    if q == BTN_SF_FROM_PORT:
        clear_pending_feature(context)
        clear_eval_context(context)
        await run_schoolfinder_from_portfolio(update, context)
        return

    if q == BTN_SF_MANUAL:
        clear_pending_feature(context)
        clear_eval_context(context)
        context.user_data["topic"] = "school_finder"
        track_topic(context, "school_finder")
        set_pending_feature(context, "schoolfinder")
        await send_with_image(
            update,
            "🏫 School Finder mode ON.\n\nSend me your GPA (or approximate), test scores (if any), budget, target countries, intended major, and constraints (e.g. need scholarship).",
            reply_markup=schoolfinder_keyboard(),
            image_key="schoolfinder_main",
        )
        return

    pending = context.user_data.get("pending_feature")
    # If the user clicks a UI button while in a pending mode, treat it as navigation
    # (do NOT treat the button label as the user's input).
    if pending and is_ui_button(q):
        clear_pending_feature(context)
    elif pending:
        clear_pending_feature(context)
        if pending == "brainstorm":
            await run_brainstorm(update, context, q)
            return
        if pending == "rewrite":
            await run_rewrite(update, context, q)
            return
        if pending == "plan":
            await run_plan(update, context, q)
            return
        if pending == "recpacket":
            await run_recpacket(update, context, q)
            return
        if pending == "schoolfinder":
            await run_schoolfinder(update, context, q)
            return
        if pending == "portfolioideas":
            await run_portfolioideas(update, context, q)
            return
        if pending == "app_tests":
            await run_app_tests(update, context, q)
            return
        if pending == "app_essays":
            await run_app_essays(update, context, q)
            return
        if pending == "app_ecs":
            await run_app_ecs(update, context, q)
            return
        if pending == "app_awards":
            await run_app_awards(update, context, q)
            return
        if pending == "app_prefs":
            await run_app_prefs(update, context, q)
            return
        if pending == "app_wellness":
            await run_app_wellness(update, context, q)
            return
        if pending == "advisor_mode":
            await run_advisor_mode(update, context, q)
            return
        if pending == "app_ecs":
            await run_app_ecs(update, context, q)
            return
        if pending == "app_awards":
            await run_app_awards(update, context, q)
            return
        if pending == "app_prefs":
            await run_app_prefs(update, context, q)
            return
        if pending == "app_wellness":
            await run_app_wellness(update, context, q)
            return
        if pending == "wowfactor_confirm":
            # User typed instead of tapping a button. Treat it as new text if long enough.
            if len((q or '').strip()) >= 120:
                await run_wowfactor(update, context, q)
                return
            await update.message.reply_text(
                "Tap ✅ Use last evaluated text, or paste a new text (120+ words).",
                reply_markup=wowfactor_confirm_keyboard(),
            )
            set_pending_feature(context, "wowfactor_confirm")
            return
        if pending == "wowfactor":
            if len((q or "").strip()) < 120:
                await update.message.reply_text(
                    "Please paste a bit more text (120+ words) so I can detect a real wow factor.",
                    reply_markup=tools_menu_keyboard(),
                )
                set_pending_feature(context, "wowfactor")
                return
            await run_wowfactor(update, context, q)
            return

    # ---- MAIN MENUS ----
    if q == BTN_ESSAY:
        clear_eval_context(context)
        await send_with_image(
            update,
            "Essays selected.\nChoose Personal Statement or Supplemental Essays.",
            reply_markup=essay_main_keyboard(),
            image_key="essays_main",
        )
        return

    if q == BTN_SAT:
        clear_eval_context(context)
        await send_with_image(
            update,
            "SAT selected. Choose a section:",
            reply_markup=sat_menu_keyboard(),
            image_key="sat_main",
        )
        return

    if q == BTN_IELTS:
        clear_eval_context(context)
        await send_with_image(
            update,
            "IELTS selected. Choose a section.",
            reply_markup=ielts_main_keyboard(),
            image_key="ielts_main",
        )
        return

    if q == BTN_PLAN_MAIN:
        clear_eval_context(context)
        context.user_data["topic"] = "application_plan"
        track_topic(context, "application_plan")
        # If the user already filled some info in My Application Portfolio, offer a 1-tap plan.
        try:
            mem = get_user_memory_cached(update, context)
            merge_usage_into_memory(context, mem)
            if has_application_portfolio_data(mem):
                clear_pending_feature(context)
                await send_with_image(
                    update,
                    "📅 Application Plan\n\nI can create a plan using your saved 📁 My Application Portfolio, or you can enter info manually.",
                    reply_markup=plan_choice_keyboard(),
                    image_key="plan_main",
                )
                return

        except Exception:
            pass

        # No (usable) portfolio yet -> ask the standard questions.
        set_pending_feature(context, "plan")
        await send_with_image(
            update,
            "📅 Application Plan mode ON.\n\nTell me your grade, target countries, intended major, test scores (if any), and your rough deadlines.",
            reply_markup=plan_keyboard(),
            image_key="plan_main",
        )
        return

    if q == BTN_PLAN_FROM_PORT:
        clear_eval_context(context)
        context.user_data["topic"] = "application_plan"
        track_topic(context, "application_plan")
        clear_pending_feature(context)
        stop_eval_mode(context)
        await run_plan_from_portfolio(update, context)
        return

    if q == BTN_PLAN_MANUAL:
        clear_eval_context(context)
        context.user_data["topic"] = "application_plan"
        track_topic(context, "application_plan")
        set_pending_feature(context, "plan")
        await send_with_image(
            update,
            "📅 Application Plan mode ON.\n\nTell me your grade, target countries, intended major, test scores (if any), and your rough deadlines.",
            reply_markup=plan_keyboard(),
            image_key="plan_main",
        )
        return

    if q == BTN_SF_MAIN:
        clear_eval_context(context)
        context.user_data["topic"] = "school_finder"
        track_topic(context, "school_finder")

        # If the user already has portfolio data, show the same 2-choice submenu as Application Plan.
        try:
            mem = get_user_memory_cached(update, context)
            merge_usage_into_memory(context, mem)
            if has_application_portfolio_data(mem):
                clear_pending_feature(context)
                await send_with_image(
                    update,
                    "🏫 School Finder\n\nChoose one:\n• 📌 Use My Portfolio\n• ✍️ Enter Details Manually",
                    reply_markup=schoolfinder_keyboard(),
                    image_key="schoolfinder_main",
                )
                return
        except Exception:
            pass

        # No (usable) portfolio yet -> go straight to manual input.
        set_pending_feature(context, "schoolfinder")
        await send_with_image(
            update,
            "🏫 School Finder mode ON.\n\nSend me your GPA (or approximate), test scores (if any), budget, target countries, intended major, and constraints (e.g. need scholarship).",
            reply_markup=schoolfinder_keyboard(),
            image_key="schoolfinder_main",
        )
        return

    # ---- TOPIC BUTTONS ----
    if q == BTN_ESSAY_PS:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        await send_with_image(
            update,
            "Great, let's work on your Personal Statement.\nAsk about structure, voice, and storytelling.",
            reply_markup=essay_ps_keyboard(),
            image_key="essays_personal",
        )
        await update.message.reply_text(
            "For detailed feedback, tap '✅ Personal Statement Evaluation' and then upload PDF/DOCX or paste the text.\n\n"
            "After I evaluate, you can ask ANY follow-up question about your essay (grammar, clarity, wording, etc.)."
        )
        return

    if q == BTN_ESSAY_SUPP:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        await send_with_image(
            update,
            "Great, let's work on your Supplemental Essays.\nAsk about 'Why us', community essays, and short prompts.",
            reply_markup=essay_supp_keyboard(),
            image_key="essays_supplemental",
        )
        await update.message.reply_text(
            "For detailed feedback, tap '✅ Supplemental Essay Evaluation' and then upload PDF/DOCX or paste the text.\n\n"
            "After I evaluate, you can ask ANY follow-up question about your essay."
        )
        return

    if q == BTN_EC:
        clear_eval_context(context)
        context.user_data["topic"] = "extracurriculars"

        track_topic(context, "extracurriculars")
        await send_with_image(
            update,
            "Great, let's talk about your Extracurricular activities.\nAsk how to present impact, leadership, and long-term involvement.",
            reply_markup=ec_keyboard(),
            image_key="extracurriculars",
        )
        await update.message.reply_text(
            "For feedback, click '✅ Extracurricular Evaluation' and then upload PDF/DOCX or paste your EC descriptions."
        )
        return

    if q == BTN_EC_PROGRAMS:
        await update.message.reply_text(
            "🌍 Here is a list of top extracurricular programs and opportunities:\n\n"
            "https://docs.google.com/spreadsheets/d/1D-UlJGrg32Ib-9Rvm9y7lKkE6jkx3EK-Kb_qJ6G3tos/edit?usp=sharing\n"
        )
        return

    if q == BTN_REC:
        clear_eval_context(context)
        context.user_data["topic"] = "recommendations"

        track_topic(context, "recommendations")
        await send_with_image(
            update,
            "Great, let's work on Recommendation Letters.\nAsk how to request them and what makes a strong letter.",
            reply_markup=rec_keyboard(),
            image_key="recommendations",
        )
        await update.message.reply_text(
            "You can:\n- Tap '✅ Rec Letter Evaluation' to get feedback on a draft.\n- Tap '📄 Rec Letter Packet' to build a brag sheet."
        )
        return

    if q == BTN_PORT:
        clear_eval_context(context)
        context.user_data["topic"] = "portfolio"

        track_topic(context, "portfolio")
        await send_with_image(
            update,
            "You're now in Portfolio.\nAsk about structure and how to present your work.",
            reply_markup=portfolio_keyboard(),
            image_key="portfolio",
        )
        await update.message.reply_text(
            "You can upload PDF/DOCX for detailed feedback (✅ Portfolio Evaluation), tap '💡 Portfolio Ideas', or open 📂 My Application Portfolio."
        )
        return

    if q in {BTN_SAT_MATH, BTN_SAT_ENGLISH}:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        nice_name = "SAT Math" if topic == "sat_math" else "SAT English"
        image_key = "sat_math" if topic == "sat_math" else "sat_english"
        await send_with_image(
            update,
            f"You're now in {nice_name}. Ask anything.",
            reply_markup=sat_menu_keyboard(),
            image_key=image_key,
        )
        return

    if q in {BTN_IELTS_READING, BTN_IELTS_LISTENING, BTN_IELTS_SPEAKING}:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        await send_with_image(
            update,
            f"You're now in {FRIENDLY_TOPIC_NAMES.get(topic, topic)}. Ask anything.",
            reply_markup=ielts_main_keyboard(),
            image_key=topic,
        )
        return

    if q == BTN_IELTS_WRITING:
        clear_eval_context(context)
        context.user_data["topic"] = "ielts_writing"

        track_topic(context, "ielts_writing")
        track_topic(context, "ielts_writing")
        await send_with_image(
            update,
            "You're now in IELTS Writing.\nAsk about Task 1/2, band 7+ strategies, or send your answer for feedback.",
            reply_markup=ielts_writing_keyboard(),
            image_key="ielts_writing",
        )
        await update.message.reply_text(
            "For evaluation, click '✅ Writing Evaluation' then send your answer as text, PDF/DOCX, or a clear photo.\n\n"
            "After evaluation, you can ask ANY follow-up question about your writing."
        )
        return

    # ---- FEATURE BUTTONS ----
    if q == BTN_BRAINSTORM:
        set_pending_feature(context, "brainstorm")
        await update.message.reply_text(
            "🧠 Brainstorm mode ON.\n\nBriefly tell me about yourself, your target major, and what you want to write about."
        )
        return

    if q == BTN_REWRITE:
        set_pending_feature(context, "rewrite")
        await update.message.reply_text("✍️ Rewrite mode ON.\n\nSend me the paragraph or essay you want me to improve.")
        return

    if q == BTN_REC_PACKET:
        set_pending_feature(context, "recpacket")
        await update.message.reply_text(
            "✉️ Rec Letter Packet mode ON.\n\nTell me which teacher will write your rec, what classes you took, your achievements, and what you want highlighted."
        )
        return

    if q == BTN_PORTFOLIO_IDEAS:
        set_pending_feature(context, "portfolioideas")
        await update.message.reply_text(
            "💡 Portfolio Ideas mode ON.\n\nTell me your field (CS, design, art, film, etc.), your skills, and target programs."
        )
        return

    # ---- APPLICATION PORTFOLIO (inside Portfolio) ----
    if q == BTN_APP_PORT:
        # Keep topic as portfolio for consistency
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        stop_eval_mode(context)
        await app_portfolio_show_menu(update, context)
        return

    if q == BTN_BACK_PORT:
        clear_pending_feature(context)
        stop_eval_mode(context)
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        await send_with_image(
            update,
            "You're now in Portfolio.\nAsk about structure and how to present your work.",
            reply_markup=portfolio_keyboard(),
            image_key="portfolio",
        )
        await update.message.reply_text(
            "You can upload PDF/DOCX for detailed feedback (✅ Portfolio Evaluation), tap '💡 Portfolio Ideas', or open 📂 My Application Portfolio."
        )
        return

    if q == BTN_APP_TESTS:
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        await run_app_tests(update, context, None)
        return

    if q == BTN_APP_ESSAYS:
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        await run_app_essays(update, context, None)
        return

    if q == BTN_APP_ADVISOR:
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        set_pending_feature(context, "advisor_mode")
        await update.message.reply_text(
            "🧭 Advisor Mode ON.\n\n"
            "Tell me: target countries, intended major, budget/need aid, and a rough school list (optional).\n"
            "If you want, just type: `use saved portfolio`.",
            reply_markup=app_portfolio_keyboard(),
        )
        return

    if q == BTN_APP_ECS:
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        await run_app_ecs(update, context, None)
        return

    if q == BTN_APP_AWARDS:
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        await run_app_awards(update, context, None)
        return

    if q == BTN_APP_PREFS:
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        await run_app_prefs(update, context, None)
        return

    if q == BTN_APP_WELLNESS:
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        clear_pending_feature(context)
        await run_app_wellness(update, context, None)
        return

    if q == BTN_APP_READINESS:
        mem = get_user_memory_cached(update, context)
        merge_usage_into_memory(context, mem)
        score, reasons = compute_portfolio_readiness(mem)
        app = _ensure_application_defaults(mem)
        app.setdefault("readiness", {})
        app["readiness"]["last_score"] = score
        app["readiness"]["last_updated"] = int(__import__("time").time())
        mem["application"] = app
        persist_user_memory(update, context)

        why = ("; ".join(reasons) if reasons else "Not enough saved yet.")
        await update.message.reply_text(
            "✅ READINESS CHECK (Portfolio)\n\n"
            f"Score: {_progress_bar(score)}\n"
            f"Based on: {why}\n\n"
            "Next step: update 🧾 Test Scores + GPA and 📝 Essays Status, then run 🧭 Advisor Mode.",
            reply_markup=app_portfolio_keyboard(),
        )
        return

    # ---- EVALUATION BUTTONS ----
    if q == BTN_PS_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "essays_personal"

        track_topic(context, "essays_personal")
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "essays_personal"
        await update.message.reply_text(
            "Personal Statement Evaluation mode ON ✅\n\nNow paste your Personal Statement (100+ words) or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question about your essay (grammar, clarity, structure, etc.).\n"
            "If you say things like 'apply this feedback' or 'rewrite the conclusion', I will revise the text."
        )
        return

    if q == BTN_SUPP_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "essays_supplemental"

        track_topic(context, "essays_supplemental")
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "essays_supplemental"
        await update.message.reply_text(
            "Supplemental Essay Evaluation mode ON ✅\n\nNow paste your essay (100+ words) or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question about your essay.\n"
            "If you say 'apply this feedback' or 'rewrite the ending', I will revise the text."
        )
        return

    if q == BTN_EC_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "extracurriculars"
        track_topic(context, "extracurriculars")
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "extracurriculars"
        await update.message.reply_text(
            "Extracurricular Evaluation mode ON ✅\n\nNow paste your EC descriptions or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    if q == BTN_REC_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "recommendations"
        track_topic(context, "recommendations")
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "recommendations"
        await update.message.reply_text(
            "Rec Letter Evaluation mode ON ✅\n\nNow paste the draft letter or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    if q == BTN_IW_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "ielts_writing"
        track_topic(context, "ielts_writing")
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "ielts_writing"
        await update.message.reply_text(
            "IELTS Writing Evaluation mode ON ✅\n\nSend your answer as text, PDF/DOCX, or a clear photo.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    if q == BTN_PORT_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "portfolio"
        await update.message.reply_text(
            "Portfolio Evaluation mode ON ✅\n\nNow paste your portfolio description or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    # ---- EVALUATION FLOW (supports follow-ups + Q&A) ----
    if context.user_data.get("eval_active", False):
        last_text = (context.user_data.get("last_eval_text") or "").strip()

        # pasted a new submission while eval mode is ON
        if looks_like_submission(q):
            await evaluate_text_for_topic(update, context)
            return

        # eval mode ON but no submission yet
        if not last_text:
            await update.message.reply_text(
                "✅ Evaluation mode is ON.\n\n"
                "Paste your full text here (100+ words) or upload a PDF/DOCX.\n"
                "After I evaluate, you can ask ANY follow-up question about your text."
            )
            return

        # ✅ If it looks like a rewrite/apply request -> do rewrite
        if is_followup_intent(q):
            await run_eval_followup(update, context, q)
            return

        # ✅ Otherwise: answer ANY follow-up question using saved text + feedback
        await update.message.reply_text("Got it! Thinking about your question…")
        await run_eval_qa(update, context, q)
        return

    # ---- NORMAL Q&A WITH RAG ----
    await show_typing(update, context)
    await update.message.reply_text("Got it! Thinking about your question…")

    topic = get_current_topic(context)
    
    col = None

    docs = []
    col = None

    # Optional latency optimization: skip RAG for low-signal messages
    use_rag = _should_use_rag(context, q)

    if use_rag:
        try:
            col = get_collection(chat_id, topic)
            emb = _get_cached_query_embedding(context, q)

            if emb is not None:
                results = col.query(query_embeddings=[emb], where={"type": "qa"}, n_results=4)
            else:
                results = col.query(query_texts=[q], where={"type": "qa"}, n_results=4)

            docs = results.get("documents", [[]])[0]
            metas = results.get("metadatas", [[]])[0]
            _save_last_retrieval(
                context,
                topic=topic,
                mode="qa_primary",
                query=q,
                items=[
                    {
                        "bucket": "qa",
                        "title": (m or {}).get("title", "Untitled"),
                        "type": (m or {}).get("type", "qa"),
                        "part": (m or {}).get("part", "—"),
                        "preview": (d or "")[:160],
                    }
                    for d, m in zip(docs or [], metas or [])
                ],
            )
        except Exception as e:
            logging.error(f"Error querying collection: {e}")
            docs = []

        if not docs and col is not None:
            try:
                # Fallback without type filter
                emb = _get_cached_query_embedding(context, q)
                if emb is not None:
                    results = col.query(query_embeddings=[emb], n_results=4)
                else:
                    results = col.query(query_texts=[q], n_results=4)
                docs = results.get("documents", [[]])[0]
                metas = results.get("metadatas", [[]])[0]
                _save_last_retrieval(
                    context,
                    topic=topic,
                    mode="qa_fallback",
                    query=q,
                    items=[
                        {
                            "bucket": (m or {}).get("type", "unknown"),
                            "title": (m or {}).get("title", "Untitled"),
                            "type": (m or {}).get("type", "unknown"),
                            "part": (m or {}).get("part", "—"),
                            "preview": (d or "")[:160],
                        }
                        for d, m in zip(docs or [], metas or [])
                    ],
                )
            except Exception as e:
                logging.error(f"Error querying collection (fallback): {e}")
                docs = []

    if use_rag and not docs:
        _save_last_retrieval(context, topic=topic, mode="qa_none", query=q, items=[])

    context_block = "\n\n---\n\n".join(docs or [])
    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)
    mem = get_user_memory_cached(update, context)
    merge_usage_into_memory(context, mem)
    # update profile signals from the question too (cheap heuristic)
    mem["profile"] = extract_profile_signals(q, mem.get("profile", {}) or {})
    decision_notes = coach_decision_notes(topic, q, mem)

    sys = coach_qa_system_prompt(topic, mem, decision_notes)

    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": f"Reference knowledge (do not copy formatting):\n{context_block}"},
        {"role": "user", "content": q},
    ]

    a = await openai_chat(model=FAST_MODEL, messages=messages, temperature=0.4)
    await send_long(update, a)

# -------- Dummy HTTP server for Render/Railway --------
def start_dummy_server():
    class Handler(BaseHTTPRequestHandler):
        def do_GET(self):
            self.send_response(200)
            self.send_header("Content-type", "text/plain")
            self.end_headers()
            self.wfile.write(b"OK")

        def log_message(self, format, *args):
            return

    port = int(os.environ.get("PORT", 10000))
    server = HTTPServer(("0.0.0.0", port), Handler)
    print(f"✅ Health check server running on port {port}")
    # Tip: to reduce Railway cold-start latency, ping this service every 5–10 minutes (e.g., UptimeRobot).
    server.serve_forever()


# -------- Global error handler (prevents silent no-reply bugs) --------
async def error_handler(update, context):
    logging.exception("Unhandled exception while handling an update: %s", getattr(context, 'error', None))
    try:
        if update is not None and getattr(update, 'effective_message', None):
            await update.effective_message.reply_text("⚠️ I hit an internal error while processing that. Please try again.")
    except Exception:
        pass
# ===== Setup Application =====
app = ApplicationBuilder().token(TELEGRAM_TOKEN).concurrent_updates(True).build()
app.add_error_handler(error_handler)

# ===== GROUP -1: PAID ACCESS GATE (RUNS FIRST) =====
app.add_handler(MessageHandler(filters.ALL, paid_access_gate), group=-1)

# ===== GROUP 0: COMMANDS ONLY =====
app.add_handler(CommandHandler("start", start), group=0)
app.add_handler(CommandHandler("pay", pay_cmd), group=0)
app.add_handler(CommandHandler("id", id_cmd), group=0)
app.add_handler(CommandHandler("mysub", mysub_cmd), group=0)
app.add_handler(CommandHandler("activate", activate_cmd), group=0)
app.add_handler(CommandHandler("deactivate", deactivate_cmd), group=0)
app.add_handler(CommandHandler("paidusers", paidusers_cmd), group=0)
app.add_handler(CommandHandler("teach", teach), group=0)
app.add_handler(CommandHandler("teachrubric", teachrubric), group=0)
app.add_handler(CommandHandler("teachfile", teachfile), group=0)
app.add_handler(CommandHandler("teachfile_eval", teachfile_eval), group=0)
app.add_handler(CommandHandler("teachlink", teachlink), group=0)
app.add_handler(CommandHandler("teachlink_eval", teachlink_eval), group=0)
app.add_handler(CommandHandler("teachimage", teachimage), group=0)
app.add_handler(CommandHandler("sources", sources), group=0)
app.add_handler(CommandHandler("sources_all", sources_all), group=0)
app.add_handler(CommandHandler("unlearn", unlearn), group=0)
app.add_handler(CommandHandler("clear", clear), group=0)
app.add_handler(CommandHandler("stats", stats_cmd), group=0)
app.add_handler(CommandHandler("brainstorm", brainstorm_cmd), group=0)
app.add_handler(CommandHandler("rewrite", rewrite_cmd), group=0)
app.add_handler(CommandHandler("plan", plan_cmd), group=0)
app.add_handler(CommandHandler("recpacket", recpacket_cmd), group=0)
app.add_handler(CommandHandler("schoolfinder", schoolfinder_cmd), group=0)
app.add_handler(CommandHandler("portfolioideas", portfolioideas_cmd), group=0)
app.add_handler(CommandHandler("backup_sources", backup_sources), group=0)
app.add_handler(CommandHandler("health", health), group=0)
app.add_handler(CommandHandler("debugsources", debugsources_cmd), group=0)
app.add_handler(CommandHandler("help", help_cmd))
app.add_handler(CommandHandler("profile", profile_cmd))
app.add_handler(CommandHandler("feedback", feedback_cmd))
app.add_handler(CommandHandler("how_to_use", how_to_use_cmd))


# ===== GROUP 1: FILES / PHOTOS =====
app.add_handler(MessageHandler(filters.Document.ALL, document_router), group=1)
app.add_handler(MessageHandler(filters.PHOTO, photo_router), group=1)

# ===== GROUP 2: NORMAL TEXT (ABSOLUTELY LAST) =====
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, answer), group=2)

# ===== Daily subscription reminders (3-day + expired) =====
try:
    app.job_queue.run_repeating(subscription_reminder_job, interval=86400, first=120)
except Exception as e:
    logging.error(f"Failed to schedule subscription reminders: {e}")


print(
    "\n" + "="*80 + "\n"
    "Bot is running with GLOBAL per-topic RAG + metadata separation (qa/evaluation) + submenus "
    "+ eval follow-ups (apply feedback) + eval Q&A (any follow-up question) + embedded tools "
    "+ Application Plan & School Finder + analytics + admin locks + backup + health "
    "+ UUID IDs + robust command parsing (incl captions)…\n"
    "="*80 + "\n"
)

if __name__ == "__main__":
    threading.Thread(target=start_dummy_server, daemon=True).start()
    print("✅ Bot starting...")
    app.run_polling()