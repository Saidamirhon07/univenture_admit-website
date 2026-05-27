# AIBOT.py
from telegram import Update, KeyboardButton, ReplyKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters,
)
from dotenv import load_dotenv
from openai import OpenAI
import chromadb
from chromadb.utils import embedding_functions
import os, io, asyncio, nest_asyncio, logging, json, base64
from collections import Counter

# -------- File extraction deps --------
from pdfminer.high_level import extract_text
from docx import Document as DocxDocument

# -------- Web page extraction (for teachlink) --------
import trafilatura

logging.basicConfig(level=logging.INFO)
load_dotenv()

TELEGRAM_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
oa = OpenAI(api_key=OPENAI_API_KEY)

# -------- Persistent local Chroma --------
chroma = chromadb.PersistentClient(path="./chroma_store")
emb_fn = embedding_functions.OpenAIEmbeddingFunction(
    api_key=OPENAI_API_KEY,
    model_name="text-embedding-3-small",
)

# -------- Analytics storage --------
STATS_FILE = "analytics.json"


def _default_stats():
    return {
        "users": [],
        "messages_total": 0,
        "messages_per_user": {},
        "topic_counts": {},
        "eval_counts": {},
    }


def load_stats():
    if not os.path.exists(STATS_FILE):
        return _default_stats()
    try:
        with open(STATS_FILE, "r") as f:
            data = json.load(f)
    except Exception:
        return _default_stats()

    base = _default_stats()
    base.update({k: data.get(k, v) for k, v in base.items()})
    return base


def save_stats(stats):
    try:
        with open(STATS_FILE, "w") as f:
            json.dump(stats, f, indent=2)
    except Exception as e:
        logging.warning(f"Could not save analytics: {e}")


def record_event(user_id, topic: str, kind: str = "message"):
    stats = load_stats()
    uid = str(user_id)

    if uid not in stats["users"]:
        stats["users"].append(uid)

    stats["messages_total"] = stats.get("messages_total", 0) + 1

    mpu = stats.setdefault("messages_per_user", {})
    mpu[uid] = mpu.get(uid, 0) + 1

    if topic:
        tc = stats.setdefault("topic_counts", {})
        tc[topic] = tc.get(topic, 0) + 1

    if kind == "eval" and topic:
        ec = stats.setdefault("eval_counts", {})
        ec[topic] = ec.get(topic, 0) + 1

    save_stats(stats)


# -------- Main menu buttons --------
BTN_ESSAY = "📝 Essays"
BTN_EC = "🎯 Extracurricular activities"
BTN_REC = "✉️ Recommendation Letters"
BTN_SAT = "📈 SAT"
BTN_IELTS = "🗣️ IELTS"
BTN_PORT = "🖼️ Portfolio Check"
BTN_PLAN_MAIN = "📅 Application Plan"
BTN_SF_MAIN = "🏫 School Finder"

# Evaluation sub-buttons
BTN_PS_EVAL = "✅ Personal Statement Evaluation"
BTN_SUPP_EVAL = "✅ Supplemental Essay Evaluation"
BTN_EC_EVAL = "✅ Extracurricular Evaluation"
BTN_REC_EVAL = "✅ Rec Letter Evaluation"
BTN_IW_EVAL = "✅ Writing Evaluation"
BTN_PORT_EVAL = "✅ Portfolio Evaluation"

# Essays sub-buttons
BTN_ESSAY_PS = "📝 Personal Statement"
BTN_ESSAY_SUPP = "📝 Supplemental Essays"

# Extra tools (as buttons inside menus)
BTN_BRAINSTORM = "🧠 Brainstorm ideas"
BTN_REWRITE = "✍️ Rewrite my text"
BTN_REC_PACKET = "📄 Rec Letter Packet"
BTN_PORTFOLIO_IDEAS = "💡 Portfolio Ideas"

# SAT sub-buttons
BTN_SAT_MATH = "📐 SAT Math"
BTN_SAT_ENGLISH = "📚 SAT English"

# IELTS sub-buttons
BTN_IELTS_READING = "📖 IELTS Reading"
BTN_IELTS_LISTENING = "👂 IELTS Listening"
BTN_IELTS_WRITING = "✍️ IELTS Writing"
BTN_IELTS_SPEAKING = "🗣️ IELTS Speaking"

# Back button
BTN_BACK = "⬅️ Back"


# -------- Keyboards --------
def main_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_ESSAY), KeyboardButton(BTN_EC)],
            [KeyboardButton(BTN_REC), KeyboardButton(BTN_SAT)],
            [KeyboardButton(BTN_IELTS), KeyboardButton(BTN_PORT)],
            [KeyboardButton(BTN_PLAN_MAIN), KeyboardButton(BTN_SF_MAIN)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def main_menu_with_back():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_ESSAY), KeyboardButton(BTN_EC)],
            [KeyboardButton(BTN_REC), KeyboardButton(BTN_SAT)],
            [KeyboardButton(BTN_IELTS), KeyboardButton(BTN_PORT)],
            [KeyboardButton(BTN_PLAN_MAIN), KeyboardButton(BTN_SF_MAIN)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def essay_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_ESSAY_PS), KeyboardButton(BTN_ESSAY_SUPP)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def essay_ps_keyboard():
    # PS menu: Eval + Brainstorm + Rewrite + Back
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PS_EVAL)],
            [KeyboardButton(BTN_BRAINSTORM)],
            [KeyboardButton(BTN_REWRITE)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def essay_supp_keyboard():
    # Supp menu: Eval + Brainstorm + Rewrite + Back
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_SUPP_EVAL)],
            [KeyboardButton(BTN_BRAINSTORM)],
            [KeyboardButton(BTN_REWRITE)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def ec_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_EC_EVAL)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def rec_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_REC_EVAL)],
            [KeyboardButton(BTN_REC_PACKET)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def sat_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_SAT_MATH), KeyboardButton(BTN_SAT_ENGLISH)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def ielts_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_IELTS_READING), KeyboardButton(BTN_IELTS_LISTENING)],
            [KeyboardButton(BTN_IELTS_WRITING), KeyboardButton(BTN_IELTS_SPEAKING)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def ielts_writing_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_IELTS_READING), KeyboardButton(BTN_IELTS_LISTENING)],
            [KeyboardButton(BTN_IELTS_SPEAKING), KeyboardButton(BTN_IW_EVAL)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


def portfolio_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PORT_EVAL)],
            [KeyboardButton(BTN_PORTFOLIO_IDEAS)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )


# -------- Topic mapping & helpers --------
TOPIC_KEYS = {
    BTN_ESSAY_PS: "essays_personal",
    BTN_ESSAY_SUPP: "essays_supplemental",
    BTN_EC: "extracurriculars",
    BTN_REC: "recommendations",
    BTN_PORT: "portfolio",
    BTN_SAT_MATH: "sat_math",
    BTN_SAT_ENGLISH: "sat_english",
    BTN_IELTS_READING: "ielts_reading",
    BTN_IELTS_LISTENING: "ielts_listening",
    BTN_IELTS_WRITING: "ielts_writing",
    BTN_IELTS_SPEAKING: "ielts_speaking",
}

# Topics where uploads are evaluated (not used as training)
EVAL_TOPICS = {
    "essays_personal",
    "essays_supplemental",
    "recommendations",
    "portfolio",
    "extracurriculars",
}

FRIENDLY_TOPIC_NAMES = {
    "essays_personal": "Personal Statement",
    "essays_supplemental": "Supplemental Essays",
    "extracurriculars": "Extracurricular Activities",
    "recommendations": "Recommendation Letters",
    "portfolio": "Portfolio",
    "sat_math": "SAT Math",
    "sat_english": "SAT English",
    "ielts_reading": "IELTS Reading",
    "ielts_listening": "IELTS Listening",
    "ielts_writing": "IELTS Writing",
    "ielts_speaking": "IELTS Speaking",
    "application_plan": "Application Plan",
    "school_finder": "School Finder",
    "general": "General admissions help",
}

DEFAULT_TOPIC = "general"


# -------- Image mapping --------
IMAGE_FILES = {
    "welcome": "images/welcome.png",
    "essays_main": "images/essays_main.png",
    "essays_personal": "images/essays_personal.png",
    "essays_supplemental": "images/essays_supplemental.png",
    "extracurriculars": "images/extracurriculars.png",
    "recommendations": "images/recommendations.png",
    "portfolio": "images/portfolio.png",
    "sat_main": "images/sat_main.png",
    "sat_math": "images/sat_math.png",
    "sat_english": "images/sat_english.png",
    "ielts_main": "images/ielts_main.png",
    "ielts_reading": "images/ielts_reading.png",
    "ielts_listening": "images/ielts_listening.png",
    "ielts_writing": "images/ielts_writing.png",
    "ielts_speaking": "images/ielts_speaking.png",
    "plan_main": "images/plan_main.png",
    "schoolfinder_main": "images/schoolfinder_main.png",
}


def get_current_topic(context: ContextTypes.DEFAULT_TYPE) -> str:
    return context.user_data.get("topic", DEFAULT_TOPIC)


def get_collection(chat_id: int, topic: str):
    return chroma.get_or_create_collection(
        name=f"chat_{chat_id}_{topic}",
        embedding_function=emb_fn,
    )


def _chunk(text: str, max_chars=1000, overlap=150):
    """Simple char-based chunker."""
    text = " ".join((text or "").split())
    if not text:
        return []
    chunks, i = [], 0
    while i < len(text):
        end = min(len(text), i + max_chars)
        chunks.append(text[i:end])
        if end == len(text):
            break
        i = max(0, end - overlap)
    return chunks


async def show_typing(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    await context.bot.send_chat_action(chat_id=chat_id, action="typing")


async def send_long(update: Update, text: str):
    MAX_LEN = 4000
    if not text:
        return
    for i in range(0, len(text), MAX_LEN):
        chunk = text[i : i + MAX_LEN]
        await update.message.reply_text(chunk)


async def send_with_image(
    update: Update,
    caption: str,
    reply_markup=None,
    image_key: str | None = None,
):
    if image_key and image_key in IMAGE_FILES:
        path = IMAGE_FILES[image_key]
        if os.path.exists(path):
            try:
                with open(path, "rb") as f:
                    await update.message.reply_photo(
                        photo=f,
                        caption=caption,
                        reply_markup=reply_markup,
                    )
                return
            except Exception as e:
                logging.warning(f"Failed to send image {path}: {e}")

    await update.message.reply_text(caption, reply_markup=reply_markup)


# -------- Vision helper --------
def extract_text_from_image_bytes(image_bytes: bytes) -> str:
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    data_url = f"data:image/jpeg;base64,{b64}"

    resp = oa.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract all readable text from this image. "
                            "Return ONLY the plain text, no extra comments."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": data_url},
                    },
                ],
            }
        ],
        temperature=0.0,
    )
    return resp.choices[0].message.content or ""


# -------- Handlers --------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["topic"] = DEFAULT_TOPIC
    context.user_data["eval_active"] = False
    context.user_data["pending_feature"] = None
    user = update.effective_user
    record_event(user.id, "start", kind="start")

    await send_with_image(
        update,
        "Hi! I'm your coached AI 🤖\nChoose a topic or ask a question.",
        reply_markup=main_menu_keyboard(),
        image_key="welcome",
    )


# ---------- TEACH (Q&A sources) ----------
async def teach(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teach")

    text = update.message.text or ""

    if "|" not in text:
        await update.message.reply_text(
            "Use format:\n/teach <title> | <content>\n\n"
            f"Current topic: {topic}"
        )
        return

    title, content = [p.strip() for p in text.split("|", 1)]
    col = get_collection(chat_id, topic)

    existing = col.get(where={"title": title})
    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' already exists in topic: {topic}. "
            "Use /unlearn '<title>' first if you want to replace it."
        )
        return

    doc_id = f"{chat_id}_{topic}_{col.count()+1}"
    col.add(
        ids=[doc_id],
        metadatas=[{"title": title, "topic": topic, "type": "qa", "source": "manual"}],
        documents=[content],
    )
    await update.message.reply_text(f"Learned '{title}' ✅ (topic: {topic}, mode: Q&A)")


# ---------- TEACH RUBRIC (EVALUATION sources) ----------
async def teachrubric(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Teach evaluation criteria / rubric as text, for evaluation mode."""
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachrubric")

    text = update.message.text or ""
    if "|" not in text:
        await update.message.reply_text(
            "Use format:\n/teachrubric <title> | <rubric / evaluation criteria>\n\n"
            f"Current topic: {topic}"
        )
        return

    title, content = [p.strip() for p in text.split("|", 1)]
    col = get_collection(chat_id, topic)

    existing = col.get(where={"title": title})
    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' already exists in topic: {topic}. "
            "Use /unlearn '<title>' first if you want to replace it."
        )
        return

    doc_id = f"{chat_id}_{topic}_{col.count()+1}"
    col.add(
        ids=[doc_id],
        metadatas=[
            {"title": title, "topic": topic, "type": "evaluation", "source": "manual"}
        ],
        documents=[content],
    )
    await update.message.reply_text(
        f"Learned evaluation rubric '{title}' ✅ (topic: {topic})"
    )


# ---------- TEACH FILE (Q&A sources) ----------
async def teachfile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Teach by sending a PDF or DOCX, for Q&A use."""
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachfile")

    await update.message.reply_text("Reading your file and extracting text to learn from it (Q&A)…")
    doc = update.message.document
    if not doc:
        await update.message.reply_text(
            "Attach a PDF or DOCX and write /teachfile in the caption to train me from it."
        )
        return

    tgfile = await doc.get_file()
    file_bytes = await tgfile.download_as_bytearray()
    name = (doc.file_name or "upload").lower()

    try:
        if name.endswith(".pdf"):
            text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for /teachfile.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    parts = _chunk(text)
    if not parts:
        await update.message.reply_text("I couldn’t find any readable text in that file.")
        return

    col = get_collection(chat_id, topic)

    existing = col.get(where={"title": name})
    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{name}' is already learned in topic: {topic}.\n"
            "Use /unlearn <title> to remove it first."
        )
        return

    base = f"{chat_id}_{topic}_{col.count()+1}"
    ids = [f"{base}_{i}" for i in range(len(parts))]
    metas = [
        {
            "title": name,
            "topic": topic,
            "part": i,
            "source": "file",
            "type": "qa",
        }
        for i in range(len(parts))
    ]
    col.add(ids=ids, metadatas=metas, documents=parts)

    await update.message.reply_text(
        f"Learned from file ✅ ({len(parts)} parts) in topic: {topic} (Q&A)"
    )


# ---------- TEACH FILE EVAL (EVALUATION sources) ----------
async def teachfile_eval(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Teach evaluation rubric by sending a PDF or DOCX."""
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachfile_eval")

    await update.message.reply_text("Reading your rubric file and extracting evaluation criteria…")
    doc = update.message.document
    if not doc:
        await update.message.reply_text(
            "Attach a PDF or DOCX and write /teachfile_eval in the caption to teach an evaluation rubric."
        )
        return

    tgfile = await doc.get_file()
    file_bytes = await tgfile.download_as_bytearray()
    name = (doc.file_name or "upload").lower()

    try:
        if name.endswith(".pdf"):
            text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for /teachfile_eval.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    parts = _chunk(text)
    if not parts:
        await update.message.reply_text("I couldn’t find any readable text in that file.")
        return

    col = get_collection(chat_id, topic)

    existing = col.get(where={"title": name})
    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{name}' is already learned in topic: {topic}.\n"
            "Use /unlearn <title> to remove it first."
        )
        return

    base = f"{chat_id}_{topic}_{col.count()+1}"
    ids = [f"{base}_eval_{i}" for i in range(len(parts))]
    metas = [
        {
            "title": name,
            "topic": topic,
            "part": i,
            "source": "file",
            "type": "evaluation",
        }
        for i in range(len(parts))
    ]
    col.add(ids=ids, metadatas=metas, documents=parts)

    await update.message.reply_text(
        f"Learned evaluation rubric from file ✅ ({len(parts)} parts) in topic: {topic}"
    )


# ---------- TEACH LINK (Q&A) ----------
async def teachlink(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachlink")

    msg_text = (update.message.text or "").strip()
    parts = msg_text.split(maxsplit=1)
    if len(parts) < 2:
        await update.message.reply_text("Use: /teachlink <url>")
        return

    url = parts[1].strip()
    await update.message.reply_text("Fetching content from link and learning from it (Q&A)…")

    try:
        downloaded = trafilatura.fetch_url(url)
        text = trafilatura.extract(downloaded)
    except Exception as e:
        await update.message.reply_text(f"Error while fetching the URL: {e}")
        return

    if not text:
        await update.message.reply_text("Couldn't extract readable text from that link.")
        return

    chunks = _chunk(text)
    if not chunks:
        await update.message.reply_text("The page did not contain enough text to learn from.")
        return

    col = get_collection(chat_id, topic)

    existing = col.get(where={"title": url})
    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"This link is already learned in topic: {topic}.\n"
            "Use /unlearn <url> to remove it first."
        )
        return

    base = f"{chat_id}_{topic}_{col.count()+1}"
    ids = [f"{base}_{i}" for i in range(len(chunks))]
    metas = [
        {
            "title": url,
            "topic": topic,
            "part": i,
            "source": "link",
            "type": "qa",
        }
        for i in range(len(chunks))
    ]
    col.add(ids=ids, metadatas=metas, documents=chunks)

    await update.message.reply_text(
        f"Learned from link ✅ ({len(chunks)} parts) in topic: {topic} (Q&A)"
    )


# ---------- TEACH LINK EVAL (EVALUATION sources) ----------
async def teachlink_eval(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachlink_eval")

    msg_text = (update.message.text or "").strip()
    parts = msg_text.split(maxsplit=1)
    if len(parts) < 2:
        await update.message.reply_text("Use: /teachlink_eval <url>")
        return

    url = parts[1].strip()
    await update.message.reply_text(
        "Fetching content from link and learning it as evaluation / rubric material…"
    )

    try:
        downloaded = trafilatura.fetch_url(url)
        text = trafilatura.extract(downloaded)
    except Exception as e:
        await update.message.reply_text(f"Error while fetching the URL: {e}")
        return

    if not text:
        await update.message.reply_text("Couldn't extract readable text from that link.")
        return

    chunks = _chunk(text)
    if not chunks:
        await update.message.reply_text("The page did not contain enough text to learn from.")
        return

    col = get_collection(chat_id, topic)

    existing = col.get(where={"title": url})
    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"This link is already learned in topic: {topic}.\n"
            "Use /unlearn <url> to remove it first."
        )
        return

    base = f"{chat_id}_{topic}_{col.count()+1}"
    ids = [f"{base}_eval_{i}" for i in range(len(chunks))]
    metas = [
        {
            "title": url,
            "topic": topic,
            "part": i,
            "source": "link",
            "type": "evaluation",
        }
        for i in range(len(chunks))
    ]
    col.add(ids=ids, metadatas=metas, documents=chunks)

    await update.message.reply_text(
        f"Learned evaluation material from link ✅ ({len(chunks)} parts) in topic: {topic}"
    )


# ---------- TEACH IMAGE (Q&A) ----------
async def teachimage(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachimage")

    photos = update.message.photo or []
    if not photos:
        await update.message.reply_text(
            "Please send a clear image with caption:\n/teachimage <title>"
        )
        return

    caption = (update.message.caption or "").strip()
    title = None
    if caption.startswith("/teachimage"):
        parts = caption.split(maxsplit=1)
        if len(parts) > 1:
            title = parts[1].strip()

    largest = photos[-1]
    tgfile = await largest.get_file()
    if not title:
        title = f"image_{tgfile.file_unique_id}"

    await update.message.reply_text(
        f"Reading your image for topic '{topic}' and extracting text to learn from it…"
    )

    img_bytes = await tgfile.download_as_bytearray()

    try:
        extracted = extract_text_from_image_bytes(img_bytes)
    except Exception as e:
        await update.message.reply_text(f"Could not extract text from image: {e}")
        return

    if not extracted.strip():
        await update.message.reply_text("I couldn't read any text from that image.")
        return

    parts = _chunk(extracted)
    if not parts:
        await update.message.reply_text("The extracted text was too short to learn from.")
        return

    col = get_collection(chat_id, topic)

    existing = col.get(where={"title": title})
    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' is already learned in topic: {topic}.\n"
            "Use /unlearn <title> to remove it first if needed."
        )
        return

    base = f"{chat_id}_{topic}_{col.count()+1}"
    ids = [f"{base}_img_{i}" for i in range(len(parts))]
    metas = [
        {
            "title": title,
            "topic": topic,
            "part": i,
            "source": "image",
            "type": "qa",
        }
        for i in range(len(parts))
    ]
    col.add(ids=ids, metadatas=metas, documents=parts)

    await update.message.reply_text(
        f"Learned from image '{title}' ✅ ({len(parts)} parts) in topic: {topic} (Q&A)"
    )


# ---------- SOURCES ----------
async def sources_all(update: Update, context: ContextTypes.DEFAULT_TYPE):
    client = chroma
    collections = client.list_collections()

    if not collections:
        await update.message.reply_text("No sources stored yet.")
        return

    source_stats = {}  # title -> {chunks, bytes}
    total_bytes = 0
    total_chunks = 0

    for col_info in collections:
        col = client.get_collection(col_info.name)
        data = col.get(include=["documents", "metadatas"])

        docs = data.get("documents", [])
        metas = data.get("metadatas", [])

        for doc, meta in zip(docs, metas):
            title = meta.get("title", "Untitled")
            size_bytes = len(doc.encode("utf-8"))

            if title not in source_stats:
                source_stats[title] = {"chunks": 0, "bytes": 0}

            source_stats[title]["chunks"] += 1
            source_stats[title]["bytes"] += size_bytes

            total_chunks += 1
            total_bytes += size_bytes

    lines = []
    for title, stats in sorted(
        source_stats.items(), key=lambda x: x[1]["bytes"], reverse=True
    ):
        mb = stats["bytes"] / (1024 * 1024)
        lines.append(f"• {title}: {stats['chunks']} chunks, {mb:.2f} MB")

    total_mb = total_bytes / (1024 * 1024)

    msg = (
        "📚 ALL BOT SOURCES (ALL TOPICS)\n\n"
        f"Total chunks: {total_chunks}\n"
        f"Total text size: {total_mb:.2f} MB\n\n"
        + "\n".join(lines)
    )

    await send_long(update, msg)

async def sources(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    topic = get_current_topic(context)
    col = get_collection(chat_id, topic)
    data = col.get(include=["metadatas"])
    metas = data.get("metadatas") or []
    counts = Counter(m.get("title", "Untitled") for m in metas)

    if not counts:
        msg = f"📌 Active topic: {topic}\nNo sources yet."
        await update.message.reply_text(msg)
        return

    type_map = {}
    for m in metas:
        title = m.get("title", "Untitled")
        t = m.get("type", "qa")
        type_map.setdefault(title, t)

    lines = []
    for title, n in counts.items():
        t = type_map.get(title, "qa")
        label = "Q&A" if t == "qa" else ("Eval" if t == "evaluation" else t)
        if n > 1:
            lines.append(f"• {title} ({n} parts, {label})")
        else:
            lines.append(f"• {title} ({label})")

    msg = f"📌 Active topic: {topic}\n" + "\n".join(lines)
    await update.message.reply_text(msg)


# ---------- UNLEARN ----------
async def unlearn(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="unlearn")

    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        await update.message.reply_text("Usage: /unlearn <exact title shown in /sources>")
        return

    title = parts[1].strip()
    col = get_collection(chat_id, topic)
    to_delete = col.get(where={"title": title})
    removed = len((to_delete or {}).get("ids") or [])

    if removed == 0:
        await update.message.reply_text(
            f"No source titled '{title}' found in topic: {topic}."
        )
        return

    col.delete(where={"title": title})
    await update.message.reply_text(
        f"Removed '{title}' ✅ ({removed} parts) from topic: {topic}"
    )


# ---------- CLEAR ----------
async def clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="clear")

    chroma.delete_collection(f"chat_{chat_id}_{topic}")
    await update.message.reply_text(f"Forgot everything 🧹 (topic: {topic})")


# ---------- EVALUATION HELPERS ----------
def _eval_context_from_collection(col, extra_query: str = ""):
    """Get combined evaluation + QA snippets for evaluation prompts."""
    eval_docs = []
    qa_docs = []
    try:
        res_eval = col.query(
            query_texts=["evaluation criteria", "guidelines", "rubric", extra_query],
            where={"type": "evaluation"},
            n_results=6,
        )
        eval_docs = res_eval.get("documents", [[]])[0]
    except Exception:
        eval_docs = []

    try:
        res_qa = col.query(
            query_texts=["tips", "examples", "advice", extra_query],
            where={"type": "qa"},
            n_results=6,
        )
        qa_docs = res_qa.get("documents", [[]])[0]
    except Exception:
        qa_docs = []

    docs = (eval_docs or []) + (qa_docs or [])
    return "\n\n---\n\n".join(docs) if docs else ""


async def evaluate_file_for_topic(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)

    record_event(user.id, topic, kind="eval")

    allowed = set(EVAL_TOPICS) | {"ielts_writing"}
    if topic not in allowed:
        await update.message.reply_text(
            "To evaluate an essay, recommendation, EC description, portfolio, or IELTS Writing, "
            "choose the correct menu first, then send the file."
        )
        return

    doc = update.message.document
    if not doc:
        await update.message.reply_text("Please attach a PDF or DOCX file.")
        return

    pretty_topic = {
        "essays_personal": "Personal Statement essay",
        "essays_supplemental": "Supplemental essay",
        "recommendations": "Recommendation letter",
        "portfolio": "Portfolio",
        "extracurriculars": "Extracurricular activities description",
        "ielts_writing": "IELTS Writing answer",
    }.get(topic, "document")

    await update.message.reply_text(f"Reading your {pretty_topic} file…")

    tgfile = await doc.get_file()
    file_bytes = await tgfile.download_as_bytearray()
    name = (doc.file_name or "document").lower()

    try:
        if name.endswith(".pdf"):
            text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text(
                "Only PDF or DOCX are supported for evaluation."
            )
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    parts = _chunk(text, max_chars=1500)
    if not parts:
        await update.message.reply_text(
            "I couldn't read enough text from that file to evaluate."
        )
        return

    student_text = "\n\n---\n\n".join(parts[:5])

    await update.message.reply_text(
        f"Analyzing your {pretty_topic} against my guidelines…"
    )
    col = get_collection(chat_id, topic)

    context_block = _eval_context_from_collection(col, extra_query=pretty_topic)

    if topic in {"essays_personal", "essays_supplemental"}:
        sys_role = (
            "You are an expert college admissions essay coach. "
            "Evaluate the student's essay. Focus on clarity, structure, voice, authenticity, and impact. "
            "Give specific, actionable feedback and suggestions."
        )
    elif topic == "recommendations":
        sys_role = (
            "You are an expert on college recommendation letters. "
            "Evaluate the letter in terms of specificity, credibility, depth of insight, and support for the student. "
            "Give constructive feedback and suggestions for improvement."
        )
    elif topic == "extracurriculars":
        sys_role = (
            "You are an expert on extracurricular strategy for college applications. "
            "Evaluate how well the activities are presented in terms of impact, leadership, continuity, and uniqueness. "
            "Give specific, practical suggestions to make the activities stand out."
        )
    elif topic == "ielts_writing":
        sys_role = (
            "You are an experienced IELTS Writing examiner. "
            "Evaluate the student's writing according to IELTS band descriptors. "
            "Comment on Task Response, Coherence and Cohesion, Lexical Resource, and Grammatical Range and Accuracy. "
            "Give an approximate band score and clear, actionable feedback."
        )
    else:  # portfolio
        sys_role = (
            "You are an expert college portfolio reviewer. "
            "Evaluate the portfolio in terms of coherence, originality, technical quality, and fit for selective colleges. "
            "Give specific, constructive feedback, not generic advice."
        )

    messages = [
        {
            "role": "system",
            "content": (
                sys_role
                + " Use the guidelines and examples in the context when relevant. "
                "If guidelines are missing or incomplete, use general best practices."
            ),
        },
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {
            "role": "user",
            "content": f"Here is the student's {pretty_topic}. Please evaluate it:\n\n{student_text}",
        },
    ]

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.3,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error during evaluation: {e}"

    await send_long(update, a)


async def evaluate_ielts_writing_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = "ielts_writing"

    record_event(user.id, topic, kind="eval")

    photos = update.message.photo or []
    if not photos:
        await update.message.reply_text("Please send a clear photo of your IELTS Writing answer.")
        return

    largest = photos[-1]
    tgfile = await largest.get_file()
    await update.message.reply_text("Reading your IELTS Writing answer from the image…")
    img_bytes = await tgfile.download_as_bytearray()

    try:
        extracted = extract_text_from_image_bytes(img_bytes)
    except Exception as e:
        await update.message.reply_text(f"Could not extract text from image: {e}")
        return

    if not extracted.strip():
        await update.message.reply_text(
            "I couldn't read enough text from that image. Please try a clearer photo."
        )
        return

    parts = _chunk(extracted, max_chars=1500)
    student_text = "\n\n---\n\n".join(parts[:5])

    await update.message.reply_text("Analyzing your IELTS Writing answer…")

    col = get_collection(chat_id, topic)
    context_block = _eval_context_from_collection(col, extra_query="IELTS Writing")

    sys_role = (
        "You are an experienced IELTS Writing examiner. "
        "Evaluate the student's writing according to IELTS Academic/General Writing band descriptors. "
        "Comment on Task Response, Coherence and Cohesion, Lexical Resource, and Grammatical Range and Accuracy. "
        "Give an approximate band score (like 6.0, 6.5, 7.0) and then clear, actionable feedback."
    )

    messages = [
        {
            "role": "system",
            "content": (
                sys_role
                + " Use any rubric/context provided when available. "
                "If the text seems incomplete or too short, mention that in your feedback."
            ),
        },
        {"role": "system", "content": f"IELTS writing rubrics and notes (may be empty):\n{context_block}"},
        {
            "role": "user",
            "content": f"Here is the student's IELTS Writing answer (from an image):\n\n{student_text}",
        },
    ]

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.3,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error during IELTS Writing evaluation: {e}"

    await send_long(update, a)


async def evaluate_text_for_topic(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Evaluate pasted text (essay, ECs, rec, portfolio, IELTS writing)."""
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="eval")

    allowed = set(EVAL_TOPICS) | {"ielts_writing"}
    if topic not in allowed:
        await update.message.reply_text(
            "Text evaluation is only available for Personal Statement, Supplemental Essays, "
            "Extracurriculars, Recommendation Letters, Portfolio, and IELTS Writing.\n\n"
            "Choose the correct topic first, then tap the Evaluation button again."
        )
        return

    raw = (update.message.text or "").strip()
    for marker in [
        BTN_PS_EVAL,
        BTN_SUPP_EVAL,
        BTN_EC_EVAL,
        BTN_REC_EVAL,
        BTN_IW_EVAL,
        BTN_PORT_EVAL,
    ]:
        raw = raw.replace(marker, "").strip()

    if len(raw) < 100:
        await update.message.reply_text(
            "The text is too short to evaluate. Please paste the full essay/letter/description."
        )
        return

    pretty_topic = {
        "essays_personal": "Personal Statement essay",
        "essays_supplemental": "Supplemental essay",
        "recommendations": "Recommendation letter",
        "portfolio": "Portfolio description",
        "extracurriculars": "Extracurricular activities description",
        "ielts_writing": "IELTS Writing answer",
    }.get(topic, "document")

    parts = _chunk(raw, max_chars=1500)
    student_text = "\n\n---\n\n".join(parts[:5])

    await update.message.reply_text(f"Evaluating your {pretty_topic}…")

    col = get_collection(chat_id, topic)
    context_block = _eval_context_from_collection(col, extra_query=pretty_topic)

    if topic in {"essays_personal", "essays_supplemental"}:
        sys_role = (
            "You are an expert college admissions essay coach. "
            "Evaluate the student's essay. Focus on clarity, structure, voice, authenticity, and impact. "
            "Give specific, actionable feedback and suggestions."
        )
    elif topic == "recommendations":
        sys_role = (
            "You are an expert on college recommendation letters. "
            "Evaluate the letter in terms of specificity, credibility, depth of insight, and support for the student. "
            "Give constructive feedback and suggestions for improvement."
        )
    elif topic == "extracurriculars":
        sys_role = (
            "You are an expert on extracurricular strategy for college applications. "
            "Evaluate how well the activities are presented in terms of impact, leadership, continuity, and uniqueness. "
            "Give specific, practical suggestions to make the activities stand out."
        )
    elif topic == "ielts_writing":
        sys_role = (
            "You are an experienced IELTS Writing examiner. "
            "Evaluate the student's writing according to IELTS band descriptors. "
            "Comment on Task Response, Coherence and Cohesion, Lexical Resource, and Grammatical Range and Accuracy. "
            "Give an approximate band score and clear, actionable feedback."
        )
    else:  # portfolio
        sys_role = (
            "You are an expert college portfolio reviewer. "
            "Evaluate the portfolio description in terms of coherence, originality, technical quality, and fit for selective colleges. "
            "Give specific, constructive feedback, not generic advice."
        )

    messages = [
        {
            "role": "system",
            "content": (
                sys_role
                + " Use the guidelines and examples in the context when available. "
                "If guidelines are missing or incomplete, use general best practices."
            ),
        },
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {
            "role": "user",
            "content": f"Here is the student's {pretty_topic}. Please evaluate it:\n\n{student_text}",
        },
    ]

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.3,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error during text evaluation: {e}"

    await send_long(update, a)


# ---------- DOCUMENT & PHOTO ROUTERS ----------
async def document_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    caption = (update.message.caption or "").strip()
    topic = get_current_topic(context)

    if caption.startswith("/teachfile_eval"):
        await teachfile_eval(update, context)
        return

    if caption.startswith("/teachfile"):
        await teachfile(update, context)
        return

    if topic in (set(EVAL_TOPICS) | {"ielts_writing"}):
        await evaluate_file_for_topic(update, context)
        return

    await update.message.reply_text(
        "If you want me to LEARN from this file, send it again and write /teachfile "
        "or /teachfile_eval in the caption.\n\n"
        "If this is an essay, recommendation, EC description, portfolio, or IELTS Writing for feedback, "
        "choose the correct topic and tap its Evaluation button."
    )


async def photo_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    caption = (update.message.caption or "").strip()
    topic = get_current_topic(context)
    user = update.effective_user

    record_event(user.id, topic, kind="photo")

    if caption.startswith("/teachimage"):
        await teachimage(update, context)
        return

    if topic == "ielts_writing":
        await evaluate_ielts_writing_image(update, context)
        return

    await update.message.reply_text(
        "I can learn from images too 🤖🖼\n\n"
        "If you want me to *learn* from this image (e.g., essay screenshot, rubric), "
        "send it again with caption:\n\n"
        "/teachimage <title>\n\n"
        "For IELTS Writing evaluation from an image, switch to IELTS Writing first."
    )


# ---------- STATS ----------
async def stats_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    stats = load_stats()
    total_users = len(stats.get("users", []))
    total_msgs = stats.get("messages_total", 0)
    topic_counts = stats.get("topic_counts", {})
    eval_counts = stats.get("eval_counts", {})

    if topic_counts:
        top_topics = sorted(
            topic_counts.items(), key=lambda kv: kv[1], reverse=True
        )[:5]
        topics_str = "\n".join(f"- {t}: {c} msgs" for t, c in top_topics)
    else:
        topics_str = "No topic data yet."

    total_evals = sum(eval_counts.values()) if eval_counts else 0

    msg = (
        f"📊 Bot analytics\n"
        f"- Unique users: {total_users}\n"
        f"- Total interactions (events): {total_msgs}\n"
        f"- Total evaluations: {total_evals}\n\n"
        f"Top topics:\n{topics_str}"
    )
    await update.message.reply_text(msg)


# ---------- NEW FEATURE HELPERS ----------
def set_pending_feature(context: ContextTypes.DEFAULT_TYPE, feature: str | None):
    context.user_data["pending_feature"] = feature


async def run_brainstorm(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    topic = get_current_topic(context)
    user = update.effective_user
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="brainstorm")

    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)

    # Pull topic-specific context (Q&A + Eval) to guide brainstorming
    col = get_collection(chat_id, topic)
    try:
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception:
        docs = []
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        f"You are an admissions mentor helping a student brainstorm ideas for {nice_topic}.\n"
        "- Give 3–5 short bullet ideas or angles.\n"
        "- Each bullet should be 1–2 concise sentences.\n"
        "- Focus on realistic, personal, and application-relevant ideas.\n"
        "- Keep the tone friendly, specific, and not generic."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {
                "role": "system",
                "content": f"Program-specific notes and examples (may be empty):\n{context_block}",
            }
        )
    messages.append(
        {
            "role": "user",
            "content": (
                "Here is the student's situation and what they are thinking about:\n\n"
                + description
            ),
        }
    )

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.5,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error during brainstorming: {e}"

    await send_long(update, a)


async def brainstorm_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "brainstorm")
        await update.message.reply_text(
            "🧠 Brainstorm mode ON.\n\n"
            "Briefly tell me about yourself, your target major, and what you want to write about."
        )
        return

    description = parts[1].strip()
    await run_brainstorm(update, context, description)


async def run_rewrite(update: Update, context: ContextTypes.DEFAULT_TYPE, text_to_fix: str):
    await show_typing(update, context)
    topic = get_current_topic(context)
    user = update.effective_user
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="rewrite")

    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)

    # Pull topic-specific context
    col = get_collection(chat_id, topic)
    try:
        res = col.query(query_texts=[text_to_fix], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception:
        docs = []
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        f"You are an admissions writing coach helping improve a student's {nice_topic} text.\n"
        "- Rewrite the text to be clearer, more natural, and slightly more mature.\n"
        "- Preserve the student's original meaning and main ideas.\n"
        "- Keep roughly similar length (do not double it).\n"
        "- Use a human, conversational but polished tone.\n"
        "- Reply ONLY with the revised text, no explanations."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {
                "role": "system",
                "content": f"Program-specific notes and examples (may be empty):\n{context_block}",
            }
        )
    messages.append({"role": "user", "content": text_to_fix})

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.4,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error during rewrite: {e}"

    await send_long(update, a)


async def rewrite_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "rewrite")
        await update.message.reply_text(
            "✍️ Rewrite mode ON.\n\n"
            "Send me the paragraph or essay you want me to improve."
        )
        return

    to_fix = parts[1].strip()
    await run_rewrite(update, context, to_fix)


async def run_plan(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="plan")

    col = get_collection(chat_id, topic)
    try:
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception:
        docs = []
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        "You are an admissions strategy mentor.\n"
        "- Based on the student's situation, create a concise application plan.\n"
        "- Organize it into short bullet points under 3 headings: "
        "Academics & Testing, Essays & Recs, Activities & Extras.\n"
        "- Keep total response around 120–200 words.\n"
        "- Focus on practical next steps, not theory."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {
                "role": "system",
                "content": f"Program-specific planning notes (may be empty):\n{context_block}",
            }
        )
    messages.append(
        {
            "role": "user",
            "content": (
                "Here is my situation (grade, target countries, majors, timeline, current profile):\n\n"
                + description
            ),
        }
    )

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.5,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error while building plan: {e}"

    await send_long(update, a)


async def plan_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["topic"] = "application_plan"
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "plan")
        await update.message.reply_text(
            "📅 Application Plan mode ON.\n\n"
            "Tell me your grade, target countries, intended major, test scores (if any), and your rough deadlines."
        )
        return

    description = parts[1].strip()
    await run_plan(update, context, description)


async def run_recpacket(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="recpacket")

    sys = (
        "You are creating a recommendation letter 'brag sheet' for a teacher.\n"
        "- Output in 3 short sections:\n"
        "  1) 3–5 sentence summary the student can give the teacher.\n"
        "  2) Bullet list of key achievements/impacts.\n"
        "  3) Bullet list of personal qualities and 2–3 specific story ideas.\n"
        "- Keep it concise and realistic for competitive admissions."
    )

    messages = [
        {"role": "system", "content": sys},
        {
            "role": "user",
            "content": (
                "Here is information about me, my relationship with the recommender, and what I hope they mention:\n\n"
                + description
            ),
        },
    ]

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.5,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error while generating rec letter packet: {e}"

    await send_long(update, a)


async def recpacket_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "recpacket")
        await update.message.reply_text(
            "✉️ Rec Letter Packet mode ON.\n\n"
            "Tell me which teacher will write your rec, what classes you took with them, "
            "your main achievements, and what you want them to highlight."
        )
        return

    description = parts[1].strip()
    await run_recpacket(update, context, description)


async def run_schoolfinder(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="schoolfinder")

    col = get_collection(chat_id, topic)
    try:
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception:
        docs = []
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        "You are a university match advisor.\n"
        "- Based on the student's stats and preferences, suggest Reach, Match, and Safety school *types* "
        "and a few example universities.\n"
        "- For each category, give 2–4 example schools and 1–2 bullets about why they fit.\n"
        "- Keep total response concise (around 150–220 words).\n"
        "- Make it clear this is an approximate starting point and they must research details themselves."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {
                "role": "system",
                "content": f"Program-specific school lists/notes (may be empty):\n{context_block}",
            }
        )
    messages.append(
        {
            "role": "user",
            "content": (
                "Here are my stats and preferences (GPA, tests, budget, countries, major, special interests):\n\n"
                + description
            ),
        },
    )

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.6,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error while suggesting schools: {e}"

    await send_long(update, a)


async def schoolfinder_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["topic"] = "school_finder"
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "schoolfinder")
        await update.message.reply_text(
            "🏫 School Finder mode ON.\n\n"
            "Send me your GPA (or approximate), test scores (if any), budget, target countries, "
            "intended major, and any constraints (e.g. need scholarship)."
        )
        return

    description = parts[1].strip()
    await run_schoolfinder(update, context, description)


async def run_portfolioideas(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="portfolioideas")

    sys = (
        "You are a portfolio mentor for university applications.\n"
        "- Based on the student's field (e.g. CS, design, art, film, business) and interests, "
        "suggest 3–6 concrete project ideas.\n"
        "- Each idea should be 1–2 sentences, focused on impact and what it shows about the student.\n"
        "- Make ideas realistic for a high school student, but impressive."
    )

    messages = [
        {"role": "system", "content": sys},
        {
            "role": "user",
            "content": (
                "Here is my background, target major/program, and what I might include in a portfolio:\n\n"
                + description
            ),
        },
    ]

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.6,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error while generating portfolio ideas: {e}"

    await send_long(update, a)


async def portfolioideas_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "portfolioideas")
        await update.message.reply_text(
            "🖼️ Portfolio Ideas mode ON.\n\n"
            "Tell me your field (CS, design, art, film, etc.), your skills, and the kind of programs you are targeting."
        )
        return

    description = parts[1].strip()
    await run_portfolioideas(update, context, description)


# ---------- MAIN ANSWER ----------
async def answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    user = update.effective_user
    q = (update.message.text or "").strip()

    topic_before = get_current_topic(context)
    record_event(user.id, topic_before, kind="message")

    # --- Handle pending feature input first ---
    pending = context.user_data.get("pending_feature")
    if pending:
        context.user_data["pending_feature"] = None
        if pending == "brainstorm":
            await run_brainstorm(update, context, q)
            return
        if pending == "rewrite":
            await run_rewrite(update, context, q)
            return
        if pending == "plan":
            await run_plan(update, context, q)
            return
        if pending == "recpacket":
            await run_recpacket(update, context, q)
            return
        if pending == "schoolfinder":
            await run_schoolfinder(update, context, q)
            return
        if pending == "portfolioideas":
            await run_portfolioideas(update, context, q)
            return

    # ---- MAIN MENUS ----
    if q == BTN_ESSAY:
        await send_with_image(
            update,
            "Essays selected.\nChoose Personal Statement or Supplemental Essays.",
            reply_markup=essay_main_keyboard(),
            image_key="essays_main",
        )
        return

    if q == BTN_SAT:
        await send_with_image(
            update,
            "SAT selected. Choose a section:",
            reply_markup=sat_menu_keyboard(),
            image_key="sat_main",
        )
        return

    if q == BTN_IELTS:
        await send_with_image(
            update,
            "IELTS selected. Choose a section.",
            reply_markup=ielts_main_keyboard(),
            image_key="ielts_main",
        )
        return

    if q == BTN_PLAN_MAIN:
        context.user_data["topic"] = "application_plan"
        context.user_data["eval_active"] = False
        set_pending_feature(context, "plan")
        await send_with_image(
            update,
            "📅 Application Plan mode ON.\n\n"
            "Tell me your grade, target countries, intended major, test scores (if any), and your rough deadlines.",
            reply_markup=main_menu_with_back(),
            image_key="plan_main",
        )
        return

    if q == BTN_SF_MAIN:
        context.user_data["topic"] = "school_finder"
        context.user_data["eval_active"] = False
        set_pending_feature(context, "schoolfinder")
        await send_with_image(
            update,
            "🏫 School Finder mode ON.\n\n"
            "Send me your GPA (or approximate), test scores (if any), budget, target countries, "
            "intended major, and any constraints (e.g. need scholarship).",
            reply_markup=main_menu_with_back(),
            image_key="schoolfinder_main",
        )
        return

    if q == BTN_BACK:
        context.user_data["eval_active"] = False
        context.user_data["pending_feature"] = None
        context.user_data["topic"] = DEFAULT_TOPIC
        await update.message.reply_text(
            "Back to main menu. You're now in general mode – you can just ask questions or choose a section again.",
            reply_markup=main_menu_keyboard(),
        )
        return

    # ---- TOPIC BUTTONS ----
    if q == BTN_ESSAY_PS:
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        await send_with_image(
            update,
            "Great, let's work on your Personal Statement.\n"
            "Ask about structure, voice, and storytelling.",
            reply_markup=essay_ps_keyboard(),
            image_key="essays_personal",
        )
        await update.message.reply_text(
            "If you want detailed feedback, tap '✅ Personal Statement Evaluation' "
            "and then upload your essay as a PDF/DOCX or paste the text.\n\n"
            "You can also use '🧠 Brainstorm ideas' to find topics, or '✍️ Rewrite my text' to polish a draft."
        )
        return

    if q == BTN_ESSAY_SUPP:
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        await send_with_image(
            update,
            "Great, let's work on your Supplemental Essays.\n"
            "Ask about 'Why us', community essays, and short prompts.",
            reply_markup=essay_supp_keyboard(),
            image_key="essays_supplemental",
        )
        await update.message.reply_text(
            "If you want detailed feedback, tap '✅ Supplemental Essay Evaluation' "
            "and then upload your essay as a PDF/DOCX or paste the text.\n\n"
            "You can also use '🧠 Brainstorm ideas' to explore angles, or '✍️ Rewrite my text' to refine your answer."
        )
        return

    if q == BTN_EC:
        topic = "extracurriculars"
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        await send_with_image(
            update,
            "Great, let's talk about your Extracurricular activities.\n"
            "Ask how to present impact, leadership, and long-term involvement.",
            reply_markup=ec_keyboard(),
            image_key="extracurriculars",
        )
        await update.message.reply_text(
            "For feedback on your activities list, click '✅ Extracurricular Evaluation' "
            "and then upload your EC descriptions as a PDF/DOCX or paste them."
        )
        return

    if q == BTN_REC:
        topic = "recommendations"
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        await send_with_image(
            update,
            "Great, let's work on Recommendation Letters.\n"
            "Ask how to request them, what to share with recommenders, and what makes a strong letter.",
            reply_markup=rec_keyboard(),
            image_key="recommendations",
        )
        await update.message.reply_text(
            "You can:\n"
            "- Tap '✅ Rec Letter Evaluation' to get feedback on a draft letter.\n"
            "- Tap '📄 Rec Letter Packet' to build a brag sheet for your teacher."
        )
        return

    if q == BTN_PORT:
        topic = "portfolio"
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        await send_with_image(
            update,
            "You're now in Portfolio Check.\nAsk about structure and how to present your work.",
            reply_markup=portfolio_keyboard(),
            image_key="portfolio",
        )
        await update.message.reply_text(
            "You can upload your portfolio as a PDF/DOCX for detailed feedback (✅ Portfolio Evaluation), "
            "or tap '💡 Portfolio Ideas' to get new project ideas."
        )
        return

    if q in {BTN_SAT_MATH, BTN_SAT_ENGLISH}:
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        nice_name = "SAT Math" if topic == "sat_math" else "SAT English"
        prompts = {
            "sat_math": "Ask about SAT Math topics, strategies, and practice plans.",
            "sat_english": "Ask about SAT Reading/Writing strategies and question types.",
        }
        image_key = "sat_math" if topic == "sat_math" else "sat_english"

        await send_with_image(
            update,
            f"You're now in {nice_name}.\n{prompts.get(topic, '')}",
            reply_markup=sat_menu_keyboard(),
            image_key=image_key,
        )
        return

    if q in {BTN_IELTS_READING, BTN_IELTS_LISTENING, BTN_IELTS_SPEAKING}:
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        pretty = {
            "ielts_reading": "IELTS Reading",
            "ielts_listening": "IELTS Listening",
            "ielts_speaking": "IELTS Speaking",
        }
        prompts = {
            "ielts_reading": "Ask about IELTS Reading strategies and question types.",
            "ielts_listening": "Ask about IELTS Listening tips and common traps.",
            "ielts_speaking": "Ask about IELTS Speaking parts, fluency, and idea generation.",
        }

        await send_with_image(
            update,
            f"You're now in {pretty.get(topic, topic)}.\n{prompts.get(topic, '')}",
            reply_markup=ielts_main_keyboard(),
            image_key=topic,
        )
        return

    if q == BTN_IELTS_WRITING:
        topic = "ielts_writing"
        context.user_data["topic"] = topic
        context.user_data["eval_active"] = False

        await send_with_image(
            update,
            "You're now in IELTS Writing.\nAsk about Task 1 & 2, band 7+ strategies, or send your answer for feedback.",
            reply_markup=ielts_writing_keyboard(),
            image_key="ielts_writing",
        )
        await update.message.reply_text(
            "For a full band-style evaluation, click '✅ Writing Evaluation' "
            "and then send your answer as text, PDF/DOCX, or a clear photo."
        )
        return

    # ---- FEATURE BUTTONS INSIDE MENUS ----
    if q == BTN_BRAINSTORM:
        set_pending_feature(context, "brainstorm")
        await update.message.reply_text(
            "🧠 Brainstorm mode ON.\n\n"
            "Briefly tell me about yourself, your target major, and what you want to write about."
        )
        return

    if q == BTN_REWRITE:
        set_pending_feature(context, "rewrite")
        await update.message.reply_text(
            "✍️ Rewrite mode ON.\n\n"
            "Send me the paragraph or essay you want me to improve."
        )
        return

    if q == BTN_REC_PACKET:
        set_pending_feature(context, "recpacket")
        await update.message.reply_text(
            "✉️ Rec Letter Packet mode ON.\n\n"
            "Tell me which teacher will write your rec, what classes you took with them, "
            "your main achievements, and what you want them to highlight."
        )
        return

    if q == BTN_PORTFOLIO_IDEAS:
        set_pending_feature(context, "portfolioideas")
        await update.message.reply_text(
            "💡 Portfolio Ideas mode ON.\n\n"
            "Tell me your field (CS, design, art, film, etc.), your skills, and the kind of programs you are targeting."
        )
        return

    # ---- EVALUATION BUTTONS ----
    if q == BTN_PS_EVAL:
        context.user_data["topic"] = "essays_personal"
        context.user_data["eval_active"] = True
        await update.message.reply_text(
            "Personal Statement Evaluation mode ON ✅\n\n"
            "Now send your Personal Statement as text in the chat, "
            "or upload it as a PDF/DOCX file.\n"
            "I'll evaluate it using your program's guidelines and give structured feedback."
        )
        return

    if q == BTN_SUPP_EVAL:
        context.user_data["topic"] = "essays_supplemental"
        context.user_data["eval_active"] = True
        await update.message.reply_text(
            "Supplemental Essay Evaluation mode ON ✅\n\n"
            "Now send your supplemental essay as text, "
            "or upload it as a PDF/DOCX file.\n"
            "I'll evaluate it using your program's guidelines and give structured feedback."
        )
        return

    if q == BTN_EC_EVAL:
        context.user_data["topic"] = "extracurriculars"
        context.user_data["eval_active"] = True
        await update.message.reply_text(
            "Extracurricular Evaluation mode ON ✅\n\n"
            "Now send your EC descriptions as text, "
            "or upload them as a PDF/DOCX file.\n"
            "I'll review them for impact, leadership, and clarity."
        )
        return

    if q == BTN_REC_EVAL:
        context.user_data["topic"] = "recommendations"
        context.user_data["eval_active"] = True
        await update.message.reply_text(
            "Rec Letter Evaluation mode ON ✅\n\n"
            "Now send your draft recommendation letter as text, "
            "or upload it as a PDF/DOCX file.\n"
            "I'll review it for specificity, credibility, and strength."
        )
        return

    if q == BTN_IW_EVAL:
        context.user_data["topic"] = "ielts_writing"
        context.user_data["eval_active"] = True
        await update.message.reply_text(
            "IELTS Writing Evaluation mode ON ✅\n\n"
            "Now send your IELTS Writing answer as text, as a PDF/DOCX, or as a clear photo.\n"
            "I'll estimate your band and give detailed feedback."
        )
        return

    if q == BTN_PORT_EVAL:
        context.user_data["topic"] = "portfolio"
        context.user_data["eval_active"] = True
        await update.message.reply_text(
            "Portfolio Evaluation mode ON ✅\n\n"
            "Now send your portfolio description as text, or upload it as a PDF/DOCX.\n"
            "I'll review it for coherence, originality, and strength for selective colleges."
        )
        return

    # ---- EVALUATION FROM TEXT ----
    if context.user_data.get("eval_active", False):
        context.user_data["eval_active"] = False
        await evaluate_text_for_topic(update, context)
        return

    # ---- NORMAL Q&A WITH RAG ----
    await show_typing(update, context)
    await update.message.reply_text("Got it! Thinking about your question…")

    topic = get_current_topic(context)
    col = get_collection(chat_id, topic)

    # Try to use QA docs only; if none, fall back to all
    try:
        results = col.query(
            query_texts=[q],
            where={"type": "qa"},
            n_results=4,
        )
        docs = results.get("documents", [[]])[0]
    except Exception:
        docs = []

    if not docs:
        try:
            results = col.query(query_texts=[q], n_results=4)
            docs = results.get("documents", [[]])[0]
        except Exception:
            docs = []

    context_block = "\n\n---\n\n".join(docs)

    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)
    sys = (
        f"You are UniVenture, a focused university admissions mentor.\n"
        f"Current topic: {nice_topic}.\n"
        "- Give short, clear, high-value answers (3–7 sentences max).\n"
        "- Speak in a natural, human tone, like a friendly but direct older student mentor.\n"
        "- Prioritize practical, actionable advice over theory.\n"
        "- Use the provided context (if any) as trusted program material and do not contradict it."
    )

    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": f"Context (may be empty):\n{context_block}"},
        {"role": "user", "content": q},
    ]

    try:
        r = oa.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.4,
        )
        a = r.choices[0].message.content
    except Exception as e:
        a = f"Error: {e}"

    await send_long(update, a)


# -------- App wiring --------
app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()
app.add_handler(CommandHandler("start", start))
app.add_handler(CommandHandler("teach", teach))
app.add_handler(CommandHandler("teachrubric", teachrubric))
app.add_handler(CommandHandler("teachfile", teachfile))
app.add_handler(CommandHandler("teachfile_eval", teachfile_eval))
app.add_handler(CommandHandler("teachlink", teachlink))
app.add_handler(CommandHandler("teachlink_eval", teachlink_eval))
app.add_handler(
    CommandHandler(
        "teachimage",
        lambda u, c: u.message.reply_text(
            "To teach me from an image, send a photo with caption:\n\n/teachimage <title>"
        ),
    )
)
app.add_handler(CommandHandler("sources", sources))
app.add_handler(CommandHandler("unlearn", unlearn))
app.add_handler(CommandHandler("clear", clear))
app.add_handler(CommandHandler("stats", stats_cmd))

# New feature commands (still available via slash, plus via buttons)
app.add_handler(CommandHandler("brainstorm", brainstorm_cmd))
app.add_handler(CommandHandler("rewrite", rewrite_cmd))
app.add_handler(CommandHandler("plan", plan_cmd))
app.add_handler(CommandHandler("recpacket", recpacket_cmd))
app.add_handler(CommandHandler("schoolfinder", schoolfinder_cmd))
app.add_handler(CommandHandler("portfolioideas", portfolioideas_cmd))

# File & photo routers
app.add_handler(MessageHandler(filters.Document.ALL, document_router))
app.add_handler(MessageHandler(filters.PHOTO, photo_router))
# Normal text goes to main answer
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, answer))
#sources handler
app.add_handler(CommandHandler("sources_all", sources_all))

print(
    "Bot is running with topic-scoped RAG + metadata separation (qa/evaluation) + submenus "
    "+ per-topic evaluation + embedded tools (brainstorm, rewrite, plan, recpacket, schoolfinder, portfolioideas) "
    "+ Application Plan & School Finder topics + analytics…"
)
nest_asyncio.apply()
asyncio.run(app.run_polling())
